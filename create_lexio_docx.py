#!/usr/bin/env python3
"""Generate LEXIO_MASTER_DOCUMENT.docx from scratch with full formatting."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color palette ──────────────────────────────────────────────────────────────
DEEP_NAVY   = RGBColor(0x1A, 0x23, 0x3A)   # headings
INDIGO      = RGBColor(0x37, 0x4A, 0xBC)   # H2 accent
TEAL        = RGBColor(0x00, 0x83, 0x80)   # H3 accent
GOLD        = RGBColor(0xD4, 0xA0, 0x17)   # callout border
LIGHT_BLUE  = RGBColor(0xE8, 0xF0, 0xFE)   # table header bg
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TEXT   = RGBColor(0x44, 0x44, 0x55)
DARK_ROW    = RGBColor(0xF3, 0xF4, 0xF6)   # alternate table row
GREEN_TICK  = RGBColor(0x16, 0xA3, 0x4A)
RED_CROSS   = RGBColor(0xDC, 0x26, 0x26)
WARN_YELLOW = RGBColor(0xD9, 0x77, 0x06)

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set cell background color ─────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── Helper: set cell borders ───────────────────────────────────────────────────
def set_table_borders(table, color="AAAAAA", sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    # remove existing
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

# ── Helper: paragraph spacing ─────────────────────────────────────────────────
def set_para_spacing(para, before=0, after=6, line_pts=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_pts:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_pts)

# ── Helper: run with font properties ──────────────────────────────────────────
def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font_name="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run

# ── Helper: add a styled heading ──────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=18, after=6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Calibri"
    run.font.color.rgb = WHITE
    # shading behind heading – use a 1×1 table trick
    # Actually set paragraph shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f"{DEEP_NAVY[0]:02X}{DEEP_NAVY[1]:02X}{DEEP_NAVY[2]:02X}")
    pPr.append(shd)
    return p

def h2(text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=14, after=4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    run.font.color.rgb = INDIGO
    # underline
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), f"{INDIGO[0]:02X}{INDIGO[1]:02X}{INDIGO[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h3(text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=10, after=3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Calibri"
    run.font.color.rgb = TEAL
    return p

def h4(text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.name = "Calibri"
    run.font.color.rgb = DEEP_NAVY
    return p

def body(text, indent=False, bold_parts=None):
    """bold_parts: list of substrings to bold."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    if bold_parts:
        parts = text
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx >= 0:
                add_run(p, remaining[:idx], color=GRAY_TEXT)
                add_run(p, bp, bold=True, color=DEEP_NAVY)
                remaining = remaining[idx + len(bp):]
        add_run(p, remaining, color=GRAY_TEXT)
    else:
        add_run(p, text, color=GRAY_TEXT)
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    set_para_spacing(p, before=1, after=3)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True, color=DEEP_NAVY)
        add_run(p, text[len(bold_prefix):], color=GRAY_TEXT)
    else:
        add_run(p, text, color=GRAY_TEXT)
    return p

def callout(text, style='info'):
    """Styled callout box as a 1-col table."""
    bg = RGBColor(0xEF, 0xF6, 0xFF) if style == 'info' else RGBColor(0xFFF, 0xBF, 0x00)
    border_color = "3B82F6" if style == 'info' else "F59E0B"
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    set_para_spacing(p, before=4, after=4)
    add_run(p, text, italic=True, color=DEEP_NAVY, size=11)
    tbl.columns[0].width = Inches(6)
    # custom border color
    set_table_borders(tbl, color=border_color, sz=6)
    doc.add_paragraph()  # spacer
    return tbl

def divider():
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_table(headers, rows, col_widths=None, header_bg=None):
    """Add a styled table with headers."""
    if header_bg is None:
        header_bg = LIGHT_BLUE
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="CBD5E1", sz=4)

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before=3, after=3)
        add_run(p, h, bold=True, color=DEEP_NAVY, size=10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg = DARK_ROW if r_idx % 2 == 0 else WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_spacing(p, before=2, after=2)
            # Color ✅ ❌ ⚠️
            txt = str(cell_text)
            if txt in ('✅',):
                add_run(p, txt, color=GREEN_TICK, size=11)
            elif txt in ('❌',):
                add_run(p, txt, color=RED_CROSS, size=11)
            elif txt in ('⚠️',):
                add_run(p, txt, color=WARN_YELLOW, size=11)
            elif c_idx == 0:
                add_run(p, txt, bold=True, color=DEEP_NAVY, size=10)
            else:
                add_run(p, txt, color=GRAY_TEXT, size=10)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacer
    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# ── Cover / Title ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=2)
add_run(p, "LEXIO", bold=True, size=40, color=DEEP_NAVY, font_name="Calibri")

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p2, before=0, after=4)
add_run(p2, "TÀI LIỆU GỐC TOÀN DIỆN", bold=True, size=18, color=INDIGO)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p3, before=0, after=2)
add_run(p3, "Nền tảng học tiếng Anh tích hợp CLIL · Cambridge A1→B1+ · Dành cho học sinh Việt Nam",
        italic=True, size=12, color=TEAL)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p4, before=2, after=12)
add_run(p4, "Phiên bản: 1.0  ·  Tháng 4/2026  ·  Bảo mật nội bộ", size=10, color=GRAY_TEXT)

divider()

callout("Tài liệu này là nguồn gốc cho toàn bộ nội dung marketing, bán hàng và truyền thông của Lexio.")

doc.add_page_break()

# ── MỤC LỤC ──────────────────────────────────────────────────────────────────
h1("MỤC LỤC")
doc.add_paragraph()
toc_items = [
    ("1.", "Tuyên ngôn định vị"),
    ("2.", "Vấn đề thị trường"),
    ("3.", "Triết lý & Nguyên lý giáo dục"),
    ("4.", "Mục tiêu đầu ra"),
    ("5.", "Cấu trúc chương trình 3 năm"),
    ("6.", "Kiến trúc ứng dụng — 10 Modules"),
    ("7.", "Flow học tập hàng tuần"),
    ("8.", "AI Tutor Nova — Gia sư riêng 24/7"),
    ("9.", "Hệ sinh thái giáo viên & phụ huynh"),
    ("10.", "Lợi thế cạnh tranh"),
    ("11.", "Người sáng lập & Thẩm quyền nội dung"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=2)
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, f"{num}  ", bold=True, color=INDIGO, size=11)
    add_run(p, title, color=DEEP_NAVY, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PHẦN I
# ══════════════════════════════════════════════════════════════════════════════
h1("PHẦN I: APP LÀ GÌ — VÀ TẠI SAO NÓ TỒN TẠI")
doc.add_paragraph()

# ── 1. Tuyên ngôn định vị ─────────────────────────────────────────────────────
h2("1. Tuyên ngôn định vị")

callout(
    "Lexio là chương trình học tiếng Anh tích hợp CLIL duy nhất tại Việt Nam được thiết kế "
    "để đưa học sinh từ A1 lên B1+ sau 3 năm — không chỉ học tiếng Anh như một môn học, "
    "mà học Toán, Khoa học, Tư duy bằng tiếng Anh."
)

body("Kết quả đầu ra được đo bằng năng lực cụ thể:")
for item in [
    "Đọc hiểu văn bản học thuật tiếng Anh",
    "Viết bài luận nghị luận 5 đoạn có cấu trúc",
    "Giải bài toán có lời văn hoàn toàn bằng tiếng Anh (Singapore Math)",
    "Tranh luận và bảo vệ quan điểm trước đám đông",
    "Đủ năng lực dùng tiếng Anh làm Medium of Instruction để học các chương trình giáo dục quốc tế (IELTS 5.0–5.5)",
]:
    bullet(item)

divider()

# ── 2. Vấn đề thị trường ──────────────────────────────────────────────────────
h2("2. Vấn đề thị trường")

h3("2.1 Nghịch lý của giáo dục tiếng Anh tại Việt Nam")
body(
    "Học sinh Việt Nam học tiếng Anh trung bình 8–10 năm nhưng vẫn không nói được một cuộc hội thoại đơn giản. "
    "Nguyên nhân không phải thiếu chăm chỉ — mà là học sai cách:",
    bold_parts=["8–10 năm", "học sai cách"]
)

add_table(
    headers=["Cách học phổ biến", "Kết quả"],
    rows=[
        ["Học ngữ pháp như quy tắc toán học", "Biết luật nhưng không dùng được"],
        ["Học từ vựng riêng lẻ không ngữ cảnh", "Quên sau 1 tuần"],
        ["Luyện đọc văn bản không liên quan thực tế", "Không chuyển hóa thành kỹ năng"],
        ["Không có cơ hội nói thật sự", "Câm tiếng Anh sau nhiều năm học"],
    ],
    col_widths=[3.2, 3.2],
)

h3("2.2 Vấn đề của các app hiện tại")
for item in [
    "Duolingo: Gamification tốt, nhưng không có lộ trình cấu trúc. Không dạy được viết, không CLIL, không kết nối với chương trình học.",
    "ELSA Speak: Chỉ tập trung phát âm, không có nội dung học thuật.",
    "Edupia / Monkey Junior: Hướng đến trẻ nhỏ, nội dung đơn giản, không có AI tutor thật sự.",
    "Các trung tâm: Tốn kém, phụ thuộc giáo viên, không có dữ liệu tiến độ.",
]:
    p = doc.add_paragraph(style='List Bullet')
    set_para_spacing(p, before=1, after=3)
    p.paragraph_format.left_indent = Cm(0.6)
    bold_end = item.index(":")
    add_run(p, item[:bold_end + 1], bold=True, color=DEEP_NAVY, size=11)
    add_run(p, item[bold_end + 1:], color=GRAY_TEXT, size=11)

h3("2.3 Khoảng trống thị trường")
body("Chưa có ứng dụng nào tại Việt Nam cung cấp đồng thời:", bold_parts=["Chưa có ứng dụng nào"])
for i, item in enumerate([
    "Lộ trình có cấu trúc từ A1 → B1+ (156 tuần, không phải 'học mãi không xong')",
    "Tích hợp CLIL — học Toán, Khoa học bằng tiếng Anh",
    "AI Tutor hội thoại thật sự (không phải chatbot câu hỏi-đáp)",
    "Hệ sinh thái đầy đủ: học sinh + giáo viên + phụ huynh",
    "Được xây bởi chuyên gia ESL 30 năm, không phải AI-generated",
], 1):
    p = doc.add_paragraph()
    set_para_spacing(p, before=1, after=3)
    p.paragraph_format.left_indent = Cm(0.6)
    add_run(p, f"{i}. ", bold=True, color=INDIGO, size=11)
    add_run(p, item, color=GRAY_TEXT, size=11)

divider()

# ── 3. Triết lý & Nguyên lý ───────────────────────────────────────────────────
h2("3. Triết lý & Nguyên lý giáo dục")

h3("3.1 Triết lý gốc: CLIL — Tiếng Anh là Phương tiện, không phải Mục tiêu")
body(
    "CLIL (Content and Language Integrated Learning) là phương pháp giảng dạy trong đó tiếng Anh không phải là môn học "
    "— mà là ngôn ngữ để học các môn khác. Học sinh học Toán, Khoa học, Xã hội học bằng tiếng Anh. "
    "Ngôn ngữ được tiếp thu tự nhiên qua việc dùng nó, không phải qua việc ghi nhớ nó.",
    bold_parts=["CLIL (Content and Language Integrated Learning)"]
)
callout('"We don\'t teach English. We teach through English."', style='info')

h3("3.2 Sáu nguyên lý cốt lõi")
principles = [
    ('① Chiến lược “Ngựa Gỗ” (Trojan Horse)',
     "Không dạy ngữ pháp khô khan dưới dạng quy tắc. Ngữ pháp được giấu trong chủ đề hấp dẫn. "
     "Học sinh dùng đúng vì ngữ cảnh tự nhiên yêu cầu — không phải vì nhớ luật."),
    ('② Nguyên tắc “Nấu Chậm” (Slow-Cook)',
     "Với học sinh mới bắt đầu (A1), một thì ngữ pháp cần ít nhất 4–6 tuần thẩm thấu trước khi chuyển sang cấu trúc mới. Không vội vã, không nhồi nhét."),
    ("③ Spaced Repetition System (SRS)",
     "Hệ thống ôn tập khoa học: từ mới được ôn lại theo chu kỳ tăng dần (1 ngày → 3 ngày → 7 ngày → 21 ngày). Não bộ ghi nhớ lâu dài, không bị quên theo kiểu 'học để thi'."),
    ("④ Scaffolding động — Giàn giáo rút dần",
     "Hỗ trợ tối đa ở giai đoạn đầu (hình ảnh, dịch song ngữ, câu mẫu), sau đó rút dần để học sinh tự chủ. Đến Phase 3, học sinh viết tự do và tranh luận không cần trợ giúp."),
    ("⑤ Production-Oriented — Sản xuất ngay từ ngày đầu",
     "Mỗi bài học kết thúc bằng một sản phẩm ngôn ngữ: một câu, một đoạn văn, một video nói. Không có khái niệm 'học trước, dùng sau'."),
    ("⑥ Zero-L1 trong giảng dạy",
     "Các khái niệm mới được giới thiệu qua hình ảnh và cử chỉ (Visual Anchor + Total Physical Response), không dùng tiếng Việt để giải thích."),
]
for title, desc in principles:
    p = doc.add_paragraph()
    set_para_spacing(p, before=5, after=2)
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, title, bold=True, color=INDIGO, size=11)
    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=0, after=5)
    p2.paragraph_format.left_indent = Cm(1.2)
    add_run(p2, desc, color=GRAY_TEXT, size=11)

h3("3.3 Chiến lược Song Sinh (Dual Mode)")
body("Mỗi tuần học có 2 phiên bản nội dung song song trên cùng một chủ đề:",
     bold_parts=["2 phiên bản nội dung song song"])

add_table(
    headers=["", "EASY MODE", "ADVANCED MODE"],
    rows=[
        ["Định vị", "Chuẩn kiến thức lớp học", "Mở rộng & thử thách"],
        ["Ngữ cảnh", "Cá nhân, cụ thể, xung quanh bản thân", "Toàn cầu, trừu tượng, học thuật"],
        ["Từ vựng", "Tier 1 & Basic Tier 2", "Tier 2 & 3 (học thuật, chuyên ngành)"],
        ["Ngữ pháp", "Câu đơn, câu ghép", "Câu phức, mệnh đề quan hệ, bị động"],
        ["Phù hợp", "Học sinh mới bắt đầu, lớp đại trà", "Học sinh khá-giỏi, có mục tiêu quốc tế"],
    ],
    col_widths=[1.5, 2.5, 2.5],
)

divider()

# ── 4. Mục tiêu đầu ra ────────────────────────────────────────────────────────
h2("4. Mục tiêu đầu ra")
h3("Sau 3 năm (156 tuần · ~624 giờ học có hướng dẫn)")
callout("Trình độ đầu ra: B1+ CEFR = IELTS 5.0–5.5")

add_table(
    headers=["Kỹ năng", "Năng lực cụ thể"],
    rows=[
        ["Đọc", "Hiểu văn bản học thuật, diễn giải dữ liệu, biểu đồ, bảng số liệu"],
        ["Viết", "Bài luận 5 đoạn (Narrative, Descriptive, Expository, Argumentative)"],
        ["Nghe", "Hiểu podcast, tài liệu giáo dục, giảng dạy bằng tiếng Anh"],
        ["Nói", "Thuyết trình, tranh luận, hội thoại học thuật tự nhiên"],
        ["Toán", "Giải bài toán có lời văn đa bước bằng tiếng Anh (Singapore Math)"],
        ["Tư duy", "Lý luận nhân quả, phân tích, so sánh, phản biện"],
    ],
    col_widths=[1.2, 5.2],
)

body("Cột mốc sử dụng thực tế:", bold_parts=["Cột mốc sử dụng thực tế:"])
for item in [
    "Sẵn sàng học các chương trình tiểu học/THCS của Mỹ, Úc, Singapore bằng tiếng Anh",
    "Đủ nền tảng để luyện IELTS trong 6–12 tháng để đạt 6.0+",
    "Có thể tham gia cuộc thi toán & khoa học bằng tiếng Anh",
]:
    bullet("✅  " + item)

divider()

# ── 5. Cấu trúc chương trình 3 năm ───────────────────────────────────────────
h2("5. Cấu trúc chương trình 3 năm")
h3("Tổng quan: 156 tuần · 3 giai đoạn")

# Overview table
add_table(
    headers=["Giai đoạn 1", "Giai đoạn 2", "Giai đoạn 3"],
    rows=[
        ["Tuần 1–42", "Tuần 43–112", "Tuần 113–156"],
        ["NỀN TẢNG", "ỨNG DỤNG HỌC THUẬT", "TỔNG HỢP & TRANH LUẬN"],
        ["A1 → A2", "A2 → B1", "B1 → B1+"],
    ],
    col_widths=[2.0, 2.5, 2.5],
)

# GIAI ĐOẠN 1
h3("GIAI ĐOẠN 1 — Nền tảng Trôi chảy (Tuần 1–42)")
body(
    "Mục tiêu: Xây dựng nền tảng từ vựng và ngữ pháp. Viết đoạn văn miêu tả, tường thuật, thủ tục. "
    "Tự tin nói những câu đơn giản.",
    bold_parts=["Mục tiêu:"]
)
body("Cấu trúc mỗi buổi học (120 phút):", bold_parts=["Cấu trúc mỗi buổi học (120 phút):"])
bullet("Session 1: 120 phút ELA (English Language Arts)")
bullet("Session 2: 120 phút ELA")

add_table(
    headers=["Chu trình", "Tuần", "Chủ đề"],
    rows=[
        ["1.1", "1–14", "Core Description — Mô tả người, địa điểm"],
        ["1.2", "15–28", "Recounting & Reasoning — Kể lại quá khứ, bày tỏ ý kiến"],
        ["1.3", "29–42", "Procedural & Expository — Hướng dẫn thủ tục, giải thích khoa học"],
    ],
    col_widths=[1.0, 1.0, 4.5],
)

body("Chủ đề ngữ pháp theo block:", bold_parts=["Chủ đề ngữ pháp theo block:"])
for b in [
    "Block A (W1–18): Subject Pronouns, To Be, Action Verbs, Present Simple/Continuous",
    "Block B (W19–36): Past Simple (Regular/Irregular), Like/Want/Need + V-ing/to-V",
    "Block C (W37–42): Review, project-based consolidation",
]:
    bullet(b)

# GIAI ĐOẠN 2
h3("GIAI ĐOẠN 2 — Ứng dụng Học thuật (Tuần 43–112)")
body(
    "Mục tiêu: Chuyển từ viết cá nhân sang viết học thuật. Dùng tiếng Anh để học và giải thích "
    "khái niệm Toán, Khoa học, Xã hội học.",
    bold_parts=["Mục tiêu:"]
)
body("Cấu trúc mỗi buổi học (120 phút):", bold_parts=["Cấu trúc mỗi buổi học (120 phút):"])
bullet("Session 1: 60 phút English Math + 60 phút ELA")
bullet("Session 2: 40 phút Science + 40 phút ELA (production) + 40 phút Social Studies")

add_table(
    headers=["Chu trình", "Tuần", "Trọng tâm"],
    rows=[
        ["2.1", "43–70", "Logic & Tư duy + Giới thiệu Toán tích hợp"],
        ["2.2", "71–98", "Học nội dung tích hợp toàn phần (CLIL đầy đủ)"],
        ["2.3", "99–112", '"My Passion Project" — Dự án nghiên cứu độc lập'],
    ],
    col_widths=[1.0, 1.0, 4.5],
)

body("Các đơn vị nội dung (Units):", bold_parts=["Các đơn vị nội dung (Units):"])
for u in [
    "Unit 5: Logical Thinking  ·  Unit 6: Communication & Data",
    "Unit 7: Persuasion  ·  Unit 8: Natural Systems",
    "Unit 9: Geography  ·  Unit 10: Advanced Science",
    "Unit 11: Government  ·  Unit 12: Basic Economics",
]:
    bullet(u)

# GIAI ĐOẠN 3
h3("GIAI ĐOẠN 3 — Tổng hợp & Tranh luận (Tuần 113–156)")
body(
    "Mục tiêu: Đạt tự chủ học thuật hoàn toàn. Nghiên cứu độc lập, viết bài luận nghị luận phức tạp, tranh luận chính thức.",
    bold_parts=["Mục tiêu:"]
)

add_table(
    headers=["Dự án", "Câu hỏi tranh luận"],
    rows=[
        ["Homework: Helper or Burden?", '"Should homework be banned in primary schools?"'],
        ["Video Games: More Than a Game?", '"Should students be allowed to play video games every day?"'],
        ["Smartphones for Kids", '"Should children under 12 have their own smartphones?"'],
    ],
    col_widths=[2.5, 4.0],
)

body("Cấu trúc mỗi chu trình dự án (8 tuần):", bold_parts=["Cấu trúc mỗi chu trình dự án (8 tuần):"])
for item in [
    "Tuần 1–2: Research + Thesis Writing",
    "Tuần 3–4: Drafting 5-paragraph essay",
    "Tuần 5–6: Editing + Counter-argument",
    "Tuần 7–8: Debate Preparation + Formal Debate",
]:
    bullet(item)
body("Kết thúc Giai đoạn 3 (Tuần 145–156):", bold_parts=["Kết thúc Giai đoạn 3 (Tuần 145–156):"])
body("Capstone presentation tổng kết toàn chương trình.", indent=True)

divider()

# ── 6. Kiến trúc ứng dụng ─────────────────────────────────────────────────────
h2("6. Kiến trúc ứng dụng — 10 Modules")
body("Mỗi tuần học trên app gồm 10 module (Tab) được thiết kế theo nguyên tắc: Input → Processing → Output.",
     bold_parts=["10 module (Tab)", "Input → Processing → Output"])

add_table(
    headers=["INPUT", "PROCESSING", "OUTPUT"],
    rows=[
        ["📖 Read & Explore", "🧠 Logic Lab", "✍️ Writing"],
        ["📝 New Words (10 từ)", "🎙️ Ask AI (Nova)", "🎤 Shadowing"],
        ["💪 Word Power", "🗺️ Mindmap Speaking", "🎮 Game Hub"],
        ["📺 Daily Watch (Video)", "🔁 Dictation", "📊 Progress Report"],
    ],
    col_widths=[2.0, 2.2, 2.0],
    header_bg=DEEP_NAVY,
)

modules = [
    ("MODULE 1 — Read & Explore", [
        ("Mục đích", "Cung cấp ngữ cảnh chính cho tuần học. Bài đọc là 'anchor' — neo tất cả ngữ pháp, từ vựng, và nội dung CLIL của tuần vào một bối cảnh nhất quán."),
        ("Easy Mode", "6–8 câu, ngữ pháp chuẩn syllabus"),
        ("Advanced Mode", "8–12 câu, ngữ pháp nâng cao (+1 level)"),
        ("Nội dung", "10 Bold Words bắt buộc in đậm; Audio giọng người thật, diễn cảm"),
        ("UX", "Bấm từ in đậm → Popup (Hình + Nghĩa + Âm thanh). Dịch từng câu riêng lẻ để tránh học vẹt."),
    ]),
    ("MODULE 2 — New Words (10 từ vựng cốt lõi)", [
        ("Mục đích", "Dạy sâu 10 từ vựng cốt lõi của tuần theo quy trình 3 bước."),
        ("Bước 1 — Listen & Type", "Nghe và gõ lại (âm thanh → văn bản)"),
        ("Bước 2 — Meaning Match", "Chọn nghĩa đúng (nhận diện)"),
        ("Bước 3 — Context", "Điền từ vào câu ví dụ (sản xuất có kiểm soát)"),
        ("Phase 1", "Từ đơn + Hình ảnh cụ thể (Concrete Nouns/Verbs)"),
        ("Phase 2", "Từ trừu tượng + Định nghĩa đơn giản"),
        ("Phase 3", "Thuật ngữ chuyên ngành (Terminology) + Định nghĩa học thuật"),
    ]),
    ("MODULE 3 — Word Power (Từ vựng Mở rộng)", [
        ("Mục đích", "Mở rộng vốn từ vựng theo chiều sâu (collocations, synonyms, idioms)."),
        ("Phase 1 (3 từ/tuần)", "Collocations đơn giản (ride a bike, take a bath)"),
        ("Phase 2 (5 từ/tuần)", "Synonyms & Antonyms (big = huge; small ↔ tiny)"),
        ("Phase 3 (7 từ/tuần)", "Idioms & Phrasal Verbs (figure out, point out)"),
        ("Quy tắc", "Mỗi từ bắt buộc đi kèm câu ví dụ đầy đủ trong ngữ cảnh thực tế."),
    ]),
    ("MODULE 4 — Daily Watch (Video tuyển chọn)", [
        ("Mục đích", "Nghe tiếng Anh tự nhiên, xây dựng listening comprehension, cung cấp thêm context trực quan."),
        ("Phase 1", "Hoạt hình giáo dục (Peppa Pig, Numberblocks, English Singsing, Cocomelon)"),
        ("Phase 2", "Tài liệu khoa học cho trẻ (SciShow Kids, NatGeo Kids, Khan Academy Kids)"),
        ("Phase 3", "Debate clips, TED-Ed, tin tức thực tế phù hợp lứa tuổi"),
    ]),
    ("MODULE 5 — Logic Lab (Toán & Tư duy)", [
        ("Mục đích", "Dạy Toán và Tư duy khoa học bằng tiếng Anh — thể hiện rõ nhất triết lý CLIL."),
        ("Sub-tab A: Logic & Science", "Patterns, deductive reasoning, khái niệm khoa học, phân loại, câu hỏi YES/NO logic"),
        ("Sub-tab B: Singapore Math", "Part-Whole, Comparison, Missing Part, Groups, Before-After — có Bar Model diagram"),
        ("Phase 1 (W1–15)", "Toán ngôn ngữ cơ bản (hình, số, phép tính)"),
        ("Phase 1+ (W16–54)", "STEM tích hợp bắt đầu (vật lý, sinh học, sinh thái)"),
        ("Phase 2 (W55–120)", "Toán ứng dụng + Pre-algebra (phương trình, tỉ số, phân số)"),
        ("Phase 3 (W121+)", "Cạnh tranh (đại số 2 biến, hình học tổng hợp, thống kê)"),
    ]),
    ("MODULE 6 — Ask AI / Nova (Luyện nói & Hỏi đáp)", [
        ("Mục đích", "Luyện kỹ năng đặt câu hỏi và hội thoại với AI theo ngữ cảnh có chủ đích."),
        ("Cơ chế", "Voice-First (mặc định microphone). Tối thiểu 5 prompts/tuần."),
        ("Phase 1 — Shadow Asking", "Nhái lại câu hỏi mẫu, luyện ngữ điệu"),
        ("Phase 2 — Guided Asking", "Tự lắp ghép câu từ gợi ý ngữ cảnh + STEM Reasoning"),
        ("Phase 3 — Free Inquiry", "Hỏi và tranh luận tự do, đánh giá độ phức tạp câu hỏi"),
    ]),
    ("MODULE 7 — Shadowing (Luyện nói đuổi)", [
        ("Mục đích", "Luyện phát âm, ngữ điệu, và tốc độ nói bằng phương pháp shadowing chuyên nghiệp."),
        ("Cơ chế", "Nghe đoạn audio từ bài đọc → nói đuổi theo đồng thời → app ghi âm → so sánh với bản gốc."),
        ("Hiệu quả", "Shadowing buộc não bộ xử lý cả âm thanh lẫn nghĩa cùng lúc — phát âm tự nhiên, không 'cứng'."),
    ]),
    ("MODULE 8 — Mindmap Speaking (Nói có cấu trúc)", [
        ("Mục đích", "Luyện sản xuất câu hoàn chỉnh và có cấu trúc thông qua bản đồ tư duy tương tác."),
        ("Cơ chế", "Chủ đề trung tâm → 4–6 nhánh (phrase gợi ý) → chọn nhánh → tạo câu hoàn chỉnh → ghi âm."),
    ]),
    ("MODULE 9 — Writing (Viết có hướng dẫn)", [
        ("Mục đích", "Luyện viết theo lộ trình từ câu đơn đến bài luận 5 đoạn."),
        ("Phase 1", "Photo Submission (viết giấy chụp ảnh) + Sentence Frames (điền vào)"),
        ("Phase 2", "Dictation Mode (nói → ghi) + Sentence Starters + Outlining (kéo thả)"),
        ("Phase 3", "Free Writing với rubric đánh giá, peer review"),
    ]),
    ("MODULE 10 — Game Hub + Dictation", [
        ("Game Hub", "Mini-games ôn luyện từ vựng và ngữ pháp trong môi trường áp lực thấp (word scramble, quick quiz, matching)."),
        ("Dictation — 3 cấp độ",
         "Nghe câu → viết lại theo 3 cấp độ tăng dần. Nội dung lấy từ bài đọc Read & Explore, "
         "tách thành câu riêng lẻ. "
         "Level 1 — Unscramble: xáo trộn từ → sắp xếp lại. "
         "Level 2 — Fill in the Blanks: điền từ còn thiếu. "
         "Level 3 — Full Dictation: nghe audio → gõ lại toàn bộ câu.\n"
         "⚠️ Lưu ý: Hiện tại chưa tích hợp SRS — câu được chọn từ bài đọc tuần, không phải từ kho từ vựng đã học theo chu kỳ."),
    ]),
]

for mod_title, fields in modules:
    h3(mod_title)
    add_table(
        headers=["Thành phần", "Chi tiết"],
        rows=fields,
        col_widths=[2.0, 4.5],
    )

divider()

# ── 7. Flow học tập ───────────────────────────────────────────────────────────
h2("7. Flow học tập hàng tuần")

h3("Vòng lặp chuẩn trong 1 tuần (Weekly Loop)")
add_table(
    headers=["TRƯỚC LỚP (Pre-class)", "TRÊN LỚP (In-class)", "SAU LỚP (Post-class)"],
    rows=[[
        "📖 Read & Explore\n📝 New Words (drill 3 bước)\n📺 Daily Watch",
        "Giáo viên dạy với lesson plan tích hợp nội dung app (CLIL + ngữ pháp)",
        "📝 Dictation (SRS ôn từ)\n🎮 Game Hub (ôn luyện)\n✍️ Writing (bài tập nhà)\n🎙️ Shadowing (phát âm)\n💬 Ask AI (hội thoại thêm)",
    ]],
    col_widths=[2.1, 2.3, 2.1],
)

h3("Vòng lặp SRS (Spaced Repetition)")
add_table(
    headers=["Thời điểm", "Hoạt động", "Kết quả"],
    rows=[
        ["Ngày 1", "Học lần đầu", "Ghi nhớ ngắn hạn"],
        ["Ngày 2", "Ôn lần 1 (Dictation / Quiz)", "Củng cố ban đầu"],
        ["Ngày 5", "Ôn lần 2", "Chuyển sang trung hạn"],
        ["Ngày 14", "Ôn lần 3", "Ghi nhớ trung hạn"],
        ["Ngày 30", "Ôn cuối", "Ghi nhớ dài hạn"],
    ],
    col_widths=[1.3, 2.5, 2.7],
)
callout("Retention rate >85% sau 3 tháng, so với ~20% của phương pháp truyền thống.")

divider()

# ── 8. AI Tutor Nova ──────────────────────────────────────────────────────────
h2("8. AI Tutor Nova — Gia sư riêng 24/7")

h3("Vai trò")
body(
    "Nova không phải chatbot câu hỏi-đáp. Nova là gia sư riêng thích ứng, hiểu ngữ cảnh tuần học hiện tại, "
    "lịch sử học của học sinh, và điều chỉnh cách tương tác theo trình độ.",
    bold_parts=["gia sư riêng thích ứng"]
)

h3("4 tab của Nova (hiện tại)")
nova_functions = [
    ("① Story Mission — Nhiệm vụ câu chuyện",
     "Nova dẫn dắt học sinh qua một câu chuyện gắn với chủ đề tuần — có nhân vật, bối cảnh, nhiệm vụ cụ thể. "
     "Học sinh đọc hiểu, trả lời câu hỏi của Nova (bằng giọng nói hoặc gõ), hoàn thành nhiệm vụ để nhận XP. "
     "Scaffolding: câu hỏi Yes/No → open-ended theo tiến độ; hệ thống hints tự động khi học sinh bị kẹt."),
    ("② Free Talk — Hội thoại có kiểm soát (Conversation Cards)",
     "Không phải chat tự do. Free Talk được điều khiển bởi Conversation Cards: các thẻ hội thoại "
     "theo chủ đề cụ thể từng tuần (gặp bạn lớp, giới thiệu gia đình, đặt đồ ăn...). Nova đóng vai và "
     "dẫn dắt theo kịch bản thẻ. Ngoài ra có Translation Help (dịch tức thì + spelling) "
     "và Ask Anything (hỏi về bài học tuần). Scaffolding: suggested hints sau mỗi lượt, giảm dần."),
    ("③ Pronunciation — Phát âm + Sentence Shadowing",
     "Hai chế độ: (1) Word Pronunciation — chọn từ → nghe mẫu → ghi âm → AI chấm từng âm vị, "
     "chỉ ra âm nào sai và cách sửa. "
     "(2) Sentence Shadowing — nghe câu hoàn chỉnh từ bài đọc hoặc syllabus ngữ pháp → nói đuổi theo → "
     "AI so sánh ngữ điệu, tốc độ, độ chính xác. "
     "Scaffolding: Phase 1 — shadowing câu ngắn 4–6 từ; Phase 3 — đoạn thoại tự nhiên."),
    ("④ Debate — Tranh luận với AI (Mở khóa từ Tuần 40)",
     "Học sinh chọn chủ đề, đứng một phía → Nova phản biện. Nova calibrate độ mạnh phản biện theo trình độ. "
     "Scaffolding: cung cấp argument starters ở giai đoạn đầu "
     "('I believe...', 'The main reason is...'), rút dần khi học sinh quen cấu trúc."),
]
for title, desc in nova_functions:
    h4(title)
    body(desc, indent=True)

h3("Tính năng đề xuất bổ sung: Story Builder (Tab thứ 5)")
body(
    "Cơ chế đề xuất: Nova và học sinh cùng viết một câu chuyện sáng tạo, câu nối câu. "
    "Nova KHÔNG viết thay học sinh — Nova chỉ đặt câu hỏi gợi ý để học sinh tự tạo nội dung tiếp theo. "
    "Kết quả: học sinh có câu chuyện hoàn chỉnh do chính mình tạo ra — portfolio sáng tạo thực sự.",
    bold_parts=["KHÔNG viết thay"]
)
body("Ví dụ flow:", bold_parts=["Ví dụ flow:"])
for item in [
    'Nova: "Nhân vật của chúng ta tên là gì? Hãy mô tả bạn ấy trong 1 câu."',
    'Học sinh: "His name is Leo. He is brave and funny."',
    'Nova: "Great! Where does Leo live? Write 1 sentence about his home."',
    'Học sinh: "Leo lives in a treehouse in the jungle." → câu chuyện tích lũy dần...',
]:
    bullet(item)

body("Phân tích: Có nên thêm Story Builder?", bold_parts=["Phân tích: Có nên thêm Story Builder?"])
pros = [
    "Output cao nhất trong thang Bloom: học sinh sáng tác tác phẩm liên tục, đòi hỏi coherence và narrative logic — không module nào hiện tại đạt tới tầng này.",
    "Portfolio vật chất: câu chuyện có thể lưu, in, chia sẻ với phụ huynh — tạo cảm giác thành tựu rõ ràng hơn điểm số.",
    "Nova giữ vai trò pure scaffolding (chỉ đặt câu hỏi, không viết thay) → hoàn toàn phù hợp nguyên lý Production-Oriented của Lexio.",
    "Phân biệt rõ với Story Mission: Story Mission = đọc hiểu câu chuyện của Nova. Story Builder = học sinh viết câu chuyện của chính mình.",
]
body("✅ Lý do ủng hộ:", bold_parts=["✅ Lý do ủng hộ:"])
for p in pros:
    bullet(p)
challenges = [
    "Đòi hỏi Nova giữ 'story memory' — nhớ toàn bộ câu đã viết để đảm bảo continuity. Với LLM context window ngắn, đây là thách thức kỹ thuật thực tế.",
    "Phase 1 học sinh chưa đủ từ vựng viết sáng tạo tự do — cần word bank và sentence frames mạnh ở giai đoạn đầu, nếu không Story Builder gây frustration thay vì motivation.",
    "Tốt nhất mở khóa từ Phase 1+ (Tuần 16+) khi học sinh đã có >100 từ vựng tích lũy. Phase 3 mới cho viết hoàn toàn tự do.",
    "Cần giải quyết câu hỏi: story được lưu ở đâu và hiển thị thế nào cho phụ huynh? Nếu không có UI lưu/xem lại, giá trị portfolio bị mất.",
]
body("⚠️ Thách thức cần cân nhắc:", bold_parts=["⚠️ Thách thức cần cân nhắc:"])
for c in challenges:
    bullet(c)

h3("Nova hiểu ngữ cảnh")
body("Nova biết:")
for item in [
    "Học sinh đang ở tuần nào, giai đoạn nào",
    "Từ vựng nào đã học, từ nào còn yếu (SRS data)",
    "Chủ đề tuần hiện tại để giữ hội thoại liên quan",
    "Lịch sử tương tác để không lặp lại",
]:
    bullet(item)

divider()

# ── 9. Hệ sinh thái giáo viên & phụ huynh ────────────────────────────────────
h2("9. Hệ sinh thái giáo viên & phụ huynh")

h3("Cho giáo viên — Teacher Panel")
add_table(
    headers=["Tính năng", "Mô tả"],
    rows=[
        ["Lesson Plans tích hợp", "3 session plan/tuần đầy đủ chi tiết. In được: layout chuẩn, có diagram SVG, có đánh số câu hỏi."],
        ["Periodic Quiz Generator", "Tạo đề kiểm tra theo khoảng tuần với 5 loại câu hỏi (multiple choice, fill-blank, matching, error correction, writing)"],
        ["WarmUp Quiz", "Đề ôn nhanh đầu buổi (5–10 câu)"],
        ["Parent Quiz", "Đề kiểm tra cho phụ huynh in về nhà"],
        ["Ghi chú học sinh", "Giáo viên ghi chú quan sát từng học sinh từng buổi — lưu lại làm basis cho báo cáo."],
    ],
    col_widths=[2.2, 4.3],
)

h3("Cho phụ huynh — Parent Dashboard")
add_table(
    headers=["Tính năng", "Mô tả"],
    rows=[
        ["Retention Rate từ vựng", "Số từ vựng đã học và tỉ lệ ghi nhớ thực tế"],
        ["Phân tích kỹ năng", "Kỹ năng nào mạnh, kỹ năng nào cần củng cố"],
        ["So sánh với mục tiêu", "So sánh với mục tiêu tuần / giai đoạn"],
        ["Gợi ý hỗ trợ", "Gợi ý hoạt động hỗ trợ tại nhà cụ thể"],
        ["Minh bạch dữ liệu", "Báo cáo dữ liệu thật — không phải 'con học tốt' mà là số liệu cụ thể"],
    ],
    col_widths=[2.2, 4.3],
)

divider()

# ── 10. Lợi thế cạnh tranh ────────────────────────────────────────────────────
h2("10. Lợi thế cạnh tranh")

h3("Ma trận so sánh")
add_table(
    headers=["Tính năng", "Lexio", "Duolingo", "Edupia", "ELSA", "Trung tâm"],
    rows=[
        ["CLIL — học môn học bằng tiếng Anh", "✅", "❌", "❌", "❌", "❌"],
        ["Lộ trình cấu trúc A1→B1+", "✅", "❌", "❌", "❌", "⚠️"],
        ["AI Tutor hội thoại thật sự", "✅", "❌", "❌", "❌", "❌"],
        ["Spaced Repetition (SRS)", "✅", "✅", "❌", "❌", "❌"],
        ["Chuẩn Cambridge (không SGK)", "✅", "❌", "❌", "❌", "⚠️"],
        ["Shadowing & phát âm AI", "✅", "❌", "✅", "✅", "❌"],
        ["Viết câu / sửa bài tức thì", "✅", "❌", "❌", "❌", "✅"],
        ["Dashboard & báo cáo phụ huynh", "✅", "❌", "✅", "❌", "❌"],
        ["Lesson Plans cho giáo viên", "✅", "❌", "❌", "❌", "✅"],
        ["Tạo đề kiểm tra (GV & PH)", "✅", "❌", "❌", "❌", "❌"],
        ["Singapore Math bằng tiếng Anh", "✅", "❌", "❌", "❌", "❌"],
        ["Nội dung chuẩn hóa HS Việt", "✅", "❌", "✅", "❌", "⚠️"],
        ["Game + Logic Lab tư duy", "✅", "✅", "✅", "❌", "❌"],
    ],
    col_widths=[3.0, 0.8, 0.9, 0.8, 0.7, 1.0],
)

h3("3 lợi thế không thể copy nhanh")
advantages = [
    ("① Chiều sâu nội dung — 156 tuần được viết tay bởi chuyên gia",
     "Không có AI nào tạo ra được 156 tuần nội dung tích hợp nhất quán theo lộ trình CLIL. "
     "Đây là công trình 30 năm kinh nghiệm được hệ thống hóa. Đối thủ muốn copy cần ít nhất 3–5 năm."),
    ("② Tích hợp hệ sinh thái: Học sinh + Giáo viên + Phụ huynh",
     "Không phải app tự học đơn lẻ. Lexio là hạ tầng vận hành cho cả lớp học. Ba bên kết nối "
     "— tạo vòng lặp accountability không thể có ở app 'học một mình'."),
    ("③ Mục tiêu đầu ra rõ ràng và đo được",
     "Lexio cam kết: B1+ sau 3 năm, đủ dùng tiếng Anh làm Medium of Instruction. "
     "Đây là tuyên bố mà không app nào dám đưa ra — vì không ai có lộ trình đủ cụ thể để đứng sau nó."),
]
for title, desc in advantages:
    h4(title)
    body(desc, indent=True)

divider()

# ── 11. Người sáng lập ────────────────────────────────────────────────────────
h2("11. Người sáng lập & Thẩm quyền nội dung")

body(
    "Lexio không phải 'AI-generated content'. Toàn bộ 156 tuần nội dung — từng bài đọc, từng câu hỏi logic, "
    "từng scenario hội thoại — được trực tiếp thiết kế bởi người sáng lập:",
    bold_parts=["Lexio không phải 'AI-generated content'."]
)

for item in [
    "30 năm kinh nghiệm dạy và nghiên cứu ESL tại Việt Nam",
    "Phiên dịch viên cabin hàng đầu — người duy nhất có thể làm phiên dịch ở cấp độ này hiểu cả hai ngôn ngữ ở mức native",
    "Là công trình cả đời — không phải sản phẩm thương mại hóa vội vã",
    "Xây dựng theo chuẩn Cambridge & phương pháp ESL tiên tiến nhất (CLIL, SRS, TPR, Scaffolding)",
]:
    bullet(item)

doc.add_paragraph()
callout(
    '"Lexio là câu trả lời cho câu hỏi: \'Nếu tôi có thể dạy tiếng Anh đúng cách cho mọi trẻ em '
    'Việt Nam, tôi sẽ dạy như thế nào?\'"'
)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "Production_FINAL/Marketing/LEXIO_MASTER_DOCUMENT.docx"
doc.save(out_path)
print(f"✅ Saved: {out_path}")
