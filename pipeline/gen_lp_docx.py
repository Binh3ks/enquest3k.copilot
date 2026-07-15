#!/usr/bin/env python3
"""
pipeline/gen_lp_docx.py
════════════════════════════════════════════════════════════════════════════════
Generate W{N}_Lesson_Plan.docx from the W{N}.json lesson plan data.

Adheres to the W28 golden standard structure:
  Cover → Section 1 (Quick Ref) → Section 2 (Methodology) →
  Section 3 (Vocab) → Section 4-6 (Session Plans) →
  Section 7 (Teacher Materials: Listening + Speaking) →
  Section 8 (Answer Key) → Section 9 (Task Cards)

Used by:
  • pipeline/build_from_docx.py  (W01-53, after DOCX → JSON step)
  • pipeline/generate_ai_week.py (W54+, after AI generation step)

Usage (standalone):
  python3 pipeline/gen_lp_docx.py 37            # single week
  python3 pipeline/gen_lp_docx.py 37 38 39      # multiple weeks
  python3 pipeline/gen_lp_docx.py 37-53         # range
  python3 pipeline/gen_lp_docx.py --all         # all existing W{N}.json
  python3 pipeline/gen_lp_docx.py --out-dir /path/to/dir 37  # custom output dir

Output:
  Production_FINAL/1. FINAL MASS PRODUCTION/2_REFERENCE_DOCS/lesson_plans/output/
    W{N}_Lesson_Plan.docx
"""

import sys, json, re
from pathlib import Path

ROOT = Path(__file__).parent.parent

# Source JSON directories (public is primary; mcp-server as fallback)
PUB_LESSONS = ROOT / 'public/data/lessons'
MCP_LESSONS = ROOT / 'mcp-server/data/lessons'

# Default output directory (mirrors existing W18-W36 lesson plans)
DEFAULT_OUT_DIR = (
    ROOT / 'Production_FINAL/1. FINAL MASS PRODUCTION/2_REFERENCE_DOCS/lesson_plans/output'
)

# ── python-docx import ────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


# ── Separator constants ───────────────────────────────────────────────────────
SEP  = '═' * 55
SEP2 = '─' * 55


# ─────────────────────────────────────────────────────────────────────────────
# Low-level helpers
# ─────────────────────────────────────────────────────────────────────────────

class LPDocBuilder:
    """Lightweight document builder matching lp_doc_builder.py style."""

    def __init__(self):
        if not _DOCX_OK:
            raise ImportError('python-docx is required: pip install python-docx')
        self.doc = Document()
        self._setup()

    def _setup(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        for sec in self.doc.sections:
            sec.top_margin    = Inches(0.75)
            sec.bottom_margin = Inches(0.75)
            sec.left_margin   = Inches(1.0)
            sec.right_margin  = Inches(1.0)

    # ── Structural ────────────────────────────────────────────────────────────
    def sep(self):      self.doc.add_paragraph(SEP)
    def sep2(self):     self.doc.add_paragraph(SEP2)
    def blank(self):    self.doc.add_paragraph('')
    def page_break(self): self.doc.add_page_break()

    # ── Headers ───────────────────────────────────────────────────────────────
    def h1(self, text):
        """Bold centered — cover title."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = True; r.font.size = Pt(13)
        return p

    def h2(self, text):
        """Bold uppercase — section header."""
        p = self.doc.add_paragraph()
        r = p.add_run(text); r.bold = True; r.font.size = Pt(12)
        return p

    def h3(self, text):
        """Bold — subsection."""
        p = self.doc.add_paragraph()
        r = p.add_run(text); r.bold = True
        return p

    def h4(self, text):
        """Bold italic — minor subhead."""
        p = self.doc.add_paragraph()
        r = p.add_run(text); r.bold = True; r.italic = True
        return p

    # ── Body ──────────────────────────────────────────────────────────────────
    def body(self, text):
        return self.doc.add_paragraph(str(text))

    def indent(self, text, level=1):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3 * level)
        p.add_run(str(text))
        return p

    def indent_bold(self, text, level=1):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3 * level)
        r = p.add_run(str(text)); r.bold = True
        return p

    def field_row(self, label, value):
        """Render a 'Label:   value' line in body style."""
        p = self.doc.add_paragraph()
        r1 = p.add_run(label.rstrip(':') + ':  ')
        r1.bold = True
        p.add_run(str(value))
        return p

    def labeled_field(self, label, blank_len=30):
        """Render 'Label: ___________' (for teacher fill-in fields)."""
        self.body(label + ' ' + '_' * blank_len)

    def save(self, path):
        self.doc.save(str(path))


# ─────────────────────────────────────────────────────────────────────────────
# Section builders (W28 golden standard layout)
# ─────────────────────────────────────────────────────────────────────────────

def _build_cover(d: LPDocBuilder, data: dict):
    week = data.get('week', '?')
    qr   = data.get('quick_ref', {})
    block = qr.get('Block', 'C')
    theme = qr.get('Theme', data.get('unit_theme', ''))

    d.sep()
    d.h1('TEACHER CONTENT PACK — WEEK %s  |  BLOCK %s' % (week, block))
    d.h1(theme)
    d.h1('Integrated English Programme  |  v3.0')
    d.sep()
    d.blank()
    d.labeled_field('Teacher name:', 30)
    d.labeled_field('Class:', 30)
    d.labeled_field('Week dates (Mon–Sat):', 22)
    d.blank()


def _build_section1(d: LPDocBuilder, data: dict):
    week = data.get('week', '?')
    qr   = data.get('quick_ref', {})

    d.h2('SECTION 1: QUICK REFERENCE TABLE')
    d.sep2()
    for key, val in qr.items():
        d.field_row(key, val)
    d.body('Session structure:   S1 = Reading + Vocab (50 min)  |  S2 = Grammar + Speaking (55 min)  |  S3 = Writing (45 min)')
    d.blank()


def _build_section2(d: LPDocBuilder, data: dict):
    methodology = data.get('methodology', [])
    if not methodology:
        return

    d.h2('SECTION 2: TEACHING METHODOLOGY')
    d.sep2()
    for section in methodology:
        title   = section.get('title', '')
        content = section.get('content', [])
        if title:
            d.h3(title)
        for line in content:
            if not line:
                d.blank()
            elif line.startswith('  ') or line.startswith('\t'):
                d.indent(line.strip(), level=2)
            else:
                d.indent(line)
    d.blank()


def _build_section3(d: LPDocBuilder, data: dict):
    vocab = data.get('vocab_tiers', [])
    if not vocab:
        return

    d.h2('SECTION 3: VOCABULARY MAP + SRS SCHEDULE')
    d.sep2()
    d.h3('Vocabulary Words — W28 Golden Standard (4-field schema)')
    d.blank()
    for v in vocab:
        word   = v.get('Word', '')
        viet   = v.get('Vietnamese', '')
        colloc = v.get('Key Collocation(s)', '')
        trick  = v.get('Memory Trick', '')
        if not word:
            continue
        d.h4('%s  (%s)' % (word, viet))
        if colloc:
            d.indent('Key Collocation: %s' % colloc)
        if trick:
            d.indent('Memory Trick:    %s' % trick)
        d.blank()


def _build_session_sections(d: LPDocBuilder, data: dict):
    sessions = data.get('sessions', [])
    if not sessions:
        return

    for sess in sessions:
        snum  = sess.get('session', '?')
        parts = sess.get('parts', [])

        d.page_break()
        d.h2('SECTION %d: SESSION %s — STUDENT WORKSHEET' % (3 + snum, snum))
        d.sep2()
        d.labeled_field('Name:', 30)
        d.labeled_field('Date:', 30)
        d.blank()

        for part in parts:
            title   = part.get('title', '')
            content = part.get('content', [])

            if title:
                d.h3(title)
            for line in content:
                if not line:
                    d.blank()
                elif line.startswith('    ') or line.startswith('\t\t'):
                    d.indent(line.strip(), level=2)
                elif line.startswith('  ') or line.startswith('\t'):
                    d.indent(line.strip())
                else:
                    d.body(line)
        d.blank()


def _build_section_teacher(d: LPDocBuilder, data: dict):
    tc_list = data.get('teacher_contents', [])
    if not tc_list:
        return

    d.page_break()
    d.h2('TEACHER MATERIALS (CONFIDENTIAL — DO NOT DISTRIBUTE TO STUDENTS)')
    d.sep2()

    for tc in tc_list:
        snum = tc.get('session', '?')
        d.h3('─── SESSION %s ─────────────────────────────────────' % snum)

        # Listening script
        ls = tc.get('listening_script', {})
        ls_text    = ls.get('text', '') if isinstance(ls, dict) else str(ls)
        ls_speed   = ls.get('speed_note', '') if isinstance(ls, dict) else ''
        ls_dictation = ls.get('dictation', []) if isinstance(ls, dict) else []

        if ls_text or ls_speed or ls_dictation:
            d.h4('Listening Script — Session %s' % snum)
            if ls_speed:
                d.indent('Speed note: %s' % ls_speed)
            if ls_text:
                d.blank()
                for line in ls_text.split('\n'):
                    if line.strip():
                        d.indent(line.strip())
            if ls_dictation:
                d.blank()
                d.indent_bold('Dictation sentences:')
                for i, sent in enumerate(ls_dictation, 1):
                    d.indent('%d. %s' % (i, sent), level=2)
        else:
            d.indent('[Listening script: see reference DOCX or generate via AI pipeline]')

        d.blank()

        # Speaking notes
        sp = tc.get('speaking_notes', '')
        if sp:
            d.h4('Speaking Notes — Session %s' % snum)
            for line in str(sp).split('\n'):
                if line.strip():
                    d.indent(line.strip())
            d.blank()

        # STEM extension
        stem = tc.get('stem_extension', '')
        if stem:
            d.h4('STEM Extension — Session %s' % snum)
            for line in str(stem).split('\n'):
                if line.strip():
                    d.indent(line.strip())
            d.blank()

        # In-class speaking
        ics = tc.get('in_class_speaking', '')
        if ics:
            d.h4('In-Class Speaking Activities — Session %s' % snum)
            for line in str(ics).split('\n'):
                if line.strip():
                    d.indent(line.strip())
            d.blank()

        # VC answer key
        vc = tc.get('vc_answer_key', '')
        if vc:
            d.h4('VC Answer Key — Session %s' % snum)
            d.indent(str(vc))
            d.blank()


def _build_section_answer_key(d: LPDocBuilder, data: dict):
    ak = data.get('answer_key_by_session', {})
    ak_full = data.get('answer_key', [])
    if not ak and not ak_full:
        return

    d.page_break()
    d.h2('ANSWER KEY (TEACHER USE ONLY)')
    d.sep2()

    if ak:
        for sess_key in ('s1', 's2', 's3'):
            session_ak = ak.get(sess_key, [])
            if session_ak:
                d.h3('Session %s' % sess_key.upper())
                for entry in session_ak:
                    if isinstance(entry, dict):
                        part  = entry.get('part', '')
                        items = entry.get('answers', entry.get('items', []))
                        if part:
                            d.h4(part)
                        if isinstance(items, list):
                            for item in items:
                                d.indent(str(item))
                        elif items:
                            d.indent(str(items))
                    else:
                        d.indent(str(entry))
                d.blank()
    elif ak_full:
        for entry in ak_full:
            if isinstance(entry, dict):
                part  = entry.get('part', '')
                items = entry.get('answers', entry.get('items', []))
                if part:
                    d.h3(part)
                if isinstance(items, list):
                    for item in items:
                        d.indent(str(item))
                elif items:
                    d.indent(str(items))
            else:
                d.indent(str(entry))
        d.blank()


def _build_section_task_cards(d: LPDocBuilder, data: dict):
    tc_by_sess = data.get('task_cards_by_session', {})
    tc_all     = data.get('task_cards', [])
    if not tc_by_sess and not tc_all:
        return

    d.page_break()
    d.h2('TASK CARDS')
    d.sep2()

    source = tc_by_sess or {}
    if source:
        for sess_key in ('1', '2', '3'):
            cards = source.get(sess_key, [])
            if cards:
                d.h3('Session %s Task Cards' % sess_key)
                for card in cards:
                    if isinstance(card, dict):
                        title   = card.get('title', card.get('card', ''))
                        content = card.get('content', card.get('text', ''))
                        if title:
                            d.h4(title)
                        if isinstance(content, list):
                            for line in content:
                                d.indent(str(line))
                        elif content:
                            d.indent(str(content))
                    else:
                        d.indent(str(card))
                d.blank()
    else:
        for card in tc_all:
            if isinstance(card, dict):
                title   = card.get('title', card.get('card', ''))
                content = card.get('content', card.get('text', ''))
                if title:
                    d.h3(title)
                if isinstance(content, list):
                    for line in content:
                        d.indent(str(line))
                elif content:
                    d.indent(str(content))
            else:
                d.indent(str(card))
        d.blank()


def _build_section_games(d: LPDocBuilder, data: dict):
    games = data.get('games', [])
    if not games:
        return

    d.page_break()
    d.h2('GAMES & ACTIVITIES')
    d.sep2()
    for game in games:
        if isinstance(game, dict):
            name  = game.get('name', game.get('title', ''))
            desc  = game.get('description', game.get('content', ''))
            if name:
                d.h3(name)
            if isinstance(desc, list):
                for line in desc:
                    d.indent(str(line))
            elif desc:
                d.indent(str(desc))
        else:
            d.indent(str(game))
    d.blank()


def _build_section_sessions5(d: LPDocBuilder, data: dict):
    """AI Tutor slot overview (sessions_5)."""
    slots = data.get('sessions_5', [])
    if not slots:
        return

    d.page_break()
    d.h2('AI TUTOR SLOT OVERVIEW (sessions_5)')
    d.sep2()
    d.body('These 5 slots map to the AI Tutor station rotation in the app.')
    d.blank()
    for slot in slots:
        if isinstance(slot, dict):
            label   = slot.get('label', slot.get('slot', ''))
            content = slot.get('content', slot.get('items', []))
            if label:
                d.h3(label)
            if isinstance(content, list):
                for line in content:
                    d.indent(str(line))
            elif content:
                d.indent(str(content))
        else:
            d.indent(str(slot))
        d.blank()


# ─────────────────────────────────────────────────────────────────────────────
# Top-level builder
# ─────────────────────────────────────────────────────────────────────────────

def generate_lp_docx(week_num: int, out_dir: Path = None) -> Path:
    """
    Load W{week_num}.json and write W{week_num}_Lesson_Plan.docx.

    Returns the path to the generated file.
    Raises FileNotFoundError if the JSON does not exist.
    """
    if not _DOCX_OK:
        raise ImportError('python-docx is required: pip install python-docx')

    # Load JSON (prefer public/data/lessons, fall back to mcp-server)
    json_path = PUB_LESSONS / ('W%d.json' % week_num)
    if not json_path.exists():
        json_path = MCP_LESSONS / ('W%d.json' % week_num)
    if not json_path.exists():
        raise FileNotFoundError('W%d.json not found in public/data/lessons or mcp-server/data/lessons' % week_num)

    with open(json_path, encoding='utf-8') as f:
        data = json.load(f)

    # Determine output path
    if out_dir is None:
        out_dir = DEFAULT_OUT_DIR
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / ('W%d_Lesson_Plan.docx' % week_num)

    # Build document
    builder = LPDocBuilder()
    _build_cover(builder, data)
    _build_section1(builder, data)
    _build_section2(builder, data)
    _build_section3(builder, data)
    _build_session_sections(builder, data)
    _build_section_teacher(builder, data)
    _build_section_answer_key(builder, data)
    _build_section_task_cards(builder, data)
    _build_section_games(builder, data)
    _build_section_sessions5(builder, data)

    builder.save(out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# CLI helpers
# ─────────────────────────────────────────────────────────────────────────────

def _parse_week_args(args):
    weeks = set()
    for arg in args:
        m = re.match(r'^(\d+)-(\d+)$', arg)
        if m:
            weeks.update(range(int(m.group(1)), int(m.group(2)) + 1))
        elif arg.isdigit():
            weeks.add(int(arg))
    return sorted(weeks)


def _all_existing_weeks():
    weeks = set()
    for p in list(PUB_LESSONS.glob('W*.json')) + list(MCP_LESSONS.glob('W*.json')):
        m = re.match(r'W(\d+)\.json$', p.name)
        if m:
            weeks.add(int(m.group(1)))
    return sorted(weeks)


def main():
    args = sys.argv[1:]

    # --out-dir
    out_dir = None
    if '--out-dir' in args:
        idx = args.index('--out-dir')
        out_dir = Path(args[idx + 1])
        args = args[:idx] + args[idx + 2:]

    # --all
    if '--all' in args:
        weeks = _all_existing_weeks()
        args = [a for a in args if a != '--all']
    else:
        weeks = _parse_week_args(args)

    if not weeks:
        print('Usage: python3 pipeline/gen_lp_docx.py [--out-dir DIR] [--all] WEEK [WEEK...] [RANGE]')
        print()
        print('Examples:')
        print('  python3 pipeline/gen_lp_docx.py 37')
        print('  python3 pipeline/gen_lp_docx.py 37-53')
        print('  python3 pipeline/gen_lp_docx.py --all')
        sys.exit(1)

    if not _DOCX_OK:
        print('ERROR: python-docx not installed. Run: pip install python-docx')
        sys.exit(1)

    print('Generating %d lesson plan DOCX file(s)...' % len(weeks))
    success, failures = [], []

    for wn in weeks:
        try:
            path = generate_lp_docx(wn, out_dir=out_dir)
            print('  W%d ✓  →  %s' % (wn, path.name))
            success.append(wn)
        except FileNotFoundError as e:
            print('  W%d ✗  SKIP: %s' % (wn, e))
            failures.append(wn)
        except Exception as e:
            print('  W%d ✗  ERROR: %s' % (wn, e))
            failures.append(wn)

    print()
    print('─' * 60)
    print('Generated: %d/%d  →  %s' % (
        len(success), len(weeks),
        str(out_dir or DEFAULT_OUT_DIR)
    ))
    if failures:
        print('SKIPPED/FAILED: W%s' % ', W'.join(str(w) for w in failures))
    if failures and not success:
        sys.exit(1)


if __name__ == '__main__':
    main()
