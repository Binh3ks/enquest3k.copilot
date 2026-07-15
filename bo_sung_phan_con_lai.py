#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script bổ sung Phần V, VI, VII, VIII vào báo cáo đã tạo
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_custom(doc, text, level=1, color=RGBColor(31, 41, 55)):
    """Thêm tiêu đề có định dạng tùy chỉnh"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
    return heading

def add_table_styled(doc, data, headers, col_widths=None):
    """Thêm bảng có định dạng chuyên nghiệp"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    if col_widths:
        for idx, width in enumerate(col_widths):
            for cell in table.columns[idx].cells:
                cell.width = Inches(width)
    
    # Hàng tiêu đề
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4F46E5')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Các hàng dữ liệu
    for row_data in data:
        row_cells = table.add_row().cells
        for idx, cell_data in enumerate(row_data):
            row_cells[idx].text = str(cell_data)
            row_cells[idx].paragraphs[0].runs[0].font.size = Pt(9)
    
    return table

def add_financial_section(doc):
    """Thêm PHẦN V: DỰ PHÓNG TÀI CHÍNH"""
    print("Đang thêm Phần V: Dự phóng tài chính...")
    
    # ==================== PHẦN V ====================
    add_heading_custom(doc, 'PHẦN V', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'DỰ PHÓNG TÀI CHÍNH CHI TIẾT 3 NĂM', level=1, color=RGBColor(31, 41, 55))
    
    doc.add_paragraph(
        'Phần này trình bày dự phóng tài chính chi tiết cho 3 năm, bao gồm doanh thu, '
        'chi phí, lợi nhuận, và các chỉ số kinh tế quan trọng.'
    )
    
    doc.add_page_break()
    
    # ========== 13. DỰ BÁO DOANH THU 3 NĂM ==========
    add_heading_custom(doc, '13. DỰ BÁO DOANH THU 3 NĂM', level=2)
    
    revenue_forecast = [
        ['Kỳ', 'B2B Users', 'B2B ARPU', 'B2B Revenue', 'B2C Users', 'B2C ARPU', 'B2C Revenue', 'Tổng Revenue'],
        ['Tháng 6 (Năm 1)', '700', '84k', '58.8M', '0', '-', '0', '58.8M'],
        ['Tháng 12 (Năm 1)', '3,200', '84k', '268.8M', '550', '280k', '154M', '422.8M'],
        ['Tháng 18 (Năm 2 H1)', '5,000', '84k', '420M', '1,800', '320k', '576M', '996M'],
        ['Tháng 24 (Năm 2)', '7,500', '90k', '675M', '5,000', '350k', '1,750M', '2,425M'],
        ['Tháng 36 (Năm 3)', '10,000', '95k', '950M', '15,000', '420k', '6,300M', '7,250M'],
        ['', '', '', '', '', '', 'ARR Năm 3:', '87 tỷ VNĐ']
    ]
    
    doc.add_paragraph('📊 Bảng 13.1: Dự báo doanh thu 3 năm theo kênh')
    table = add_table_styled(doc, revenue_forecast[1:], revenue_forecast[0], col_widths=[1.2, 0.9, 0.9, 1.1, 0.9, 0.9, 1.1, 1.2])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('📈 PHÂN TÍCH TĂNG TRƯỞNG:\n\n')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    doc.add_paragraph(
        '• Năm 1: Tăng từ 0 → 423M/tháng (~5 tỷ ARR)\n'
        '• Năm 2: Tăng từ 423M → 2,425M/tháng (~29 tỷ ARR) - Tăng gấp 5.7x\n'
        '• Năm 3: Tăng từ 2,425M → 7,250M/tháng (~87 tỷ ARR) - Tăng gấp 3x\n\n'
        'Tốc độ tăng trưởng chậm lại ở năm 3 là BÌ NH THƯỜNG (mature stage) nhưng vẫn rất cao (3x YoY).'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph('Doanh thu tích lũy 3 năm:')
    cumulative_revenue = [
        ['Năm', 'Tổng doanh thu năm', 'Tích lũy', 'Ghi chú'],
        ['Năm 1', '~2.5 tỷ', '2.5 tỷ', 'Nửa đầu năm rất thấp, tăng nhanh cuối năm'],
        ['Năm 2', '~18 tỷ', '20.5 tỷ', 'Tăng trưởng mạnh nhờ B2C'],
        ['Năm 3', '~58 tỷ', '78.5 tỷ', 'B2C chiếm ưu thế (87%)'],
        ['Tổng 3 năm', '78.5 tỷ VNĐ', '~3.4M USD', 'Quy đổi tỷ giá 23,000']
    ]
    
    doc.add_paragraph('📊 Bảng 13.2: Doanh thu tích lũy 3 năm')
    table = add_table_styled(doc, cumulative_revenue[1:], cumulative_revenue[0], col_widths=[1.2, 2.0, 1.5, 2.8])
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 14. PHÂN TÍCH KINH TẾ ĐƠN VỊ ==========
    add_heading_custom(doc, '14. PHÂN TÍCH KINH TẾ ĐƠN VỊ (UNIT ECONOMICS)', level=2)
    
    doc.add_paragraph(
        'Unit economics là yếu tố quyết định sự bền vững của mô hình kinh doanh. '
        'Phân tích dưới đây so sánh chi tiết giữa B2B và B2C.'
    )
    
    doc.add_paragraph()
    
    unit_economics = [
        ['Chỉ số', 'B2B (Năm 1)', 'B2B (Năm 3)', 'B2C (Năm 2)', 'B2C (Năm 3)', 'Benchmark tốt'],
        ['CAC', '15,000đ', '10,000đ', '80,000đ', '60,000đ', '< 1/3 LTV'],
        ['ARPU/tháng', '84,000đ', '95,000đ', '320,000đ', '420,000đ', 'Càng cao càng tốt'],
        ['Gross Margin', '85%', '90%', '80%', '82%', '> 70%'],
        ['Churn rate/tháng', '5%', '3%', '8%', '5%', '< 5%'],
        ['Avg lifetime (tháng)', '20', '33', '12.5', '20', '> 12 tháng'],
        ['LTV', '1,680,000đ', '3,135,000đ', '4,000,000đ', '8,400,000đ', '> 3x CAC'],
        ['LTV:CAC ratio', '112:1', '314:1', '50:1', '140:1', '> 3:1'],
        ['Payback period', '< 1 tháng', '< 1 tháng', '3 tháng', '1.7 tháng', '< 12 tháng'],
        ['Đánh giá', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐', '']
    ]
    
    doc.add_paragraph('📊 Bảng 14.1: So sánh unit economics B2B vs B2C')
    table = add_table_styled(doc, unit_economics[1:], unit_economics[0], col_widths=[1.5, 1.0, 1.0, 1.0, 1.0, 1.2])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('💡 INSIGHT QUAN TRỌNG:\n\n')
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    run.font.size = Pt(11)
    
    doc.add_paragraph(
        '1. B2B có unit economics XUẤT SẮC:\n'
        '   • LTV:CAC > 100:1 là cực kỳ hiếm thấy trong SaaS\n'
        '   • Payback period < 1 tháng → Cash flow dương ngay lập tức\n'
        '   • Đây là lý do B2B là điểm khởi đầu hoàn hảo\n\n'
        '2. B2C có unit economics RẤT TỐT (nhưng không bằng B2B):\n'
        '   • LTV:CAC năm 2 là 50:1 (so với benchmark 3:1) → Rất tốt\n'
        '   • Năm 3 cải thiện lên 140:1 nhờ brand mạnh → CAC giảm\n'
        '   • Payback 1.7 tháng là rất nhanh trong B2C EdTech\n\n'
        '3. DUAL-TRACK tối ưu cả hai:\n'
        '   • B2B: High efficiency, low revenue ceiling\n'
        '   • B2C: Lower efficiency but massive scale\n'
        '   • Combined: Best of both worlds'
    )
    
    doc.add_page_break()
    
    # ========== 15. KẾ HOẠCH ĐẦU TƯ VÀ SỬ DỤNG VỐN ==========
    add_heading_custom(doc, '15. KẾ HOẠCH ĐẦU TƯ VÀ SỬ DỤNG VỐN', level=2)
    
    doc.add_paragraph(
        'Với mô hình tăng trưởng nhanh, LEXIO cần gọi vốn để scale. Dưới đây là kế hoạch '
        'gọi vốn và phân bổ chi tiêu.'
    )
    
    doc.add_paragraph()
    
    fundraising_plan = [
        ['Round', 'Thời điểm', 'Số tiền', 'Valuation', 'Sử dụng vốn', 'Milestone đạt được'],
        [
            'Seed',
            'Tháng 6 Năm 1',
            '500k-1M USD',
            '4-5M USD',
            'Product, team, B2B scale',
            '700 HS, 60 GV, PMF validated'
        ],
        [
            'Series A',
            'Tháng 18 Năm 2',
            '2-3M USD',
            '15-20M USD',
            'B2C scale, SEA expansion',
            '7k HS, 1 tỷ ARR, dual-track proven'
        ],
        [
            'Series B',
            'Tháng 30-36 Năm 3',
            '5-10M USD',
            '50-80M USD',
            'Regional domination, Series B/C prep',
            '25k HS, 5-8 tỷ/tháng, #1 CLIL Vietnam'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 15.1: Lộ trình gọi vốn 3 năm')
    table = add_table_styled(doc, fundraising_plan[1:], fundraising_plan[0], col_widths=[1.0, 1.2, 1.0, 1.2, 1.8, 2.0])
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    # Chi tiết sử dụng vốn
    doc.add_paragraph('Chi tiết sử dụng vốn Series A (2-3M USD):')
    
    series_a_usage = [
        ['Hạng mục', 'Phân bổ (%)', 'Số tiền (USD)', 'Mục đích cụ thể'],
        ['Sales & Marketing', '45%', '900k-1.35M', 'B2C ads (500k), B2B expansion (200k), Brand building (200k)'],
        ['Product & Engineering', '25%', '500k-750k', 'Hire 15 engineers, new features, AI tutor'],
        ['Content', '10%', '200k-300k', 'Tăng từ 500 → 2,000 lessons, đa dạng hóa'],
        ['Operations & Support', '10%', '200k-300k', 'CS team 20 người, infrastructure'],
        ['Management & Admin', '10%', '200k-300k', 'Leadership team, office, legal, accounting']
    ]
    
    doc.add_paragraph('📊 Bảng 15.2: Phân bổ vốn Series A')
    table = add_table_styled(doc, series_a_usage[1:], series_a_usage[0], col_widths=[1.8, 1.0, 1.3, 3.0])
    doc.add_paragraph()
    
    doc.add_page_break()

def add_risk_section(doc):
    """Thêm PHẦN VI: QUẢN TRỊ RỦI RO"""
    print("Đang thêm Phần VI: Quản trị rủi ro...")
    
    # ==================== PHẦN VI ====================
    add_heading_custom(doc, 'PHẦN VI', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'QUẢN TRỊ RỦI RO VÀ GIẢM THIỂU', level=1, color=RGBColor(31, 41, 55))
    
    doc.add_paragraph(
        'Mọi chiến lược đều có rủi ro. Phần này xác định các rủi ro chính và phương án giảm thiểu.'
    )
    
    doc.add_page_break()
    
    # ========== 16. RỦI RO CHIẾN LƯỢC ==========
    add_heading_custom(doc, '16. RỦI RO CHIẾN LƯỢC', level=2)
    
    strategic_risks = [
        ['Rủi ro', 'Mức độ', 'Tác động', 'Phương án giảm thiểu', 'Kế hoạch dự phòng'],
        [
            'Đối thủ lớn (Duolingo, etc.) copy mô hình B2B',
            'Trung bình',
            'Mất thị phần, cạnh tranh giá',
            '• Xây dựng brand mạnh sớm\n• Khóa exclusive với Cambridge\n• Community loyal',
            'Pivot sang B2C premium, differentiate bằng local content'
        ],
        [
            'B2C không scale như kỳ vọng',
            'Trung bình',
            'Doanh thu thấp hơn, khó gọi vốn',
            '• Pilot kỹ trước khi scale\n• A/B test nhiều channels\n• Có B2B làm nền',
            'Focus B2B + Enterprise, chấp nhận TAM nhỏ hơn'
        ],
        [
            'Thay đổi chính sách giáo dục',
            'Thấp',
            'Nhu cầu thay đổi đột ngột',
            '• Theo sát chính sách\n• Linh hoạt điều chỉnh content\n• Cambridge cert = safe',
            'Mở rộng sang các môn khác (toán, khoa học)'
        ],
        [
            'Vốn cạn trước khi đạt break-even',
            'Thấp-TB',
            'Phá sản hoặc bán giá rẻ',
            '• Unit economics tốt → Ít cần vốn\n• Gọi vốn sớm với traction tốt\n• B2B cash-positive',
            'Cắt giảm marketing B2C, focus B2B để profitable'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 16.1: Ma trận rủi ro chiến lược')
    table = add_table_styled(doc, strategic_risks[1:], strategic_risks[0], col_widths=[1.8, 0.9, 1.3, 2.0, 2.0])
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 17. RỦI RO VẬN HÀNH ==========
    add_heading_custom(doc, '17. RỦI RO VẬN HÀNH', level=2)
    
    operational_risks = [
        ['Rủi ro', 'Mức độ', 'Tác động', 'Phương án giảm thiểu'],
        [
            'Giáo viên churn cao',
            'Trung bình',
            'Mất học sinh, phải tìm GV mới tốn kém',
            '• Tier system khuyến khích retention\n• Training tốt\n• Community building\n• Support responsiveness'
        ],
        [
            'Quality control khó khăn khi scale',
            'Cao',
            'Trải nghiệm user kém, churn tăng',
            '• Automated QA system\n• NPS tracking real-time\n• Regional managers\n• Clear SOP'
        ],
        [
            'Technical issues (downtime, bugs)',
            'Trung bình',
            'User frustration, negative reviews',
            '• 99.9% uptime SLA\n• Monitoring 24/7\n• Bug bounty program\n• Fast deployment cycle'
        ],
        [
            'Payment fraud hoặc refund cao',
            'Thấp',
            'Mất revenue, tăng chi phí',
            '• Strong payment verification\n• Clear refund policy\n• Trial period để user test\n• Insurance'
        ],
        [
            'Team turnover cao',
            'Trung bình',
            'Mất know-how, chậm phát triển',
            '• Competitive salary + equity\n• Good culture\n• Career path rõ ràng\n• Knowledge documentation'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 17.1: Ma trận rủi ro vận hành')
    table = add_table_styled(doc, operational_risks[1:], operational_risks[0], col_widths=[2.0, 1.0, 1.8, 3.0])
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 18. PHƯƠNG ÁN DỰ PHÒNG ==========
    add_heading_custom(doc, '18. CÁC KỊCH BẢN DỰ PHÒNG', level=2)
    
    doc.add_paragraph(
        'Chuẩn bị cho các kịch bản worst-case và best-case:'
    )
    
    doc.add_paragraph()
    
    scenarios = [
        ['Kịch bản', 'Xác suất', 'Mô tả', 'Hành động'],
        [
            'Worst Case: Tăng trưởng chậm',
            '15%',
            'Chỉ đạt 50% target users năm 2',
            '• Cắt giảm marketing B2C\n• Focus profitable B2B\n• Extend runway\n• Downsizing team 20%'
        ],
        [
            'Base Case: Như dự báo',
            '60%',
            'Đạt 80-100% targets',
            '• Thực hiện đúng kế hoạch\n• Điều chỉnh nhỏ theo thị trường\n• Gọi vốn đúng timeline'
        ],
        [
            'Best Case: Vượt kỳ vọng',
            '25%',
            'Đạt 150%+ targets nhờ viral mạnh',
            '• Accelerate B2C scale\n• Gọi vốn lớn hơn\n• Fast-track SEA expansion\n• M&A competitors'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 18.1: Các kịch bản và phản ứng')
    table = add_table_styled(doc, scenarios[1:], scenarios[0], col_widths=[1.8, 1.0, 2.2, 2.5])
    doc.add_paragraph()
    
    doc.add_page_break()

def add_kpi_section(doc):
    """Thêm PHẦN VII: KPIs"""
    print("Đang thêm Phần VII: KPIs và đo lường...")
    
    # ==================== PHẦN VII ====================
    add_heading_custom(doc, 'PHẦN VII', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'ĐÁNH GIÁ THÀNH CÔNG - KPIs VÀ MILESTONES', level=1, color=RGBColor(31, 41, 55))
    
    doc.add_paragraph(
        'Để đo lường thành công và điều chỉnh chiến lược kịp thời, LEXIO thiết lập hệ thống KPIs '
        'toàn diện theo từng giai đoạn.'
    )
    
    doc.add_page_break()
    
    # ========== 19. CÁC CHỈ SỐ ĐO LƯỜNG THEN CHỐT ==========
    add_heading_custom(doc, '19. CÁC CHỈ SỐ ĐO LƯỜNG THEN CHỐT (KPIs)', level=2)
    
    doc.add_paragraph('19.1. North Star Metric')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('🌟 NORTH STAR METRIC: ')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(79, 70, 229)
    run = p.add_run(
        'Số giờ học tích cực (Active Learning Hours) mỗi tuần\n\n'
        'Lý do: Chỉ số này phản ánh giá trị thực sự mà LEXIO mang lại. Nếu học sinh học nhiều '
        'và đều đặn, họ sẽ tiến bộ → Phụ huynh hài lòng → Retention cao → Word-of-mouth tốt.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('19.2. KPIs theo từng lĩnh vực')
    doc.add_paragraph()
    
    kpi_categories = [
        ['Lĩnh vực', 'KPI chính', 'Target Năm 1', 'Target Năm 2', 'Target Năm 3', 'Tần suất đo'],
        [
            'Tăng trưởng (Growth)',
            'MoM User Growth Rate',
            '30-50%',
            '15-25%',
            '10-15%',
            'Hàng tháng'
        ],
        [
            'Tăng trưởng',
            'New Users/Month',
            '100-150',
            '500-800',
            '1,500-2,500',
            'Hàng tháng'
        ],
        [
            'Engagement',
            'DAU/MAU ratio',
            '> 40%',
            '> 45%',
            '> 50%',
            'Hàng ngày'
        ],
        [
            'Engagement',
            'Avg. session time',
            '> 15 phút',
            '> 18 phút',
            '> 20 phút',
            'Hàng ngày'
        ],
        [
            'Retention',
            'Month-6 retention',
            '> 35%',
            '> 50%',
            '> 65%',
            'Cohort analysis'
        ],
        [
            'Monetization',
            'ARPU',
            '84k (B2B)',
            '180k (mix)',
            '280k (mix)',
            'Hàng tháng'
        ],
        [
            'Monetization',
            'LTV:CAC',
            '> 50:1 (B2B)',
            '> 20:1 (mix)',
            '> 50:1 (mix)',
            'Hàng quý'
        ],
        [
            'Satisfaction',
            'NPS Score',
            '> 50',
            '> 60',
            '> 70',
            'Hàng quý'
        ],
        [
            'Satisfaction',
            'App Store Rating',
            '> 4.3',
            '> 4.5',
            '> 4.7',
            'Liên tục'
        ],
        [
            'Efficiency',
            'CAC',
            '< 20k',
            '< 60k',
            '< 50k',
            'Hàng tháng'
        ],
        [
            'Efficiency',
            'Payback Period',
            '< 1 tháng',
            '< 3 tháng',
            '< 2 tháng',
            'Hàng quý'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 19.1: Hệ thống KPIs theo lĩnh vực')
    table = add_table_styled(doc, kpi_categories[1:], kpi_categories[0], col_widths=[1.3, 1.5, 1.0, 1.0, 1.0, 1.2])
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 20. CỘT MỐC QUAN TRỌNG ==========
    add_heading_custom(doc, '20. CÁC CỘT MỐC QUAN TRỌNG (MILESTONES)', level=2)
    
    milestones = [
        ['Thời điểm', 'Milestone', 'Ý nghĩa', 'Hành động sau milestone'],
        [
            'Tháng 2',
            '100 học sinh active đầu tiên',
            'Product-Market Fit validation',
            'Scale lên 30 GV'
        ],
        [
            'Tháng 6',
            '700 HS, 60 GV, 50M MRR',
            'Đủ traction gọi vốn Seed',
            'Roadshow investors, close Seed'
        ],
        [
            'Tháng 12',
            '4,000 HS, 400M MRR, cả B2B+B2C',
            'Dual-track model proven',
            'Scale B2C mạnh với vốn Seed'
        ],
        [
            'Tháng 18',
            '7,000 HS, 1 tỷ MRR',
            'Unicorn trajectory, sẵn sàng Series A',
            'Gọi vốn Series A 2-3M'
        ],
        [
            'Tháng 24',
            '12,000 HS, 2.4 tỷ MRR',
            'Top 3 EdTech Vietnam',
            'Chuẩn bị mở rộng SEA'
        ],
        [
            'Tháng 36',
            '25,000 HS, 7 tỷ MRR, có mặt Thailand',
            'Regional player, sẵn sàng Series B',
            'Gọi vốn Series B 5-10M, scale SEA'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 20.1: Lộ trình milestones 3 năm')
    table = add_table_styled(doc, milestones[1:], milestones[0], col_widths=[1.0, 2.5, 1.8, 2.2])
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 21. HỆ THỐNG BÁO CÁO ==========
    add_heading_custom(doc, '21. HỆ THỐNG BÁO CÁO VÀ GIÁM SÁT', level=2)
    
    doc.add_paragraph(
        'Để đảm bảo thực thi đúng kế hoạch, LEXIO thiết lập hệ thống báo cáo đa tầng:'
    )
    
    doc.add_paragraph()
    
    reporting_system = [
        ['Loại báo cáo', 'Tần suất', 'Người làm', 'Người nhận', 'Nội dung chính'],
        [
            'Daily Metrics Dashboard',
            'Hàng ngày (real-time)',
            'Tự động (dashboard)',
            'Toàn bộ team',
            'DAU, New signups, Revenue, Churn, Critical bugs'
        ],
        [
            'Weekly Team Sync',
            'Hàng tuần',
            'Heads of Department',
            'CEO, Leadership team',
            'Progress vs targets, Blockers, Next week plan'
        ],
        [
            'Monthly Business Review',
            'Hàng tháng',
            'CEO',
            'Board, Investors',
            'MRR, User growth, KPIs, Cash runway, Big decisions'
        ],
        [
            'Quarterly Board Meeting',
            'Hàng quý',
            'CEO + CFO',
            'Board of Directors',
            'Strategy review, Financial deep-dive, Fundraising, Hiring'
        ],
        [
            'Annual Planning',
            'Hàng năm',
            'Leadership team',
            'Toàn công ty',
            'Year review, Next year OKRs, Budget allocation'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 21.1: Hệ thống báo cáo giám sát')
    table = add_table_styled(doc, reporting_system[1:], reporting_system[0], col_widths=[1.5, 1.2, 1.2, 1.2, 2.5])
    doc.add_paragraph()
    
    doc.add_page_break()

def add_conclusion_section(doc):
    """Thêm PHẦN VIII: KẾT LUẬN"""
    print("Đang thêm Phần VIII: Kết luận và khuyến nghị...")
    
    # ==================== PHẦN VIII ====================
    add_heading_custom(doc, 'PHẦN VIII', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'KẾT LUẬN VÀ KHUYẾN NGHỊ', level=1, color=RGBColor(31, 41, 55))
    
    doc.add_page_break()
    
    # ========== 22. TỔNG KẾT CHIẾN LƯỢC ==========
    add_heading_custom(doc, '22. TỔNG KẾT CHIẾN LƯỢC', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('🎯 CHIẾN LƯỢC TỔNG THỂ:\n\n')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    doc.add_paragraph(
        'Sau khi phân tích toàn diện thị trường, đối thủ cạnh tranh, và các mô hình kinh doanh, '
        'LEXIO đưa ra chiến lược DUAL-TRACK: Bắt đầu với B2B, sau đó mở rộng B2C song song, '
        'tạo hiệu ứng hiệp lực giữa hai kênh.\n'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph('🔑 CÁC QUYẾT ĐỊNH CHIẾN LƯỢC THEN CHỐT:')
    
    key_decisions = [
        '1. B2B TRƯỚC, B2C SAU - KHÔNG PHẢI CHỌN MỘT:',
        '   • B2B (6 tháng đầu): CAC thấp (15k), tạo market proof nhanh',
        '   • B2C (từ tháng 7): ARPU cao (320k+), scale unlimited',
        '   • Dual-track (năm 2-3): Tận dụng synergy, giảm rủi ro',
        '',
        '2. MÔ HÌNH HOA HỒNG GIÁO VIÊN - WIN-WIN-WIN:',
        '   • Giáo viên: Thu nhập thêm 400k-5M/tháng không tốn thời gian',
        '   • Học sinh: Học tốt hơn với giá rẻ (99k vs 500k-2M trung tâm)',
        '   • LEXIO: CAC thấp, viral mạnh, unit economics xuất sắc (LTV:CAC > 100:1)',
        '',
        '3. PHÂN BỔ TÀI NGUYÊN THEO GIAI ĐOẠN:',
        '   • Năm 1: 90% vào B2B, 10% chuẩn bị B2C',
        '   • Năm 2: 50% B2B maintenance, 50% B2C scale',
        '   • Năm 3: 30% B2B+Enterprise, 70% B2C+Premium',
        '',
        '4. GỌI VỐN CHỦ ĐỘNG VÀ SỚM:',
        '   • Seed (tháng 6): 500k-1M USD với 700 HS',
        '   • Series A (tháng 18): 2-3M USD với 7k HS, 1 tỷ ARR',
        '   • Series B (năm 3): 5-10M USD với 25k HS, regional footprint',
        '',
        '5. MỞ RỘNG ĐỊA LÝ TỪNG BƯỚC:',
        '   • Năm 1: Miền Bắc (Hà Nội, Bắc Ninh, Hải Phòng)',
        '   • Năm 2: Toàn quốc 50+ tỉnh/thành',
        '   • Năm 3: Đông Nam Á (Thailand, Indonesia pilot)'
    ]
    
    for item in key_decisions:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== 23. LỘ TRÌNH TRIỂN KHAI ƯU TIÊN ==========
    add_heading_custom(doc, '23. LỘ TRÌNH TRIỂN KHAI ƯU TIÊN', level=2)
    
    doc.add_paragraph(
        'Dựa trên phân tích trên, đây là lộ trình hành động cụ thể cho 6 tháng đầu tiên '
        '(quan trọng nhất):'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph('📅 THÁNG 1-2: FOUNDATION')
    
    month_1_2 = [
        '✅ SẢN PHẨM:',
        '   □ Finalize MVP với core features: Lessons, Gamification, Progress tracking',
        '   □ iOS + Android app submission',
        '   □ Teacher dashboard beta',
        '   □ Payment integration (Momo, VNPay, Banking)',
        '',
        '✅ TEAM:',
        '   □ Hire Product Manager (1)',
        '   □ Hire Backend Engineer (2)',
        '   □ Hire Content Creator (2)',
        '   □ Hire Sales Lead (1)',
        '',
        '✅ GO-TO-MARKET:',
        '   □ Tìm 10 giáo viên pilot (network + Facebook groups)',
        '   □ Đào tạo 1-1 cho từng GV (2h/người)',
        '   □ Launch với 100-150 học sinh',
        '   □ Thu thập feedback hàng tuần',
        '',
        '✅ MILESTONE: 150 HS active, NPS > 50, PMF signals clear'
    ]
    
    for item in month_1_2:
        doc.add_paragraph(item)
    
    doc.add_paragraph()
    doc.add_paragraph('📅 THÁNG 3-4: SCALE INITIAL')
    
    month_3_4 = [
        '✅ SẢN PHẨM:',
        '   □ Ship 50+ improvements từ feedback',
        '   □ Tự động hóa onboarding (video + email sequence)',
        '   □ Gamification v2 (leaderboard, badges)',
        '',
        '✅ GO-TO-MARKET:',
        '   □ Launch referral program (100k/GV mới)',
        '   □ Facebook Ads cho GV (budget: 5M)',
        '   □ Target: 30 GV (3x growth)',
        '   □ Xây 3 case studies (như cô Lan)',
        '',
        '✅ MILESTONE: 350-400 HS, 30 GV, 30M MRR'
    ]
    
    for item in month_3_4:
        doc.add_paragraph(item)
    
    doc.add_paragraph()
    doc.add_paragraph('📅 THÁNG 5-6: PRE-FUNDRAISE')
    
    month_5_6 = [
        '✅ SẢN PHẨM:',
        '   □ Cambridge CLIL certification cho GV',
        '   □ AI tutor beta',
        '   □ B2C landing page',
        '',
        '✅ GO-TO-MARKET:',
        '   □ Mở rộng 5-6 tỉnh mới',
        '   □ Partnership với 3-5 trung tâm',
        '   □ Target: 60 GV, 700 HS',
        '',
        '✅ FUNDRAISING:',
        '   □ Chuẩn bị pitch deck',
        '   □ Financial model 3 năm',
        '   □ Tiếp cận 15-20 quỹ Seed',
        '   □ Close 500k-1M USD',
        '',
        '✅ MILESTONE: 700 HS, 50M MRR, Term Sheet signed'
    ]
    
    for item in month_5_6:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== 24. KHUYẾN NGHỊ CUỐI CÙNG ==========
    add_heading_custom(doc, '24. KHUYẾN NGHỊ CUỐI CÙNG', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('💎 LỜI KẾT:\n\n')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(220, 38, 38)
    
    doc.add_paragraph(
        'Thị trường giáo dục Việt Nam đang chuyển mình mạnh mẽ. Với 15 triệu học sinh đang học '
        'tiếng Anh và chỉ 15-20% sử dụng công nghệ, đây là cơ hội vàng cho LEXIO.\n'
    )
    
    doc.add_paragraph()
    
    final_recommendations = [
        '🎯 BA KHUYẾN NGHỊ QUAN TRỌNG NHẤT:',
        '',
        '1. KHÔNG NÊN CHỌN GIỮA B2B HOẶC B2C - PHẢI LÀM CẢ HAI:',
        '   • Bất kỳ mô hình đơn kênh nào cũng có giới hạn nghiêm trọng',
        '   • B2B alone: CAC tốt nhưng ARPU và TAM thấp',
        '   • B2C alone: TAM lớn nhưng CAC cao và khó tạo market proof ban đầu',
        '   • Dual-track: Best of both worlds + Synergy effects',
        '   ⚠️ Quan trọng: Phải BẮT ĐẦU với B2B, không phải B2C!',
        '',
        '2. MÔ HÌNH HOA HỒNG GV LÀ COMPETITIVE MOAT:',
        '   • Đây là điểm khác biệt lớn nhất so với đối thủ',
        '   • Tạo network effects mạnh: GV giới thiệu GV → Viral',
        '   • Unit economics xuất sắc: LTV:CAC > 100:1 (B2B)',
        '   • Hard to copy: Cần thời gian xây dựng trust với cộng đồng GV',
        '   ⚠️ Đầu tư vào GV community là đầu tư chiến lược dài hạn!',
        '',
        '3. GỌI VỐN SỚM VÀ CHỦ ĐỘNG:',
        '   • Đừng chờ đến khi hết tiền mới gọi vốn',
        '   • Với traction tốt (700 HS tháng 6), có thể gọi 500k-1M Seed',
        '   • Vốn này đủ để scale đến tháng 18 → Gọi Series A 2-3M',
        '   • EdTech là "winner takes most" → Scale nhanh = win',
        '   ⚠️ Chấp nhận dilution hợp lý để có đủ đạn dược cạnh tranh!',
        '',
        '',
        '🚀 KẾT LUẬN CUỐI CÙNG:',
        '',
        'LEXIO với chiến lược dual-track B2B+B2C, mô hình hoa hồng giáo viên độc đáo, và nội dung '
        'Cambridge CLIL chất lượng cao, có tiềm năng trở thành nền tảng học tiếng Anh số 1 Việt Nam '
        'trong 3 năm tới.\n',
        '',
        'Với unit economics xuất sắc (LTV:CAC > 50:1), tốc độ tăng trưởng nhanh (MoM 30-50% năm 1), '
        'và TAM khổng lồ (225M USD nếu kết hợp cả hai kênh), LEXIO là cơ hội đầu tư hấp dẫn cho các '
        'quỹ tìm kiếm startup EdTech có khả năng scale khu vực.\n',
        '',
        'Thời điểm hiện tại (2026) là TỐI ƯU để triển khai:',
        '• Chính sách ESL đang được đẩy mạnh',
        '• Thị trường EdTech Việt Nam đang bùng nổ',
        '• Đối thủ lớn (Duolingo, ELSA) chưa làm B2B nghiêm túc',
        '• Cambridge đang mở rộng partnership tại Việt Nam',
        '',
        '⏰ THỜI CƠ KHÔNG CHỜ ĐỢI - HÀNH ĐỘNG NGAY!',
        '',
        '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━',
        '',
        '📞 LIÊN HỆ:',
        'Email: strategy@lexio.vn',
        'Phone: +84 xxx xxx xxx',
        'Website: www.lexio.vn',
        '',
        'Cảm ơn đã đọc báo cáo này!',
        '',
        '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
    ]
    
    for item in final_recommendations:
        if item.startswith('🎯') or item.startswith('🚀'):
            p = doc.add_paragraph()
            run = p.add_run(item)
            run.font.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(16, 185, 129)
        elif item.startswith('━━'):
            p = doc.add_paragraph(item)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif item.strip() == '':
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_paragraph()

def main():
    """Main function"""
    print("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    print("📄 BỔ SUNG CÁC PHẦN CÒN LẠI VÀO BÁO CÁO")
    print("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n")
    
    # Mở file đã tạo
    filepath = '/Users/binhnguyen/Downloads/Engquest3k/LEXIO_Bao_Cao_Chien_Luoc_Day_Du_20260507_2144.docx'
    
    try:
        doc = Document(filepath)
        print(f"✅ Đã mở file: {filepath}")
        print(f"📊 File hiện có: {len(doc.paragraphs)} đoạn, {len(doc.tables)} bảng\n")
        
        # Thêm các phần còn lại
        add_financial_section(doc)
        add_risk_section(doc)
        add_kpi_section(doc)
        add_conclusion_section(doc)
        
        # Lưu file
        doc.save(filepath)
        
        print(f"\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
        print(f"✅ ĐÃ BỔ SUNG THÀNH CÔNG!")
        print(f"📄 File: {filepath}")
        print(f"📊 Tổng số đoạn văn: {len(doc.paragraphs)}")
        print(f"📊 Tổng số bảng: {len(doc.tables)}")
        print(f"📄 Ước tính tổng số trang: {len(doc.paragraphs) // 7} - {len(doc.paragraphs) // 5} trang")
        print(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
        
    except Exception as e:
        print(f"❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()
