#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Bổ sung Phần IV-VII vào roadmap
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_custom(doc, text, level=1, color=RGBColor(31, 41, 55)):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return heading

def add_table_styled(doc, data, headers, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    if col_widths:
        for idx, width in enumerate(col_widths):
            for cell in table.columns[idx].cells:
                cell.width = Inches(width)
    
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4F46E5')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for row_data in data:
        row_cells = table.add_row().cells
        for idx, cell_data in enumerate(row_data):
            row_cells[idx].text = str(cell_data)
            row_cells[idx].paragraphs[0].runs[0].font.size = Pt(9)
    
    return table

def add_remaining_sections(doc):
    """Thêm các phần còn lại"""
    
    print("Đang thêm Phần IV: Analytics & Reporting...")
    
    # ==================== PHẦN IV: ANALYTICS ====================
    add_heading_custom(doc, 'PHẦN IV', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'HỆ THỐNG ANALYTICS VÀ REPORTING', level=1)
    
    doc.add_paragraph(
        'Hệ thống đo lường và báo cáo toàn diện để theo dõi KPIs và đưa ra quyết định dựa trên dữ liệu.'
    )
    
    doc.add_page_break()
    
    # ========== 4.1 USER ANALYTICS ==========
    add_heading_custom(doc, '4.1. User Analytics Dashboard', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('📊 METRICS CHI TIẾT:\n\n')
    run.font.bold = True
    run.font.size = Pt(11)
    
    analytics_features = [
        ['Nhóm metrics', 'Chi tiết', 'Công cụ', 'Timeline'],
        [
            'Acquisition',
            '• Daily/Monthly Active Users\n• New signups by source\n• Conversion funnel\n• CAC by channel',
            'Mixpanel + Google Analytics 4',
            'Tháng 1-2'
        ],
        [
            'Engagement',
            '• DAU/MAU ratio\n• Session duration\n• Lessons completed\n• Feature usage\n• Retention cohorts',
            'Amplitude hoặc Mixpanel',
            'Tháng 2-3'
        ],
        [
            'Monetization',
            '• ARPU by tier\n• LTV analysis\n• Churn rate\n• Upgrade/downgrade flows',
            'Custom dashboard + Stripe data',
            'Tháng 3-4'
        ],
        [
            'Content',
            '• Lesson completion rate\n• Quiz scores\n• Time per lesson\n• Drop-off points',
            'Custom tracking in DB',
            'Tháng 2-3'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng: Nhóm metrics cần tracking')
    table = add_table_styled(doc, analytics_features[1:], analytics_features[0], col_widths=[1.5, 2.5, 1.8, 1.0])
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    doc.add_paragraph('🛠️ IMPLEMENTATION:')
    implementation = [
        '• Event tracking: Instrument 50+ key events (app_open, lesson_start, lesson_complete, quiz_submit, etc.)',
        '• User properties: Store user attributes (age, level, tier, signup_date, etc.)',
        '• Custom dashboards: Grafana hoặc Metabase cho internal team',
        '• Real-time alerts: Slack notifications cho critical metrics (churn spike, payment failures)',
        '• A/B testing framework: Optimizely hoặc custom solution',
    ]
    
    for item in implementation:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_paragraph('👥 RESOURCES: 1 Data Engineer, 1 Analytics Engineer, 2 tuần setup + ongoing maintenance')
    
    doc.add_page_break()
    
    # ========== 4.2 BUSINESS INTELLIGENCE ==========
    add_heading_custom(doc, '4.2. Business Intelligence & Reporting', level=2)
    
    bi_reports = [
        ['Báo cáo', 'Tần suất', 'Người nhận', 'Nội dung'],
        [
            'Daily Metrics',
            'Hàng ngày 9am',
            'CEO, Product Lead',
            'DAU, Revenue, New signups, Churn, Critical bugs'
        ],
        [
            'Weekly Business Review',
            'Thứ 2 hàng tuần',
            'Leadership team',
            'KPIs vs targets, Growth trends, Top issues'
        ],
        [
            'Monthly Investor Report',
            'Ngày 5 hàng tháng',
            'Board, Investors',
            'MRR, User growth, Burn rate, Milestones, Risks'
        ],
        [
            'Quarterly Deep Dive',
            'Đầu quý',
            'All hands',
            'Strategy review, OKRs progress, Market analysis'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng: Hệ thống báo cáo doanh nghiệp')
    table = add_table_styled(doc, bi_reports[1:], bi_reports[0], col_widths=[1.5, 1.3, 1.5, 2.5])
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    doc.add_paragraph('🛠️ TOOLS:')
    tools = [
        '• Data warehouse: Google BigQuery hoặc AWS Redshift',
        '• ETL: Fivetran hoặc Airbyte để sync data từ các nguồn',
        '• Visualization: Looker, Tableau, hoặc Metabase',
        '• Automated reports: Python scripts + Slack/Email integration',
    ]
    
    for tool in tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('👥 RESOURCES: 1 BI Analyst, 1 Data Engineer, tháng 4-6')
    
    doc.add_page_break()
    
    print("Đang thêm Phần V: Infrastructure & DevOps...")
    
    # ==================== PHẦN V: INFRASTRUCTURE ====================
    add_heading_custom(doc, 'PHẦN V', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'INFRASTRUCTURE VÀ DEVOPS', level=1)
    
    doc.add_paragraph(
        'Hạ tầng kỹ thuật và quy trình DevOps để đảm bảo scalability, reliability, và security.'
    )
    
    doc.add_page_break()
    
    # ========== 5.1 CLOUD INFRASTRUCTURE ==========
    add_heading_custom(doc, '5.1. Cloud Infrastructure', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('☁️ KIẾN TRÚC TỔNG THỂ:\n\n')
    run.font.bold = True
    run.font.size = Pt(11)
    
    infrastructure = [
        ['Component', 'Technology', 'Purpose', 'Scalability'],
        [
            'Hosting',
            'AWS hoặc Google Cloud',
            'Main cloud provider',
            'Auto-scaling groups'
        ],
        [
            'App Server',
            'Node.js (Express) hoặc Python (FastAPI)',
            'API backend',
            'Kubernetes cluster, horizontal scaling'
        ],
        [
            'Database',
            'PostgreSQL (primary) + Redis (cache)',
            'Persistent data + caching',
            'Read replicas, sharding if needed'
        ],
        [
            'File Storage',
            'AWS S3 hoặc GCS',
            'Media files (audio, video, images)',
            'Unlimited, CDN cached'
        ],
        [
            'CDN',
            'CloudFront hoặc Cloudflare',
            'Fast content delivery',
            'Global edge locations'
        ],
        [
            'Real-time',
            'Firebase Firestore hoặc Socket.io',
            'Chat, notifications',
            'Firebase auto-scales'
        ],
        [
            'Video',
            'Zoom API / Agora.io',
            'Live tutoring',
            'Cloud-based, no hosting needed'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng: Các components hạ tầng')
    table = add_table_styled(doc, infrastructure[1:], infrastructure[0], col_widths=[1.3, 1.8, 1.8, 1.8])
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    doc.add_paragraph('💰 CHI PHÍ ƯỚC TÍNH:')
    cost_estimates = [
        '• Năm 1 (700 users): ~$500-800/tháng',
        '• Năm 2 (7,000 users): ~$2,000-3,000/tháng',
        '• Năm 3 (25,000 users): ~$6,000-10,000/tháng',
        '• LLM API costs (GPT-4): ~$1,000-3,000/tháng (năm 2-3)',
    ]
    
    for cost in cost_estimates:
        doc.add_paragraph(cost, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 5.2 DEVOPS & CI/CD ==========
    add_heading_custom(doc, '5.2. DevOps & CI/CD Pipeline', level=2)
    
    doc.add_paragraph('🔄 CI/CD WORKFLOW:')
    
    cicd_steps = [
        '1. Code commit: Developer push code to GitHub',
        '2. Automated tests: Unit tests, integration tests run on GitHub Actions',
        '3. Code review: Peer review, automated linters (ESLint, Black)',
        '4. Build: Docker image built',
        '5. Deploy to staging: Auto-deploy to staging environment',
        '6. QA testing: Manual + automated tests on staging',
        '7. Deploy to production: Manual approval → Deploy',
        '8. Monitoring: Track errors, performance via Sentry + Datadog',
    ]
    
    for step in cicd_steps:
        doc.add_paragraph(step)
    
    doc.add_paragraph()
    
    doc.add_paragraph('🛠️ TOOLS:')
    devops_tools = [
        '• Version control: GitHub',
        '• CI/CD: GitHub Actions hoặc GitLab CI',
        '• Containerization: Docker',
        '• Orchestration: Kubernetes (k8s)',
        '• Infrastructure as Code: Terraform',
        '• Monitoring: Datadog, New Relic, hoặc Grafana + Prometheus',
        '• Error tracking: Sentry',
        '• Log management: ELK stack (Elasticsearch, Logstash, Kibana)',
    ]
    
    for tool in devops_tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('👥 RESOURCES: 1 DevOps Engineer full-time từ tháng 3')
    
    doc.add_page_break()
    
    # ========== 5.3 SECURITY ==========
    add_heading_custom(doc, '5.3. Security & Compliance', level=2)
    
    security_measures = [
        ['Area', 'Measures', 'Timeline'],
        [
            'Data encryption',
            '• At rest: Database encryption\n• In transit: TLS/SSL\n• Sensitive data: Field-level encryption (PII)',
            'Tháng 1 (must-have)'
        ],
        [
            'Authentication',
            '• JWT tokens\n• 2FA for sensitive operations\n• OAuth for social login\n• Password: bcrypt hashing',
            'Tháng 1-2'
        ],
        [
            'Authorization',
            '• Role-based access control (RBAC)\n• API rate limiting\n• IP whitelisting for admin',
            'Tháng 2'
        ],
        [
            'Compliance',
            '• GDPR: Data privacy, right to delete\n• COPPA: Parental consent for kids <13\n• PCI DSS: Payment security (via Stripe)',
            'Tháng 3-6'
        ],
        [
            'Monitoring',
            '• Penetration testing: Annual\n• Security audits: Quarterly\n• Vulnerability scanning: Weekly',
            'Ongoing'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng: Biện pháp bảo mật')
    table = add_table_styled(doc, security_measures[1:], security_measures[0], col_widths=[1.5, 3.5, 1.5])
    doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_paragraph('👥 RESOURCES: 1 Security Engineer (consultant hoặc part-time), tháng 3-6')
    
    doc.add_page_break()
    
    print("Đang thêm Phần VI: QA & Testing...")
    
    # ==================== PHẦN VI: QA & TESTING ====================
    add_heading_custom(doc, 'PHẦN VI', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'QUALITY ASSURANCE VÀ TESTING', level=1)
    
    doc.add_paragraph(
        'Quy trình đảm bảo chất lượng để deliver sản phẩm ổn định, ít bugs.'
    )
    
    doc.add_page_break()
    
    # ========== 6.1 TESTING STRATEGY ==========
    add_heading_custom(doc, '6.1. Testing Strategy', level=2)
    
    testing_pyramid = [
        ['Level', 'Type', 'Coverage', 'Tools', 'Frequency'],
        [
            'Unit Tests',
            'Test individual functions',
            '70% code coverage',
            'Jest (JS), Pytest (Python)',
            'Every commit'
        ],
        [
            'Integration Tests',
            'Test API endpoints, DB interactions',
            '50% coverage',
            'Supertest, Postman',
            'Every merge to main'
        ],
        [
            'E2E Tests',
            'Test full user flows',
            '20 critical flows',
            'Cypress, Playwright',
            'Before production deploy'
        ],
        [
            'Manual QA',
            'Exploratory testing, UX review',
            'All new features',
            'Human testers',
            'Before production deploy'
        ],
        [
            'Performance Tests',
            'Load testing, stress testing',
            'Key endpoints',
            'k6, JMeter',
            'Monthly'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng: Testing pyramid')
    table = add_table_styled(doc, testing_pyramid[1:], testing_pyramid[0], col_widths=[1.2, 1.8, 1.3, 1.5, 1.2])
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    doc.add_paragraph('🎯 QUALITY METRICS:')
    quality_kpis = [
        '• Bug density: < 1 bug per 1000 lines of code',
        '• Critical bugs: 0 in production',
        '• Test coverage: > 70% code coverage',
        '• Deployment frequency: 2-3 times per week',
        '• Mean time to recovery (MTTR): < 2 hours',
    ]
    
    for kpi in quality_kpis:
        doc.add_paragraph(kpi, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('👥 RESOURCES: 2 QA Engineers full-time từ tháng 2')
    
    doc.add_page_break()
    
    # ========== 6.2 BETA TESTING ==========
    add_heading_custom(doc, '6.2. Beta Testing Program', level=2)
    
    doc.add_paragraph('📱 BETA TESTING WORKFLOW:')
    
    beta_phases = [
        '• Internal Alpha: Team members test new features (1 tuần)',
        '• Closed Beta: 20-50 selected teachers/students (2 tuần)',
        '• Open Beta: Opt-in for all users (2 tuần)',
        '• Production: Full rollout',
    ]
    
    for phase in beta_phases:
        doc.add_paragraph(phase, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_paragraph('🛠️ TOOLS:')
    beta_tools = [
        '• TestFlight (iOS) / Google Play Beta (Android)',
        '• Feature flags: LaunchDarkly để enable/disable features',
        '• Feedback collection: In-app surveys, bug report button',
        '• Crash reporting: Crashlytics (Firebase)',
    ]
    
    for tool in beta_tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_page_break()
    
    print("Đang thêm Phần VII: Timeline & Resources...")
    
    # ==================== PHẦN VII: TIMELINE ====================
    add_heading_custom(doc, 'PHẦN VII', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'TIMELINE TỔNG THỂ VÀ RESOURCES', level=1)
    
    doc.add_paragraph(
        'Lộ trình tổng thể và team size theo từng giai đoạn.'
    )
    
    doc.add_page_break()
    
    # ========== 7.1 MASTER TIMELINE ==========
    add_heading_custom(doc, '7.1. Master Timeline 3 Năm', level=2)
    
    master_timeline = [
        ['Giai đoạn', 'Tháng', 'Focus chính', 'Key deliverables', 'Team size'],
        [
            'MVP',
            '1-2',
            'Core app + Teacher onboarding',
            '• Basic lessons\n• Teacher portal\n• Referral codes\n• Payment integration',
            '8 người'
        ],
        [
            'Pilot',
            '3-4',
            'B2B commission system',
            '• Commission engine\n• Payout system\n• Teacher dashboard\n• 10 GV pilot',
            '10 người'
        ],
        [
            'Scale B2B',
            '5-6',
            'Scale to 60 GV',
            '• Tier system\n• Referral program\n• Analytics v1\n• 60 GV, 700 HS',
            '12 người'
        ],
        [
            'B2C Launch',
            '7-12',
            'B2C beta + Content expansion',
            '• B2C pricing tiers\n• Parent dashboard\n• IELTS content start\n• B2C 500 HS',
            '18 người'
        ],
        [
            'Dual-Track',
            '13-18',
            'Scale cả B2B và B2C',
            '• AI tutor beta\n• Premium features\n• 7k total users',
            '30 người'
        ],
        [
            'Premium',
            '19-24',
            'Live tutoring + Cambridge',
            '• 1-on-1 tutoring\n• Cambridge cert prep\n• Mầm non content',
            '45 người'
        ],
        [
            'Regional',
            '25-36',
            'Expansion SEA',
            '• Thailand launch\n• Indonesia prep\n• Business English\n• 25k users',
            '80 người'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng: Master timeline theo giai đoạn')
    table = add_table_styled(doc, master_timeline[1:], master_timeline[0], col_widths=[1.0, 0.8, 1.8, 2.2, 1.0])
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 7.2 TEAM STRUCTURE ==========
    add_heading_custom(doc, '7.2. Team Structure và Hiring Plan', level=2)
    
    doc.add_paragraph('👥 TEAM BREAKDOWN BY YEAR:')
    
    team_by_year = [
        ['Function', 'Năm 1', 'Năm 2', 'Năm 3', 'Notes'],
        [
            'Engineering',
            '5 (2 FE, 2 BE, 1 mobile)',
            '12 (4 FE, 5 BE, 2 mobile, 1 ML)',
            '25 (10 FE, 10 BE, 3 mobile, 2 ML)',
            'Hire senior engineers early'
        ],
        [
            'Product',
            '1 PM',
            '2 PMs',
            '4 PMs',
            'PM per product area'
        ],
        [
            'Design',
            '1 designer',
            '2 (1 UI/UX, 1 graphic)',
            '4',
            'Focus on UX research'
        ],
        [
            'Content',
            '2 creators',
            '6 (+ 2 IELTS specialists)',
            '12 (diverse specialties)',
            'Quality > quantity'
        ],
        [
            'QA',
            '1',
            '3',
            '6',
            'Mix manual + automation'
        ],
        [
            'DevOps',
            '0 (outsource)',
            '1',
            '2',
            'Hire when scale hits'
        ],
        [
            'Data',
            '0',
            '2 (1 engineer, 1 analyst)',
            '5',
            'Critical for decisions'
        ],
        [
            'Sales & Marketing',
            '2',
            '8',
            '20',
            'Regional managers năm 2'
        ],
        [
            'Customer Support',
            '1',
            '5',
            '15',
            'Tiered support'
        ],
        [
            'Operations',
            '1',
            '3',
            '8',
            'Finance, HR, Legal'
        ],
        [
            'TOTAL',
            '14',
            '44',
            '101',
            ''
        ]
    ]
    
    doc.add_paragraph('📊 Bảng: Team size theo function và năm')
    table = add_table_styled(doc, team_by_year[1:], team_by_year[0], col_widths=[1.5, 1.5, 1.5, 1.5, 1.5])
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    doc.add_paragraph('💰 TOTAL COST ESTIMATE:')
    salary_costs = [
        '• Năm 1: 14 người × avg 20M/năm = 280M (~12k USD/năm)',
        '• Năm 2: 44 người × avg 25M/năm = 1.1 tỷ (~48k USD/năm)',
        '• Năm 3: 101 người × avg 30M/năm = 3 tỷ (~130k USD/năm)',
        '• + Benefits, office, tools: ~30% overhead',
    ]
    
    for cost in salary_costs:
        doc.add_paragraph(cost, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 7.3 PRIORITIES ==========
    add_heading_custom(doc, '7.3. Phân chia Priorities', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('⚠️ NGUYÊN TẮC PHÂN PRIORITIES:\n\n')
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    
    priority_rules = [
        '🔴 P0 - CRITICAL (Must-have):',
        '• Block launch hoặc block revenue',
        '• Security vulnerabilities',
        '• Legal compliance',
        '• Examples: Payment integration, commission calculator, data encryption',
        '',
        '🟠 P1 - HIGH (Should-have):',
        '• Tăng conversion hoặc retention đáng kể',
        '• Competitive differentiation',
        '• Examples: AI tutor, teacher referral program, parent dashboard',
        '',
        '🟡 P2 - MEDIUM (Nice-to-have):',
        '• Improve UX hoặc efficiency',
        '• Support scale',
        '• Examples: Advanced analytics, tier system, class management',
        '',
        '🟢 P3 - LOW (Future):',
        '• Strategic but not urgent',
        '• Experimental',
        '• Examples: Business English, advanced gamification, social features',
    ]
    
    for rule in priority_rules:
        if rule.startswith('🔴') or rule.startswith('🟠') or rule.startswith('🟡') or rule.startswith('🟢'):
            p = doc.add_paragraph()
            run = p.add_run(rule)
            run.font.bold = True
        else:
            doc.add_paragraph(rule)
    
    doc.add_page_break()
    
    print("Đang thêm PHỤ LỤC: Tech Stack...")
    
    # ==================== PHỤ LỤC ====================
    add_heading_custom(doc, 'PHỤ LỤC', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'STACK CÔNG NGHỆ VÀ KIẾN TRÚC HỆ THỐNG', level=1)
    
    doc.add_page_break()
    
    add_heading_custom(doc, 'A. Tech Stack Đề Xuất', level=2)
    
    tech_stack = [
        ['Layer', 'Technology', 'Lý do chọn'],
        [
            'Mobile App',
            'React Native',
            '• Single codebase cho iOS + Android\n• Fast development\n• Mature ecosystem\n• Dễ tìm developers'
        ],
        [
            'Web App',
            'React + Next.js',
            '• Same language as mobile (JS)\n• SEO-friendly\n• Fast rendering\n• Great for landing pages'
        ],
        [
            'Backend API',
            'Node.js + Express HOẶC Python + FastAPI',
            '• Node: Same language stack, event-driven\n• Python: Better for ML/AI integration\n• Cả hai đều mature và scalable'
        ],
        [
            'Database',
            'PostgreSQL',
            '• Mature, reliable\n• ACID compliance\n• JSON support\n• Great for complex queries'
        ],
        [
            'Cache',
            'Redis',
            '• In-memory speed\n• Pub/sub for real-time\n• Session storage'
        ],
        [
            'File Storage',
            'AWS S3 / Google Cloud Storage',
            '• Cheap, unlimited\n• CDN integration\n• 99.999% durability'
        ],
        [
            'Real-time',
            'Firebase Firestore + Cloud Functions',
            '• Real-time sync\n• Offline support\n• Easy to use\n• Auto-scales'
        ],
        [
            'Auth',
            'Firebase Auth hoặc Auth0',
            '• Managed service\n• Social login\n• 2FA built-in'
        ],
        [
            'Payment',
            'Stripe (international) + VNPay/Momo (local)',
            '• PCI compliant\n• Recurring billing\n• Great APIs'
        ],
        [
            'Email',
            'SendGrid hoặc AWS SES',
            '• High deliverability\n• Templates\n• Analytics'
        ],
        [
            'SMS',
            'Twilio hoặc AWS SNS',
            '• Global coverage\n• Programmable'
        ],
        [
            'Video',
            'Zoom API / Agora.io',
            '• Turnkey solution\n• Good quality\n• Recordings'
        ],
        [
            'LLM',
            'OpenAI API (GPT-4)',
            '• State-of-the-art\n• Easy API\n• Fast iteration'
        ],
        [
            'Analytics',
            'Mixpanel / Amplitude',
            '• User-centric\n• Funnels, cohorts\n• Integrations'
        ],
        [
            'Monitoring',
            'Sentry + Datadog',
            '• Error tracking\n• Performance monitoring\n• Alerts'
        ],
        [
            'CI/CD',
            'GitHub Actions',
            '• Free for private repos\n• Easy setup\n• Good integrations'
        ],
        [
            'Hosting',
            'AWS hoặc Google Cloud',
            '• Mature, reliable\n• Global infrastructure\n• Full stack of services'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng: Tech stack chi tiết')
    table = add_table_styled(doc, tech_stack[1:], tech_stack[0], col_widths=[1.3, 1.8, 3.5])
    doc.add_paragraph()
    
    doc.add_page_break()
    
    add_heading_custom(doc, 'B. Kiến Trúc Hệ Thống (High-level)', level=2)
    
    doc.add_paragraph('🏗️ KIẾN TRÚC TỔNG QUAN:')
    doc.add_paragraph()
    
    architecture_description = [
        '┌─────────────────────┐',
        '│   Mobile Apps       │',
        '│  (React Native)     │',
        '└─────────┬───────────┘',
        '          │',
        '┌─────────▼───────────┐',
        '│    API Gateway      │',
        '│   (Load Balancer)   │',
        '└─────────┬───────────┘',
        '          │',
        '     ┌────┴────┐',
        '     │         │',
        '┌────▼───┐ ┌──▼─────┐',
        '│Backend │ │Firebase│',
        '│ API    │ │(Real-  │',
        '│(Node/  │ │time)   │',
        '│Python) │ │        │',
        '└────┬───┘ └────────┘',
        '     │',
        ' ┌───┴───┬───────┬─────────┐',
        ' │       │       │         │',
        '┌▼────┐┌▼────┐┌▼──────┐┌──▼───┐',
        '│Post-││Redis││S3/GCS ││LLM   │',
        '│greSQL││     ││(Media)││API   │',
        '└─────┘└─────┘└───────┘└──────┘',
    ]
    
    p = doc.add_paragraph()
    run = p.add_run('\n'.join(architecture_description))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_paragraph('📝 GIẢI THÍCH:')
    architecture_notes = [
        '• Users tương tác qua Mobile/Web apps',
        '• API Gateway: Routing, rate limiting, authentication',
        '• Backend API: Business logic, database operations',
        '• Firebase: Real-time features (chat, notifications)',
        '• PostgreSQL: Primary data store',
        '• Redis: Caching, session management',
        '• S3/GCS: Static assets (images, audio, video)',
        '• LLM API: AI tutor functionality',
    ]
    
    for note in architecture_notes:
        doc.add_paragraph(note, style='List Bullet')
    
    doc.add_page_break()
    
    add_heading_custom(doc, 'C. Database Schema (Core Tables)', level=2)
    
    doc.add_paragraph('🗄️ TABLES CHÍNH:')
    
    core_tables = [
        '1. users: id, email, name, role (student/teacher/admin), created_at',
        '2. teachers: user_id, referral_code, tier, total_students, total_earnings',
        '3. students: user_id, age, level, parent_email, teacher_id (nullable)',
        '4. subscriptions: user_id, plan (free/basic/pro/premium), status, start_date, end_date',
        '5. referrals: referrer_id, referee_id, type (student/teacher), signup_date, status',
        '6. commission_transactions: teacher_id, month, gross_revenue, commission_amount, status, paid_date',
        '7. lessons: id, title, content, level, category, duration',
        '8. progress: user_id, lesson_id, status (not_started/in_progress/completed), score, completed_at',
        '9. payments: user_id, amount, method, status, created_at',
        '10. support_tickets: user_id, subject, description, status, assigned_to, created_at',
    ]
    
    for table in core_tables:
        doc.add_paragraph(table)
    
    doc.add_paragraph()
    doc.add_paragraph('💡 NOTE: Đây là schema cơ bản. Sẽ có thêm nhiều tables khi phát triển.')
    
    doc.add_page_break()
    
    # ==================== KẾT LUẬN ====================
    add_heading_custom(doc, 'KẾT LUẬN', level=1, color=RGBColor(220, 38, 38))
    
    p = doc.add_paragraph()
    run = p.add_run('🎯 TỔNG KẾT:\n\n')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    doc.add_paragraph(
        'Roadmap này chi tiết hóa TẤT CẢ công việc cần làm để biến chiến lược dual-track thành '
        'hiện thực. Với 77 tasks lớn nhỏ được phân chia theo 7 phần, đây là blueprint hoàn chỉnh '
        'cho 3 năm phát triển sản phẩm.\n'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph('📊 TỔNG QUAN CON SỐ:')
    summary_stats = [
        '• Tổng số features: 77 tasks',
        '• Timeline: 36 tháng',
        '• Team growth: 14 → 101 người',
        '• Investment: ~10-15M USD total qua 3 năm',
        '• Expected outcome: 25k users, 87 tỷ ARR, #1 CLIL platform Vietnam',
    ]
    
    for stat in summary_stats:
        doc.add_paragraph(stat, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('⚡ 3 ĐIỀU QUAN TRỌNG NHẤT:\n\n')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(220, 38, 38)
    
    critical_points = [
        '1. FOCUS ON B2B FIRST (tháng 1-6):',
        '   Commission engine phải hoàn hảo. Đây là tim của mô hình.',
        '   Không rush qua B2C nếu B2B chưa vững.',
        '',
        '2. HIRE SMART, HIRE SLOW:',
        '   Mỗi người early hire ảnh hưởng lớn đến culture và output.',
        '   Ưu tiên senior engineers có experience scale startups.',
        '',
        '3. DATA-DRIVEN DECISIONS:',
        '   Build analytics từ ngày 1. Mọi quyết định phải dựa trên data.',
        '   A/B test mọi thứ có thể. Trust the numbers, not opinions.',
    ]
    
    for point in critical_points:
        if point.startswith(('1.', '2.', '3.')):
            p = doc.add_paragraph()
            run = p.add_run(point)
            run.font.bold = True
        else:
            doc.add_paragraph(point)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LET\'S BUILD SOMETHING AMAZING! 🚀\n\n')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(16, 185, 129)
    
    p = doc.add_paragraph()
    run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    print("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    print("📄 BỔ SUNG CÁC PHẦN CÒN LẠI")
    print("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n")
    
    filepath = '/Users/binhnguyen/Downloads/Engquest3k/LEXIO_Product_Development_Roadmap_20260507_2239.docx'
    
    try:
        doc = Document(filepath)
        print(f"✅ Đã mở file: {filepath}")
        print(f"📊 File hiện có: {len(doc.paragraphs)} đoạn, {len(doc.tables)} bảng\n")
        
        add_remaining_sections(doc)
        
        doc.save(filepath)
        
        print(f"\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
        print(f"✅ ĐÃ BỔ SUNG THÀNH CÔNG!")
        print(f"📄 File: {filepath}")
        print(f"📊 Tổng số đoạn văn: {len(doc.paragraphs)}")
        print(f"📊 Tổng số bảng: {len(doc.tables)}")
        print(f"📄 Ước tính: {len(doc.paragraphs) // 7}-{len(doc.paragraphs) // 5} trang")
        print(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
        
    except Exception as e:
        print(f"❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()
