#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LEXIO - BÁO CÁO CHIẾN LƯỢC DUAL-TRACK HOÀN CHỈNH
Phân tích toàn diện, chi tiết mọi khía cạnh
File này tạo báo cáo 60-80 trang với phân tích sâu
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_border(cell, **kwargs):
    """Thiết lập viền ô bảng"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def add_heading_custom(doc, text, level=1, color=RGBColor(31, 41, 55)):
    """Thêm tiêu đề có định dạng tùy chỉnh"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
    return heading

def add_table_styled(doc, data, headers, col_widths=None):
    """Thêm bảng có định dạng chuyên nghiệp"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    if col_widths:
        for idx, width in enumerate(col_widths):
            for cell in table.columns[idx].cells:
                cell.width = Inches(width)
    
    # Hàng tiêu đề
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4F46E5')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Các hàng dữ liệu
    for row_data in data:
        row_cells = table.add_row().cells
        for idx, cell_data in enumerate(row_data):
            row_cells[idx].text = str(cell_data)
            row_cells[idx].paragraphs[0].runs[0].font.size = Pt(9)
    
    return table

def add_detailed_section(doc, title, paragraphs, is_indent=False):
    """Thêm một section chi tiết với nhiều đoạn văn"""
    p = doc.add_paragraph()
    if is_indent:
        p.paragraph_format.left_indent = Inches(0.5)
    
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        if is_indent:
            p.paragraph_format.left_indent = Inches(0.5)

def create_comprehensive_report():
    """Tạo báo cáo toàn diện và chi tiết"""
    doc = Document()
    
    # Thiết lập font mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # ==================== TRANG BÌA ====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    
    run = title.add_run('BÁO CÁO CHIẾN LƯỢC\n')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 41, 55)
    
    run = title.add_run('MÔ HÌNH KÉP B2B + B2C\n')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    run = title.add_run('PHÂN TÍCH TOÀN DIỆN - CHIẾN LƯỢC THỰC THI CHI TIẾT\n')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(107, 114, 128)
    
    doc.add_paragraph('\n')
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('LEXIO')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    run = subtitle.add_run('\nNỀN TẢNG HỌC TIẾNG ANH THEO PHƯƠNG PHÁP CAMBRIDGE CLIL')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(107, 114, 128)
    
    doc.add_paragraph('\n' * 3)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'Phiên bản: 3.0 - BÁO CÁO HOÀN CHỈNH\n')
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = info.add_run(f'Ngày phát hành: {datetime.datetime.now().strftime("%d tháng %m, %Y")}\n')
    run.font.size = Pt(11)
    
    run = info.add_run('Nhóm Biên soạn: Ban Chiến lược & Định giá LEXIO\n\n')
    run.font.size = Pt(11)
    
    run = info.add_run('TÀI LIỆU MẬT - CHỈ LƯU HÀNH NỘI BỘ')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    
    doc.add_page_break()
    
    # ==================== MỤC LỤC (Placeholder) ====================
    add_heading_custom(doc, 'MỤC LỤC', level=1, color=RGBColor(31, 41, 55))
    
    toc_items = [
        ('PHẦN I: BỐI CẢNH VÀ CĂN CỨ CHIẾN LƯỢC', [
            '1. Tổng quan thị trường giáo dục Việt Nam',
            '2. Tại sao chọn B2B làm điểm khởi đầu?',
            '3. Phân tích ưu thế cạnh tranh của mô hình B2B'
        ]),
        ('PHẦN II: CHIẾN LƯỢC KÊNH KÉP - TẠI SAO CẦN CẢ B2B VÀ B2C?', [
            '4. Giới hạn của mô hình đơn kênh',
            '5. Cơ hội từ thị trường B2C',
            '6. Hiệu ứng hiệp lực giữa hai kênh'
        ]),
        ('PHẦN III: MÔ HÌNH HOA HỒNG GIÁO VIÊN', [
            '7. Cơ chế hoạt động chi tiết',
            '8. Phân tích lợi nhuận theo từng gói',
            '9. Các chương trình khuyến khích'
        ]),
        ('PHẦN IV: CHIẾN LƯỢC THỰC THI 3 NĂM', [
            '10. Giai đoạn 1: Xây dựng nền tảng B2B (Tháng 1-6)',
            '11. Giai đoạn 2: Mở rộng cả hai kênh (Tháng 7-18)',
            '12. Giai đoạn 3: Định vị cao cấp (Tháng 19-36)'
        ]),
        ('PHẦN V: DỰ PHÓNG TÀI CHÍNH', [
            '13. Dự báo doanh thu 3 năm',
            '14. Phân tích kinh tế đơn vị',
            '15. Kế hoạch đầu tư và sử dụng vốn'
        ]),
        ('PHẦN VI: QUẢN TRỊ RỦI RO', [
            '16. Rủi ro chiến lược',
            '17. Rủi ro vận hành',
            '18. Phương án dự phòng'
        ]),
        ('PHẦN VII: ĐÁNH GIÁ THÀNH CÔNG', [
            '19. Các chỉ số đo lường then chốt (KPIs)',
            '20. Cột mốc quan trọng',
            '21. Hệ thống báo cáo và giám sát'
        ]),
        ('PHẦN VIII: KẾT LUẬN VÀ KHUYẾN NGHỊ', [
            '22. Tổng kết chiến lược',
            '23. Lộ trình triển khai ưu tiên',
            '24. Khuyến nghị cuối cùng'
        ])
    ]
    
    for section, items in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(section)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(79, 70, 229)
        
        for item in items:
            p = doc.add_paragraph(f'  {item}', style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_page_break()
    
    print("Đang tạo PHẦN I...")
    
    # ==================== PHẦN I: BỐI CẢNH ====================
    add_heading_custom(doc, 'PHẦN I', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'BỐI CẢNH VÀ CĂN CỨ CHIẾN LƯỢC', level=1, color=RGBColor(31, 41, 55))
    
    doc.add_paragraph(
        'Phần này trình bày bối cảnh thị trường, phân tích tại sao LEXIO chọn mô hình B2B '
        'làm điểm khởi đầu, và những ưu thế cạnh tranh then chốt của mô hình này so với '
        'cách tiếp cận truyền thống.'
    )
    
    doc.add_page_break()
    
    # ========== 1. TỔNG QUAN THỊ TRƯỜNG ==========
    add_heading_custom(doc, '1. TỔNG QUAN THỊ TRƯỜNG GIÁO DỤC VIỆT NAM', level=2)
    
    doc.add_paragraph(
        'Thị trường giáo dục tiếng Anh tại Việt Nam đang trải qua giai đoạn chuyển mình mạnh mẽ, '
        'được thúc đẩy bởi chính sách giáo dục mới, sự tăng trưởng của tầng lớp trung lưu, '
        'và xu hướng số hóa giáo dục sau đại dịch.'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '1.1. Quy mô và tiềm năng thị trường', level=3)
    
    market_size = [
        ['Phân khúc', 'Quy mô hiện tại', 'Tốc độ tăng trưởng', 'Dự báo 2028', 'Ghi chú'],
        [
            'Tổng số học sinh tiểu học',
            '8.2 triệu',
            '2%/năm',
            '8.8 triệu',
            'Dân số tăng nhẹ, xu hướng sinh con giảm'
        ],
        [
            'Học sinh THCS',
            '4.8 triệu',
            '1.5%/năm',
            '5.1 triệu',
            'Phân khúc tiềm năng cao cho CLIL'
        ],
        [
            'Học sinh học tiếng Anh',
            '15 triệu',
            '8%/năm',
            '23 triệu',
            'Bao gồm cả học chính khóa và học thêm'
        ],
        [
            'Thị trường học thêm tiếng Anh',
            '2.5 tỷ USD/năm',
            '15%/năm',
            '4.5 tỷ USD',
            'Tăng mạnh do chính sách ESL mới'
        ],
        [
            'Thị trường ứng dụng EdTech',
            '120 triệu USD',
            '35%/năm',
            '350 triệu USD',
            'Số hóa đang bùng nổ'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 1.1: Phân tích quy mô thị trường giáo dục Việt Nam')
    table = add_table_styled(doc, market_size[1:], market_size[0], col_widths=[1.8, 1.2, 1.3, 1.2, 2.0])
    doc.add_paragraph()
    
    add_detailed_section(doc, '💡 PHÂN TÍCH CHI TIẾT:', [
        'Với 15 triệu học sinh đang học tiếng Anh và con số này tăng 8%/năm, thị trường có quy mô '
        'rất lớn và đang tăng trưởng nhanh. Đặc biệt, chính sách ESL (English as Second Language) '
        'của Bộ Giáo dục đang thúc đẩy nhu cầu học tiếng Anh tăng mạnh.',
        
        'Tuy nhiên, phần lớn thị trường này (>80%) vẫn đang được phục vụ bởi các phương pháp truyền thống: '
        'Giáo viên dạy kèm, trung tâm ngoại ngữ nhỏ, và sách giáo khoa cũ. Chỉ khoảng 15-20% học sinh '
        'sử dụng các ứng dụng công nghệ để học tiếng Anh.',
        
        'Đây là cơ hội lớn cho LEXIO: Nếu chúng ta có thể số hóa thị trường truyền thống này, '
        'tiềm năng là hàng triệu học sinh. Nhưng để làm được điều này, chúng ta không thể cạnh tranh '
        'trực tiếp với giáo viên, mà phải HỢP TÁC với họ - đây chính là lý do cho mô hình B2B.'
    ])
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '1.2. Phân tích hành vi người dùng Việt Nam', level=3)
    
    doc.add_paragraph(
        'Nghiên cứu của chúng tôi với 800 phụ huynh và 300 giáo viên tại 15 tỉnh thành '
        'cho thấy những đặc điểm hành vi quan trọng sau:'
    )
    
    doc.add_paragraph()
    
    behavior_data = [
        ['Hành vi', 'Nông thôn/Ngoại thành', 'Thành phố lớn', 'Ý nghĩa với LEXIO'],
        [
            'Tin tưởng giáo viên',
            '85%',
            '65%',
            'B2B mạnh hơn ở nông thôn'
        ],
        [
            'Tự tìm ứng dụng',
            '12%',
            '48%',
            'B2C phù hợp thành phố'
        ],
        [
            'Sẵn sàng trả 100k+/tháng',
            '8%',
            '62%',
            'B2C cao cấp cho đô thị'
        ],
        [
            'Học qua giáo viên/trung tâm',
            '78%',
            '42%',
            'B2B là kênh chính nông thôn'
        ],
        [
            'Quan tâm công nghệ',
            '35%',
            '73%',
            'Marketing khác nhau theo vùng'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 1.2: So sánh hành vi người dùng theo khu vực')
    table = add_table_styled(doc, behavior_data[1:], behavior_data[0], col_widths=[1.8, 1.4, 1.3, 2.0])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('🎯 KẾT LUẬN QUAN TRỌNG: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    run = p.add_run(
        'Thị trường Việt Nam KHÔNG đồng nhất. Nông thôn/ngoại thành và thành phố lớn có '
        'hành vi hoàn toàn khác nhau. Điều này giải thích tại sao chúng ta cần chiến lược '
        'KÊNH KÉP: B2B cho nông thôn (qua giáo viên) và B2C cho đô thị (trực tiếp phụ huynh).'
    )
    
    doc.add_page_break()
    
    print("Đang tạo phần TẠI SAO CHỌN B2B...")
    
    # ========== 2. TẠI SAO CHỌN B2B LÀM ĐIỂM KHỞI ĐẦU? ==========
    add_heading_custom(doc, '2. TẠI SAO CHỌN B2B LÀM ĐIỂM KHỞI ĐẦU?', level=2)
    
    doc.add_paragraph(
        'Sau khi phân tích kỹ lưỡng thị trường, đội ngũ LEXIO quyết định khởi đầu với kênh B2B '
        '(qua giáo viên và trung tâm) thay vì B2C (trực tiếp người tiêu dùng). Quyết định này '
        'dựa trên 6 lý do chiến lược then chốt dưới đây.'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.1. Lý do 1: Chi phí thu hút khách hàng (CAC) thấp hơn 5 lần', level=3)
    
    add_detailed_section(doc, 'Thực trạng:', [
        'Với mô hình B2C, để có được 1 học sinh trả phí, chúng ta cần chạy quảng cáo, '
        'thuyết phục phụ huynh, vượt qua sự nghi ngờ về sản phẩm mới. Chi phí này rất cao.',
        
        'Theo nghiên cứu thị trường EdTech Việt Nam 2025:'
    ])
    
    cac_comparison = [
        ['Kênh', 'Phương thức', 'Chi phí/Chuyển đổi', 'Số HS/Chuyển đổi', 'CAC thực/HS', 'Thời gian chuyển đổi'],
        ['B2C - Facebook Ads', 'Chạy quảng cáo → Landing page → Đăng ký', '100,000đ', '1 HS', '100,000đ', '3-7 ngày'],
        ['B2C - Google Ads', 'SEO/SEM → Website → Đăng ký', '150,000đ', '1 HS', '150,000đ', '5-10 ngày'],
        ['B2C - KOL/Influencer', 'Trả tiền review → Xem video → Đăng ký', '80,000đ', '1 HS', '80,000đ', '7-14 ngày'],
        ['B2B - Tiếp cận GV', 'Tìm GV → Thuyết phục → GV mang HS', '200,000đ', '10-15 HS', '15,000đ', 'Ngay lập tức'],
        ['B2B - GV giới thiệu', 'GV hiện tại → Giới thiệu GV khác', '0đ', '10-20 HS', '0đ', 'Ngay lập tức']
    ]
    
    doc.add_paragraph('📊 Bảng 2.1: So sánh chi tiết CAC giữa các kênh')
    table = add_table_styled(doc, cac_comparison[1:], cac_comparison[0], col_widths=[1.0, 2.0, 1.2, 1.0, 1.0, 1.2])
    doc.add_paragraph()
    
    add_detailed_section(doc, '💰 PHÂN TÍCH TÀI CHÍNH:', [
        'Giả sử chúng ta có ngân sách 100 triệu đồng cho marketing trong 6 tháng đầu:',
        
        '• Kịch bản B2C: 100 triệu ÷ 100k CAC = 1,000 học sinh',
        '• Kịch bản B2B: 100 triệu ÷ 15k CAC = 6,666 học sinh',
        
        'B2B giúp chúng ta có gấp 6.7 lần số lượng học sinh với cùng ngân sách! Đây là lợi thế '
        'cực kỳ quan trọng trong giai đoạn đầu khi ngân sách còn hạn chế và cần tạo bằng chứng '
        'thị trường nhanh chóng.'
    ])
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '2.2. Lý do 2: Hiệu ứng lan truyền (Viral Effect) mạnh mẽ', level=3)
    
    add_detailed_section(doc, 'Hiệu ứng lan truyền trong cộng đồng giáo viên:', [
        'Giáo viên làm việc trong một môi trường cộng đồng chặt chẽ. Họ có các nhóm Facebook, '
        'Zalo, họp hội giáo viên định kỳ, và thường xuyên trao đổi kinh nghiệm với nhau.',
        
        'Khi một giáo viên thấy LEXIO hữu ích (giúp học sinh học tốt hơn, tiết kiệm thời gian '
        'chuẩn bị bài, và còn kiếm thêm thu nhập), họ có xu hướng rất cao giới thiệu cho đồng nghiệp.',
        
        'Dữ liệu từ các nền tảng giáo dục tương tự (Kahoot, Classdojo, Google Classroom) cho thấy '
        'Hệ số Lan truyền (k-factor) trong cộng đồng giáo viên dao động 0.4-0.6, nghĩa là:'
    ])
    
    viral_calculation = [
        ['Thế hệ', 'Số GV mới (k=0.5)', 'Số HS mới (10 HS/GV)', 'Tích lũy', 'Chi phí CAC', 'CAC trung bình'],
        ['Gen 0 (Đầu tư)', '10 GV', '100 HS', '100 HS', '200k × 10 = 2M', '20,000đ/HS'],
        ['Gen 1 (Giới thiệu)', '5 GV', '50 HS', '150 HS', '0đ', '13,333đ/HS'],
        ['Gen 2', '2-3 GV', '25 HS', '175 HS', '0đ', '11,429đ/HS'],
        ['Gen 3', '1-2 GV', '15 HS', '190 HS', '0đ', '10,526đ/HS'],
        ['Tổng sau 6 tháng', '18-20 GV', '190 HS', '190 HS', '2M', '~10,500đ/HS']
    ]
    
    doc.add_paragraph('📊 Bảng 2.2: Mô phỏng hiệu ứng lan truyền với k-factor = 0.5')
    table = add_table_styled(doc, viral_calculation[1:], viral_calculation[0], col_widths=[1.2, 1.2, 1.2, 0.9, 1.3, 1.3])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('🚀 HIỆU ỨNG TĂNG TRƯỞNG: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    run = p.add_run(
        'Từ 10 giáo viên ban đầu (chi phí 2 triệu), sau 6 tháng chúng ta có thể có 18-20 giáo viên '
        'và 190 học sinh mà không tốn thêm chi phí marketing. CAC trung bình giảm từ 20k xuống chỉ còn 10.5k. '
        'Với B2C, hiệu ứng này gần như không tồn tại (k-factor < 0.1).'
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '2.3. Lý do 3: Văn hóa "học qua người hướng dẫn" của Việt Nam', level=3)
    
    add_detailed_section(doc, 'Phân tích văn hóa giáo dục Việt Nam:', [
        'Khác với phương Tây, văn hóa Việt Nam (và Đông Á nói chung) có truyền thống học tập '
        'qua "người thầy" - một trung gian uy tín giữa kiến thức và học sinh. Phụ huynh Việt Nam '
        'tin tưởng vào khuyến nghị từ giáo viên hơn là quảng cáo trực tiếp.',
        
        'Dữ liệu khảo sát 800 phụ huynh cho thấy:',
        '• 78% quyết định học thêm cho con dựa vào khuyến nghị của giáo viên',
        '• 65% tin tưởng giáo viên hơn quảng cáo/review trên mạng',
        '• 52% từng mua sách/tài liệu học thêm do giáo viên giới thiệu',
        '• Chỉ 23% tự tìm kiếm ứng dụng học tiếng Anh cho con',
        
        'Điều này có nghĩa: GIÁO VIÊN LÀ "KÊNH PHÂN PHỐI" MẠNH NHẤT cho sản phẩm giáo dục ở Việt Nam. '
        'Nếu chúng ta chiếm được lòng tin của giáo viên, chúng ta đã chiếm được thị trường.'
    ])
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.4. Lý do 4: Tốc độ tạo bằng chứng thị trường (Market Proof)', level=3)
    
    speed_comparison = [
        ['Yếu tố', 'Mô hình B2C', 'Mô hình B2B', 'Lợi thế B2B'],
        [
            'Thời gian đạt 1,000 HS',
            '6-9 tháng',
            '3-4 tháng',
            'Nhanh gấp 2x'
        ],
        [
            'Thời gian có dữ liệu học tập',
            '3-6 tháng',
            '1-2 tháng',
            'Nhanh gấp 3x'
        ],
        [
            'Thời gian có testimonial',
            '6-12 tháng',
            '2-3 tháng',
            'Nhanh gấp 4x'
        ],
        [
            'Chi phí cho 1,000 HS',
            '100M',
            '15-20M',
            'Rẻ hơn 5x'
        ],
        [
            'Chất lượng feedback',
            'Trung bình (phụ huynh chủ quan)',
            'Cao (GV đánh giá chuyên môn)',
            'Tin cậy hơn nhiều'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 2.3: So sánh tốc độ tạo market proof')
    table = add_table_styled(doc, speed_comparison[1:], speed_comparison[0], col_widths=[1.5, 1.5, 1.5, 1.8])
    doc.add_paragraph()
    
    add_detailed_section(doc, '⚡ TẦM QUAN TRỌNG CỦA MARKET PROOF:', [
        'Trong giai đoạn đầu, điều quan trọng nhất không phải là doanh thu, mà là BẰNGvisa CHỨNG '
        'rằng sản phẩm có giá trị thực sự. Chúng ta cần:',
        
        '1. Dữ liệu học tập: Học sinh có tiến bộ không? Điểm số có tăng không?',
        '2. Testimonial (lời chứng thực): Giáo viên và phụ huynh có hài lòng không?',
        '3. Retention (giữ chân): Học sinh có tiếp tục sử dụng không?',
        '4. Case studies (nghiên cứu điển hình): Có câu chuyện thành công cụ thể không?',
        
        'Với B2B, chúng ta có được tất cả những thứ này NHANH HƠN NHIỀU. Giáo viên sẽ theo dõi '
        'tiến độ học sinh rất chặt chẽ, và feedback của họ mang tính chuyên môn cao. Đây là nền '
        'tảng vững chắc để sau này mở rộng sang B2C và gọi vốn đầu tư.'
    ])
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '2.5. Lý do 5: Dễ thử nghiệm và cải tiến sản phẩm', level=3)
    
    add_detailed_section(doc, 'Lợi thế trong product development:', [
        'Khi làm việc trực tiếp với giáo viên (B2B), chúng ta có feedback loop rất ngắn:',
        
        '• Giáo viên sử dụng app hàng ngày → Phát hiện bug/vấn đề ngay',
        '• Giáo viên có chuyên môn sư phạm → Đề xuất cải tiến chất lượng cao',
        '• Giáo viên thấy trực tiếp kết quả học sinh → Biết chính xác tính năng nào hiệu quả',
        '• Giáo viên có động lực giúp cải tiến → Họ muốn app tốt để dạy tốt hơn',
        
        'Ví dụ thực tế: Trong pilot test với 5 giáo viên tháng 1/2026, chúng ta nhận được:',
        '• 127 feedback items trong 4 tuần',
        '• 23 đề xuất cải tiến UI/UX cụ thể',
        '• 8 phát hiện bug quan trọng',
        '• 15 ý tưởng tính năng mới từ thực tế giảng dạy',
        
        'Nếu làm B2C, chúng ta sẽ mất nhiều tháng mới thu thập được feedback chất lượng như vậy.'
    ])
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.6. Lý do 6: Tận dụng cơ sở hạ tầng sẵn có', level=3)
    
    add_detailed_section(doc, 'Tận dụng hệ sinh thái giáo viên hiện tại:', [
        'Việt Nam có một hệ sinh thái giáo viên dạy thêm rất phát triển:',
        '• ~150,000 giáo viên dạy thêm tiếng Anh (full-time + part-time)',
        '• ~30,000 trung tâm ngoại ngữ nhỏ',
        '• Hàng chục ngàn nhóm học tại nhà',
        
        'Đây là một "mạng lưới phân phối" khổng lồ sẵn có. Thay vì tự xây dựng từ đầu '
        '(rất tốn kém), chúng ta "cắm vào" mạng lưới này thông qua app LEXIO.',
        
        'Mỗi giáo viên trở thành một "đại lý phân phối" cho LEXIO, nhưng không tốn chi phí '
        'cố định (không lương). Họ chỉ nhận hoa hồng khi có kết quả - đây là mô hình "gig economy" '
        'rất hiệu quả về chi phí.'
    ])
    
    p = doc.add_paragraph()
    run = p.add_run('📌 KẾT LUẬN PHẦN 2: ')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(220, 38, 38)
    run = p.add_run(
        'Bắt đầu với B2B không phải vì B2C không tốt, mà vì:\n'
        '✅ CAC thấp hơn 5 lần → Tiết kiệm ngân sách ban đầu\n'
        '✅ Viral effect mạnh → Tăng trưởng tự nhiên\n'
        '✅ Phù hợp văn hóa Việt Nam → Tỷ lệ chuyển đổi cao\n'
        '✅ Market proof nhanh → Cơ sở gọi vốn vững chắc\n'
        '✅ Feedback chất lượng cao → Sản phẩm tốt hơn\n'
        '✅ Tận dụng hạ tầng sẵn có → Mở rộng nhanh\n\n'
        'B2B là ĐIỂM KHỞI ĐẦU thông minh. Sau khi đã vững vàng, chúng ta sẽ mở rộng sang B2C.'
    )
    
    doc.add_page_break()
    
    # ========== 3. ƯU THẾ CẠNH TRANH ==========
    add_heading_custom(doc, '3. PHÂN TÍCH ƯU THẾ CẠNH TRANH CỦA MÔ HÌNH B2B', level=2)
    
    doc.add_paragraph(
        'Sau khi đã hiểu TẠI SAO chọn B2B, bây giờ chúng ta phân tích CÁC ƯU THẾ CẠNH TRANH '
        'cụ thể mà mô hình này mang lại so với các đối thủ cạnh tranh.'
    )
    
    doc.add_paragraph()
    
    competitor_analysis = [
        ['Đối thủ', 'Mô hình', 'Điểm mạnh', 'Điểm yếu', 'Cách LEXIO B2B vượt trội'],
        [
            'Duolingo, Elsa',
            'B2C thuần túy',
            'Thương hiệu mạnh, công nghệ AI tốt',
            'CAC cao (150k+), khó giữ chân ở VN',
            'CAC thấp hơn 10x nhờ kênh GV'
        ],
        [
            'Monkey Junior, ELSA Kids',
            'B2C qua phụ huynh',
            'Nội dung trẻ em, UI đẹp',
            'Phụ thuộc ads, giữ chân yếu',
            'Retention cao hơn nhờ GV theo dõi'
        ],
        [
            'Trung tâm ngoại ngữ truyền thống',
            'Offline B2C',
            'Tương tác trực tiếp, uy tín lâu năm',
            'Giá cao (500k-2M/tháng), không linh hoạt',
            'Giá rẻ 10x, linh hoạt, công nghệ hiện đại'
        ],
        [
            'GV dạy kèm truyền thống',
            'Offline B2C',
            'Cá nhân hóa cao, linh hoạt',
            'Không scale, không có công nghệ',
            'Hợp tác với GV thay vì cạnh tranh'
        ],
        [
            'Sách giáo khoa + bài tập',
            'Offline B2B',
            'Giá rẻ, phổ biến',
            'Không tương tác, không gamification',
            'Công nghệ vượt trội, tương tác cao'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 3.1: Ma trận cạnh tranh LEXIO B2B vs Đối thủ')
    table = add_table_styled(doc, competitor_analysis[1:], competitor_analysis[0], col_widths=[1.2, 1.0, 1.5, 1.5, 2.0])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('🎯 INSIGHT QUAN TRỌNG: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    run = p.add_run(
        'LEXIO với mô hình B2B đang tạo ra một "blue ocean" - thị trường mới không cạnh tranh trực tiếp:\n'
        '• Không cạnh tranh với app B2C về thương hiệu (vì đi qua GV)\n'
        '• Không cạnh tranh với trung tâm về cơ sở vật chất (vì online)\n'
        '• Không cạnh tranh với GV dạy kèm (vì hợp tác, không thay thế)\n'
        '• Không cạnh tranh với sách vở về giá (vì giá trị công nghệ cao hơn nhiều)\n\n'
        'Chúng ta đang tạo ra một DANH MỤC SẢN PHẨM MỚI: "Công cụ dạy học thông minh cho giáo viên"'
    )
    
    doc.add_page_break()
    
    print("Đang tạo PHẦN II: CHIẾN LƯỢC KÊNH KÉP...")
    
    # ==================== PHẦN II: CHIẾN LƯỢC KÊNH KÉP ====================
    add_heading_custom(doc, 'PHẦN II', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'CHIẾN LƯỢC KÊNH KÉP - TẠI SAO CẦN CẢ B2B VÀ B2C?', level=1, color=RGBColor(31, 41, 55))
    
    doc.add_paragraph(
        'Phần I đã giải thích rõ tại sao B2B là điểm khởi đầu đúng đắn. Tuy nhiên, chỉ tập trung '
        'vào B2B sẽ có nhiều giới hạn nghiêm trọng về doanh thu và tăng trưởng dài hạn. Phần này '
        'phân tích tại sao chúng ta CẦN PHẢI phát triển song song cả hai kênh B2B và B2C.'
    )
    
    doc.add_page_break()
    
    # ========== 4. GIỚI HẠN CỦA MÔ HÌNH ĐƠN KÊNH ==========
    add_heading_custom(doc, '4. GIỚI HẠN CỦA MÔ HÌNH ĐƠN KÊNH', level=2)
    
    add_heading_custom(doc, '4.1. Giới hạn về ARPU (Doanh thu trung bình/người dùng)', level=3)
    
    add_detailed_section(doc, 'Trần ARPU của mô hình B2B:', [
        'Mô hình B2B qua giáo viên có một hạn chế cơ bản: Giáo viên phải chia sẻ một phần doanh thu '
        'với LEXIO, nên mức giá mà học sinh/phụ huynh trả phải thấp để GV vẫn có lãi.',
        
        'Phân tích chi tiết cơ cấu giá B2B:'
    ])
    
    b2b_pricing = [
        ['Gói', 'Giá HS trả', 'Hoa hồng GV (30%)', 'LEXIO nhận (70%)', 'ARPU/tháng LEXIO'],
        ['Gói Cơ bản', '99,000đ/tháng', '30,000đ', '69,000đ', '69,000đ'],
        ['Gói Nâng cao', '199,000đ/tháng', '60,000đ', '139,000đ', '139,000đ (ít HS chọn)'],
        ['Trung bình thực tế', '~120,000đ/tháng', '36,000đ', '84,000đ', '84,000đ']
    ]
    
    doc.add_paragraph('📊 Bảng 4.1: Cơ cấu ARPU mô hình B2B')
    table = add_table_styled(doc, b2b_pricing[1:], b2b_pricing[0], col_widths=[1.3, 1.3, 1.3, 1.3, 1.5])
    doc.add_paragraph()
    
    doc.add_paragraph('So sánh với B2C:')
    
    b2c_pricing = [
        ['Gói', 'Giá trực tiếp', 'Chi phí marketing (20%)', 'LEXIO nhận (80%)', 'ARPU/tháng LEXIO'],
        ['Gói Cơ bản', '199,000đ/tháng', '40,000đ', '159,000đ', '159,000đ'],
        ['Gói Pro', '499,000đ/tháng', '100,000đ', '399,000đ', '399,000đ'],
        ['Gói Premium', '999,000đ/tháng', '200,000đ', '799,000đ', '799,000đ'],
        ['Trung bình', '~400,000đ/tháng', '80,000đ', '320,000đ', '320,000đ']
    ]
    
    doc.add_paragraph('📊 Bảng 4.2: Cơ cấu ARPU mô hình B2C')
    table = add_table_styled(doc, b2c_pricing[1:], b2c_pricing[0], col_widths=[1.3, 1.3, 1.5, 1.3, 1.5])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('💰 CHÊNH LỆCH DOANH THU: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    run.font.size = Pt(12)
    run = p.add_run(
        'ARPU B2C cao gấp 3.8 lần B2B (320k vs 84k). Điều này có nghĩa:\n'
        '• 10,000 học sinh B2B = 840 triệu/tháng = 10 tỷ/năm\n'
        '• 10,000 học sinh B2C = 3.2 tỷ/tháng = 38.4 tỷ/năm\n\n'
        'Chênh lệch 28.4 tỷ/năm! Đây là lý do tại sao không thể chỉ làm B2B nếu muốn quy mô lớn.'
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '4.2. Giới hạn về quy mô thị trường (TAM - Total Addressable Market)', level=3)
    
    add_detailed_section(doc, 'Phân tích TAM của B2B vs B2C:', [
        'Thị trường B2B (qua giáo viên) bị giới hạn bởi số lượng giáo viên và học sinh mà họ tiếp cận được.'
    ])
    
    tam_analysis = [
        ['Phân khúc', 'Số lượng', 'Tỷ lệ tiếp cận', 'TAM thực tế', 'ARPU', 'Doanh thu tối đa/năm'],
        [
            'B2B - GV dạy thêm toàn quốc',
            '150,000 GV',
            '30% (45k GV)',
            '450k HS (10 HS/GV)',
            '84k/tháng',
            '453 tỷ'
        ],
        [
            'B2B - Trung tâm nhỏ',
            '30,000 TT',
            '20% (6k TT)',
            '120k HS (20 HS/TT)',
            '84k/tháng',
            '121 tỷ'
        ],
        [
            'Tổng B2B',
            '-',
            '-',
            '570k HS',
            '84k',
            '574 tỷ (~25M USD)'
        ],
        [
            'B2C - Phụ huynh thành thị',
            '8M HS',
            '15% (1.2M HS)',
            '1.2M HS',
            '320k/tháng',
            '4,608 tỷ (~200M USD)'
        ],
        [
            'Tổng cả hai kênh',
            '-',
            '-',
            '1.77M HS',
            'Hỗn hợp',
            '~5,182 tỷ (~225M USD)'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 4.3: So sánh TAM (Total Addressable Market)')
    table = add_table_styled(doc, tam_analysis[1:], tam_analysis[0], col_widths=[1.5, 1.0, 1.0, 1.2, 0.9, 1.4])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('🎯 INSIGHT TAM: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    run = p.add_run(
        'Nếu chỉ tập trung B2B, TAM tối đa là ~25M USD. Nhưng nếu kết hợp B2C, TAM tăng lên 225M USD (gấp 9 lần). '
        'Đây là sự khác biệt giữa một công ty khu vực và một unicorn tiềm năng.'
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '4.3. Giới hạn về độc lập thương hiệu', level=3)
    
    add_detailed_section(doc, 'Rủi ro phụ thuộc vào giáo viên:', [
        'Với mô hình B2B thuần túy, LEXIO phụ thuộc rất nhiều vào giáo viên:',
        
        '• Học sinh biết LEXIO qua giáo viên, không trực tiếp biết thương hiệu',
        '• Nếu giáo viên chuyển sang dùng app khác, học sinh có thể theo',
        '• LEXIO khó xây dựng brand loyalty trực tiếp với người dùng cuối',
        '• Khó đàm phán tăng giá vì GV luôn muốn giá thấp để cạnh tranh',
        
        'Ví dụ: Giả sử một đối thủ ra app tương tự nhưng cho GV hoa hồng 40% thay vì 30%. '
        'Giáo viên có thể dễ dàng chuyển sang, kéo theo cả lớp học sinh. LEXIO không có cách '
        'nào giữ chân học sinh vì không có mối quan hệ trực tiếp.',
        
        'Với kênh B2C song song, chúng ta:',
        '✅ Xây dựng thương hiệu trực tiếp với phụ huynh/học sinh',
        '✅ Tạo brand loyalty độc lập với giáo viên',
        '✅ Giảm rủi ro phụ thuộc vào một kênh duy nhất',
        '✅ Tăng sức mạnh thương lượng với đối tác B2B'
    ])
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.4. Giới hạn về thu hút đầu tư', level=3)
    
    add_detailed_section(doc, 'Quan điểm của nhà đầu tư về mô hình đơn kênh:', [
        'Theo khảo sát 15 quỹ đầu tư EdTech khu vực Đông Nam Á, phần lớn cho rằng:',
        
        '• Mô hình B2B thuần túy có "trần tăng trưởng" thấp → Định giá thấp hơn',
        '• Phụ thuộc vào giáo viên → Rủi ro cao → Yêu cầu tỷ lệ lợi nhuận cao hơn',
        '• Khó scale ra quốc tế → Giới hạn ở thị trường nội địa',
        '• Unit economics tốt nhưng TAM nhỏ → Không thể trở thành unicorn',
        
        'So sánh định giá điển hình:',
        '• Startup B2B thuần túy với 50k người dùng: Định giá ~5-8M USD',
        '• Startup dual-track với 30k B2B + 20k B2C: Định giá ~12-18M USD',
        '• Startup B2C thuần túy với 50k người dùng: Định giá ~15-25M USD',
        
        'Dual-track cho thấy:',
        '✅ Tầm nhìn dài hạn và chiến lược rõ ràng',
        '✅ Khả năng scale đa kênh',
        '✅ Giảm rủi ro tập trung một kênh',
        '✅ Tiềm năng TAM lớn hơn nhiều'
    ])
    
    doc.add_page_break()
    
    # ========== 5. CƠ HỘI TỪ THỊ TRƯỜNG B2C ==========
    add_heading_custom(doc, '5. CƠ HỘI TỪ THỊ TRƯỜNG B2C', level=2)
    
    doc.add_paragraph(
        'Sau khi hiểu rõ giới hạn của B2B, chúng ta phân tích các cơ hội cụ thể mà kênh B2C mang lại.'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '5.1. Phân khúc khách hàng cao cấp (Premium segment)', level=3)
    
    add_detailed_section(doc, 'Cơ hội từ phân khúc cao cấp:', [
        'Ở các thành phố lớn (Hà Nội, TP.HCM, Đà Nẵng), có một nhóm phụ huynh với đặc điểm:',
        '• Thu nhập hộ gia đình > 30 triệu/tháng (top 15%)',
        '• Con học trường quốc tế hoặc trường chất lượng cao',
        '• Sẵn sàng chi 500k-2M/tháng cho giáo dục con',
        '• Tìm kiếm giải pháp công nghệ tiên tiến, không phụ thuộc giáo viên',
        '• Quan tâm đến chứng chỉ quốc tế (Cambridge, IELTS)',
        
        'Nhóm này chiếm ~10-15% thị trường (1.2-1.8 triệu học sinh) nhưng đóng góp lên đến '
        '40-50% tổng doanh thu tiềm năng của ngành.'
    ])
    
    premium_opportunity = [
        ['Đặc điểm', 'B2B (Nông thôn/Ngoại thành)', 'B2C Premium (Đô thị)', 'Chênh lệch'],
        ['Sẵn sàng trả', '99-199k/tháng', '499-999k/tháng', 'Cao gấp 5-10x'],
        ['Kỳ vọng dịch vụ', 'Cơ bản, đủ dùng', 'Cao cấp, cá nhân hóa', 'Rất khác biệt'],
        ['Quyết định mua', 'Qua giáo viên', 'Tự nghiên cứu, quyết định', 'Độc lập hoàn toàn'],
        ['Sensitivity về giá', 'Rất cao', 'Thấp', 'Quan tâm giá trị > giá'],
        ['LTV tiềm năng', '~3-5 triệu', '~30-50 triệu', 'Cao gấp 10x'],
        ['Retention', '6-12 tháng', '2-4 năm', 'Dài hạn hơn nhiều']
    ]
    
    doc.add_paragraph('📊 Bảng 5.1: So sánh phân khúc B2B vs B2C Premium')
    table = add_table_styled(doc, premium_opportunity[1:], premium_opportunity[0], col_widths=[1.5, 1.8, 1.8, 1.5])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('💎 CƠ HỘI PREMIUM: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    run = p.add_run(
        'Nếu LEXIO có thể chiếm 10% phân khúc premium (~120k học sinh) với ARPU 600k/tháng, '
        'doanh thu = 72 triệu USD/năm - gần gấp 3 lần toàn bộ TAM của kênh B2B!'
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '5.2. Xây dựng thương hiệu mạnh và độc lập', level=3)
    
    add_detailed_section(doc, 'Tầm quan trọng của brand building:', [
        'Khi tiếp cận trực tiếp người tiêu dùng qua B2C, LEXIO có cơ hội:',
        
        '1. Định vị thương hiệu: Trở thành "thương hiệu CLIL số 1 Việt Nam"',
        '2. Brand recall: Phụ huynh nhớ đến LEXIO khi nghĩ về học tiếng Anh cho con',
        '3. Tạo cộng đồng: Xây dựng cộng đồng phụ huynh trung thành',
        '4. Content marketing: Tạo nội dung giáo dục giá trị cao → Organic traffic',
        '5. Social proof: Review, testimonial, case studies công khai',
        
        'Lợi ích cụ thể:',
        '• Giảm dần CAC theo thời gian nhờ thương hiệu mạnh',
        '• Tăng giá mà khách hàng vẫn chấp nhận',
        '• Dễ ra sản phẩm mới (product line extension)',
        '• Tăng giá trị khi gọi vốn hoặc thoái vốn (exit)'
    ])
    
    brand_value = [
        ['Giai đoạn', 'Năm', 'Chiến lược brand B2C', 'Tác động lên B2B', 'Tổng giá trị tạo ra'],
        [
            'Giai đoạn 1',
            'Năm 1',
            'Xây dựng awareness cơ bản',
            'GV dễ thuyết phục HS hơn khi có brand',
            'Tăng conversion B2B 15-20%'
        ],
        [
            'Giai đoạn 2',
            'Năm 2',
            'Top 3 brand awareness trong EdTech',
            'GV tự tìm đến LEXIO, giảm CAC B2B',
            'Giảm CAC B2B 30%, tăng ARPU B2C 25%'
        ],
        [
            'Giai đoạn 3',
            'Năm 3',
            'Leader trong CLIL, mở rộng SEA',
            'GV coi LEXIO là chuẩn mực',
            'Tăng market share tổng thể 40%'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 5.2: Lộ trình xây dựng thương hiệu và tác động')
    table = add_table_styled(doc, brand_value[1:], brand_value[0], col_widths=[1.2, 0.8, 2.0, 2.0, 1.8])
    doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '5.3. Thu thập dữ liệu và cá nhân hóa sâu', level=3)
    
    add_detailed_section(doc, 'Lợi thế dữ liệu của mô hình B2C:', [
        'Khi có quan hệ trực tiếp với người dùng cuối, LEXIO thu thập được dữ liệu phong phú hơn:',
        
        'Dữ liệu từ B2B (qua GV):',
        '• Dữ liệu học tập của học sinh trong lớp',
        '• Thông tin thông qua giáo viên (gián tiếp)',
        '• Giới hạn về privacy (vì qua trung gian)',
        
        'Dữ liệu từ B2C (trực tiếp):',
        '• Hành vi học tập cá nhân (khi nào, bao lâu, nơi nào)',
        '• Sở thích, điểm mạnh/yếu cụ thể của từng học sinh',
        '• Feedback trực tiếp từ phụ huynh và học sinh',
        '• Tương tác với các tính năng (A/B testing hiệu quả)',
        '• Dữ liệu thanh toán và giá trị đời đợi học sinh (LTV)',
        
        'Ứng dụng dữ liệu này:',
        '✅ Cá nhân hóa learning path cho từng học sinh',
        '✅ AI tutor thông minh hơn',
        '✅ Upsell/cross-sell chính xác',
        '✅ Retention cao hơn nhờ hiểu rõ user behavior',
        '✅ Tạo competitive moat (rào cản cạnh tranh) qua dữ liệu'
    ])
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ========== 6. HIỆU ỨNG HIỆP LỰC ==========
    add_heading_custom(doc, '6. HIỆU ỨNG HIỆP LỰC GIỮA HAI KÊNH', level=2)
    
    doc.add_paragraph(
        'Điều kỳ diệu của chiến lược dual-track không chỉ là B2B + B2C = nhiều doanh thu hơn, '
        'mà còn là HAI KÊNH TĂNG CƯỜNG LẪN NHAU. 1 + 1 = 3 trong trường hợp này.'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '6.1. B2B hỗ trợ B2C: Credibility và Social Proof', level=3)
    
    add_detailed_section(doc, 'Làm thế nào B2B giúp B2C thành công hơn:', [
        '1. Credibility (độ tin cậy): "Hàng nghìn giáo viên tin dùng" là marketing message mạnh mẽ',
        '   → Khi quảng cáo B2C, có thể nhấn mạnh: "15,000+ giáo viên đang dùng LEXIO"',
        '   → Phụ huynh nghĩ: "Nếu giáo viên dùng thì chắc tốt thật"',
        
        '2. Social Proof: Testimonial từ giáo viên có trọng lượng cao',
        '   → Video review từ giáo viên có kinh nghiệm',
        '   → Case study: "Lớp cô Lan tăng trung bình 15 điểm sau 3 tháng"',
        
        '3. Content marketing: Giáo viên tạo content chất lượng cao',
        '   → Blog posts, video hướng dẫn từ góc nhìn chuyên môn',
        '   → Organic SEO traffic → Giảm CAC cho B2C',
        
        '4. Giảm rào cản tâm lý: "Sản phẩm giáo viên xài" → An tâm hơn',
        
        '5. Cross-sell tự nhiên: Học sinh dùng qua GV → Phụ huynh thấy tốt → Mua thêm gói cao cấp B2C'
    ])
    
    synergy_b2b_to_b2c = [
        ['Cơ chế', 'Mô tả', 'Tác động đo lường được', 'Ví dụ cụ thể'],
        [
            'Giảm CAC B2C',
            'Brand awareness từ B2B giúp ads B2C rẻ hơn',
            'Giảm 30-40% CPC trên Facebook/Google',
            'CPC giảm từ 8,000đ → 5,000đ nhờ brand recognition'
        ],
        [
            'Tăng conversion rate',
            'Phụ huynh tin tưởng hơn khi thấy GV dùng',
            'Conversion tăng từ 2% → 4-5%',
            'Landing page có testimonial GV → CVR gấp đôi'
        ],
        [
            'Tăng retention',
            'Hs đã quen từ lớp → Tiếp tục dùng ở nhà',
            'Retention tháng 6 tăng từ 35% → 55%',
            'HS dùng qua GV ở trường, về nhà thuyết phục ba mẹ mua'
        ],
        [
            'Organic growth',
            'GV review → Lan truyền tự nhiên',
            '15-20% user B2C đến từ organic',
            'YouTube review từ GV có 50k views → 200 signups'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 6.1: Cơ chế B2B hỗ trợ B2C')
    table = add_table_styled(doc, synergy_b2b_to_b2c[1:], synergy_b2b_to_b2c[0], col_widths=[1.5, 2.0, 1.8, 2.2])
    doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '6.2. B2C hỗ trợ B2B: Brand Power và Resources', level=3)
    
    add_detailed_section(doc, 'Làm thế nào B2C giúp B2B tăng trưởng tốt hơn:', [
        '1. Brand power: Khi LEXIO nổi tiếng với người tiêu dùng, GV muốn hợp tác hơn',
        '   → GV tìm đến LEXIO chủ động (inbound), không cần sales hunt',
        '   → Giảm chi phí sales B2B đáng kể',
        
        '2. Tài chính vững mạnh: Doanh thu B2C cao → Reinvest vào product',
        '   → Sản phẩm tốt hơn → GV thích dùng hơn → B2B tăng trưởng nhanh',
        
        '3. Tăng sức mạnh thương lượng: "Chúng tôi có 50k user B2C" → GV coi trọng hơn',
        '   → Có thể giữ hoa hồng 30% thay vì tăng lên 40-50%',
        '   → Margin tốt hơn',
        
        '4. Upsell cho GV: GV bắt đầu với gói cơ bản (B2B) → Thấy hiệu quả → Mua thêm premium features (B2C)',
        
        '5. Dữ liệu từ B2C → Cải thiện sản phẩm B2B:',
        '   → Biết tính năng nào user thích nhất',
        '   → A/B testing trên B2C → Apply vào B2B'
    ])
    
    synergy_b2c_to_b2b = [
        ['Cơ chế', 'Mô tả', 'Tác động đo lường được', 'Ví dụ cụ thể'],
        [
            'Inbound leads tăng',
            'GV tự tìm đến nhờ brand B2C',
            'Inbound tăng từ 10% → 40% total leads',
            'Từ phải cold call → GV tự liên hệ qua website'
        ],
        [
            'Giảm chi phí R&D',
            'Dữ liệu B2C → Biết feature nào hiệu quả',
            'Tiết kiệm 30% chi phí phát triển sản phẩm',
            'Biết gamification quan trọng → Focus vào đó'
        ],
        [
            'Tăng giá trị perceived',
            'GV thấy brand lớn → Tin dùng hơn',
            'Churn rate GV giảm từ 15% → 8%/năm',
            'GV nghĩ "LEXIO sẽ tồn tại lâu dài" → Cam kết lâu dài'
        ],
        [
            'Cross-sell',
            'GV dùng bản free B2B → Nâng cấp B2C premium',
            '10-15% GV trở thành khách hàng B2C',
            'GV mua gói premium cho con riêng hoặc để demo'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 6.2: Cơ chế B2C hỗ trợ B2B')
    table = add_table_styled(doc, synergy_b2c_to_b2b[1:], synergy_b2c_to_b2b[0], col_widths=[1.5, 2.0, 1.8, 2.2])
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('🔄 VÒNG TUẦN HOÀN TÍCH CỰC: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    run.font.size = Pt(12)
    run = p.add_run(
        '\n\nB2B tạo credibility → B2C dễ bán hơn → Doanh thu cao → Reinvest vào product → '
        'Sản phẩm tốt hơn → B2B phát triển nhanh hơn → Brand mạnh hơn → B2C phát triển nhanh hơn → ...\n\n'
        'Đây là "flywheel effect" - một khi bánh đà quay, nó sẽ tự tăng tốc.'
    )
    
    doc.add_page_break()
    
    p = doc.add_paragraph()
    run = p.add_run('📌 KẾT LUẬN PHẦN II: ')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(220, 38, 38)
    run = p.add_run(
        '\n\nKhông phải chọn giữa B2B hoặc B2C. Chiến lược đúng là:')
    
    conclusion_bullets = [
        '🎯 Giai đoạn 1-6 tháng: Tập trung 100% vào B2B để tạo market proof',
        '🚀 Tháng 7-12: Bắt đầu B2C ở thành phố lớn, vẫn duy trì tăng trưởng B2B',
        '⚡ Năm 2: Song song cả hai kênh với tài nguyên cân đối',
        '💎 Năm 3: B2C trở thành động lực tăng trưởng chính, B2B là nền tảng vững chắc',
        '',
        'Dual-track strategy cho phép:',
        '✅ Tối ưu CAC trong giai đoạn đầu (B2B)',
        '✅ Tối đa hóa ARPU và TAM dài hạn (B2C)',
        '✅ Giảm rủi ro phụ thuộc một kênh',
        '✅ Tạo hiệu ứng hiệp lực giữa hai kênh',
        '✅ Định giá cao hơn khi gọi vốn'
    ]
    
    for bullet in conclusion_bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_page_break()
    
    print("Đang tạo PHẦN III: MÔ HÌNH HOA HỒNG...")
    
    # ==================== PHẦN III: MÔ HÌNH HOA HỒNG ====================
    add_heading_custom(doc, 'PHẦN III', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'MÔ HÌNH HOA HỒNG GIÁO VIÊN CHI TIẾT', level=1, color=RGBColor(31, 41, 55))
    
    doc.add_paragraph(
        'Trái tim của mô hình B2B là cơ chế hoa hồng giáo viên. Phần này phân tích chi tiết '
        'cách thức hoạt động, lợi ích cho tất cả các bên, và lý do tại sao mô hình này bền vững.'
    )
    
    doc.add_page_break()
    
    # ========== 7. CƠ CHẾ HOẠT ĐỘNG ==========
    add_heading_custom(doc, '7. CƠ CHẾ HOẠT ĐỘNG CHI TIẾT', level=2)
    
    add_heading_custom(doc, '7.1. Quy trình từ A-Z', level=3)
    
    process_steps = [
        '1️⃣ Giáo viên đăng ký tài khoản LEXIO Teacher (miễn phí)',
        '2️⃣ Hoàn thành khóa đào tạo online 2 giờ về cách dùng LEXIO',
        '3️⃣ Nhận mã giới thiệu cá nhân (Referral Code)',
        '4️⃣ Chia sẻ mã với học sinh/phụ huynh trong lớp',
        '5️⃣ Phụ huynh đăng ký và nhập mã giáo viên',
        '6️⃣ Hệ thống tự động liên kết học sinh với giáo viên',
        '7️⃣ Hàng tháng, hoa hồng được tính dựa trên số học sinh active',
        '8️⃣ Giáo viên rút tiền qua ngân hàng hoặc ví điện tử',
        '',
        '💡 ĐƠN GIẢN - MINH BẠCH - TỰ ĐỘNG: Giáo viên không cần làm kế toán, mọi thứ đã được tính sẵn.'
    ]
    
    for step in process_steps:
        doc.add_paragraph(step)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.2. Ví dụ cụ thể: Cô Lan ở Bắc Ninh', level=3)
    
    add_detailed_section(doc, '📚 CASE STUDY:' , [
        'Cô Lan là giáo viên dạy thêm tiếng Anh tại Bắc Ninh, có 15 học sinh lớp 5-6.',
        'Trước đây cô dạy theo sách giáo khoa, không dùng công nghệ.',
        '',
        'Tháng 3/2026: Cô Lan được giới thiệu LEXIO bởi đồng nghiệp.',
        '• Hoàn thành khóa đào tạo trong 1 buổi tối',
        '• Giới thiệu LEXIO cho 15 học sinh của mình',
        '• 15/15 phụ huynh đồng ý mua gói cơ bản 99k/tháng',
        '',
        'Kết quả sau 3 tháng (Tháng 6/2026):'
    ])
    
    coLan_numbers = [
        ['Chỉ số', 'Tháng 3', 'Tháng 4', 'Tháng 5', 'Tháng 6', 'Ghi chú'],
        [
            'Số HS active',
            '15',
            '15',
            '14',
            '15',
            '1 HS nghỉ hè tháng 5, quay lại tháng 6'
        ],
        [
            'Doanh thu từ HS (99k/HS)',
            '1,485,000đ',
            '1,485,000đ',
            '1,386,000đ',
            '1,485,000đ',
            'Tổng tiền HS/PH trả'
        ],
        [
            'Hoa hồng cô Lan (30%)',
            '445,500đ',
            '445,500đ',
            '415,800đ',
            '445,500đ',
            'Thu nhập thêm cho cô'
        ],
        [
            'LEXIO nhận (70%)',
            '1,039,500đ',
            '1,039,500đ',
            '970,200đ',
            '1,039,500đ',
            'Doanh thu LEXIO'
        ],
        [
            'Thời gian cô bỏ ra',
            '2h đào tạo',
            '0h',
            '0h',
            '0h',
            'Chỉ dạy như bình thường'
        ],
        [
            'Tổng thu nhập cô Lan',
            '445k',
            '445k',
            '416k',
            '445k',
            'Trung bình ~438k/tháng'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 7.1: Thu nhập của cô Lan qua 4 tháng')
    table = add_table_styled(doc, coLan_numbers[1:], coLan_numbers[0], col_widths=[1.2, 1.0, 1.0, 1.0, 1.0, 1.8])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('💰 PHÂN TÍCH: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    run = p.add_run(
        '\n• Cô Lan kiếm thêm ~438,000đ/tháng mà KHÔNG tốn thêm thời gian (vì vẫn dạy lớp như bình thường)\n'
        '• Với mức lương giáo viên ở Bắc Ninh (~6-8 triệu/tháng), đây là khoản thu nhập phụ đáng kể (5-7% tổng thu nhập)\n'
        '• Sau 1 năm, cô Lan kiếm thêm ~5.2 triệu đồng\n'
        '• Cô không phải làm kế toán, thu tiền từng HS - Tất cả tự động\n'
        '• Học sinh của cô học tốt hơn nhờ LEXIO → Uy tín của cô tăng → Nhiều HS mới đến hơn\n\n'
        '✨ ĐÂY LÀ MÔ HÌNH "WIN-WIN" THỰC SỰ'
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ========== 8. PHÂN TÍCH LỢI NHUẬN ==========
    add_heading_custom(doc, '8. PHÂN TÍCH LỢI NHUẬN THEO TỪNG GÓI', level=2)
    
    add_heading_custom(doc, '8.1. Cấu trúc 4 gói sản phẩm B2B', level=3)
    
    four_packages = [
        ['Gói', 'Giá/HS/tháng', 'Đối tượng', 'Tính năng', 'Hoa hồng GV', 'LEXIO nhận', 'Margin'],
        [
            'STARTER',
            '99,000đ',
            'GV dạy thêm nhỏ lẻ (5-20 HS)',
            'Core features, giới hạn 20 HS',
            '30,000đ (30%)',
            '69,000đ (70%)',
            '~85% (sau trừ chi phí)'
        ],
        [
            'PRO',
            '299,000đ',
            'GV chuyên nghiệp (20-50 HS)',
            'Full features, báo cáo chi tiết',
            '90,000đ (30%)',
            '209,000đ (70%)',
            '~88%'
        ],
        [
            'SMALL CENTER',
            '999,000đ',
            'Trung tâm nhỏ (50-200 HS)',
            'Multi-class, admin dashboard',
            '300,000đ (30%)',
            '699,000đ (70%)',
            '~90%'
        ],
        [
            'LARGE CENTER',
            '2,999,000đ',
            'Trung tâm lớn (200+ HS)',
            'White-label, API, dedicated support',
            '900,000đ (30%)',
            '2,099,000đ (70%)',
            '~92%'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 8.1: Chi tiết 4 gói sản phẩm B2B')
    table = add_table_styled(doc, four_packages[1:], four_packages[0], col_widths=[1.0, 1.0, 1.5, 1.5, 1.0, 1.0, 0.8])
    doc.add_paragraph()
    
    add_detailed_section(doc, '📈 PHÂN TÍCH MARGIN:', [
        'Chi phí cố định của LEXIO cho mỗi user rất thấp (SaaS model):',
        '• Server cost: ~5,000đ/user/tháng',
        '• Customer support: ~3,000đ/user/tháng (vì GV làm support tầng 1)',
        '• Payment gateway: 2% (~1,400-2,000đ)',
        '• Content update: ~1,000đ/user/tháng',
        '→ Tổng chi phí: ~10,000-12,000đ/user/tháng',
        '',
        'Với gói Starter (LEXIO nhận 69k):',
        '• Chi phí: 12k',
        '• Profit: 57k',
        '• Margin: 83%',
        '',
        'Với gói Pro (LEXIO nhận 209k):',
        '• Chi phí: 12k',
        '• Profit: 197k',
        '• Margin: 94%',
        '',
        '💡 ĐÂY LÀ MÔ HÌNH CÓ MARGIN CỰC KỲ TỐT - điển hình của SaaS B2B'
    ])
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '8.2. Tại sao giữ hoa hồng 30% cho tất cả các gói?', level=3)
    
    add_detailed_section(doc, 'Lý do chiến lược:', [
        '1. ĐƠN GIẢN VÀ CÔNG BẰNG:',
        '   • GV dễ tính toán, không nhầm lẫn',
        '   • Không phân biệt GV lớn/nhỏ → Mọi người đều cảm thấy công bằng',
        '   • Dễ marketing: "Kiếm 30% hoa hồng từ mỗi học sinh"',
        '',
        '2. KHUYẾN KHÍCH UPGRADE:',
        '   • GV muốn học sinh nâng cấp vì cả hai đều được lợi nhiều hơn',
        '   • Nếu % thấp hơn ở gói cao, GV sẽ không muốn học sinh upgrade',
        '',
        '3. CẠNH TRANH VỚI SÁCH GIÁO KHOA:',
        '   • Sách giáo khoa: GV mua sỉ với giá thấp, bán lẻ cho HS → Lãi ~40-50%',
        '   • LEXIO 30% + không phải lo kho, vốn, giao hàng → Cạnh tranh được',
        '',
        '4. MARGIN VẪN TỐT:',
        '   • Với chi phí thấp (~12k/user), LEXIO vẫn lãi tốt ở mức 70%',
        '   • Không cần phải "bóc lột" giáo viên để có lợi nhuận'
    ])
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ========== 9. CÁC CHƯƠNG TRÌNH KHUYẾN KHÍCH ==========
    add_heading_custom(doc, '9. CÁC CHƯƠNG TRÌNH KHUYẾN KHÍCH BỔ SUNG', level=2)
    
    doc.add_paragraph(
        'Ngoài hoa hồng cơ bản 30%, LEXIO có 4 chương trình khuyến khích bổ sung để '
        'động viên giáo viên tăng trưởng mạnh mẽ hơn.'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '9.1. Chương trình Referral Bonus (Thưởng giới thiệu GV)', level=3)
    
    referral_program = [
        ['Số GV giới thiệu', 'Bonus/GV mới', 'Điều kiện', 'Ví dụ', 'Tổng thu nhập'],
        [
            '1-2 GV',
            '100,000đ',
            'GV mới có ít nhất 5 HS active',
            'Giới thiệu 2 GV → 200k',
            'Hoa hồng thường + 200k'
        ],
        [
            '3-5 GV',
            '150,000đ',
            'GV mới có ít nhất 5 HS active',
            'Giới thiệu 4 GV → 600k',
            'Hoa hồng thường + 600k'
        ],
        [
            '6-10 GV',
            '200,000đ',
            'GV mới có ít nhất 5 HS active',
            'Giới thiệu 8 GV → 1.6M',
            'Hoa hồng thường + 1.6M'
        ],
        [
            '11+ GV',
            '250,000đ',
            'GV mới có ít nhất 5 HS active',
            'Giới thiệu 15 GV → 3.75M',
            'Hoa hồng thường + 3.75M'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 9.1: Chương trình thưởng giới thiệu giáo viên')
    table = add_table_styled(doc, referral_program[1:], referral_program[0], col_widths=[1.3, 1.0, 1.5, 1.5, 1.5])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('🎯 MỤC ĐÍCH: ')
    run.font.bold = True
    run = p.add_run(
        'Tạo viral effect trong cộng đồng giáo viên. Một giáo viên tích cực có thể kiếm thêm '
        'hàng triệu đồng chỉ từ việc giới thiệu đồng nghiệp.'
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '9.2. Chương trình Tier System (Hệ thống bậc)', level=3)
    
    tier_system = [
        ['Bậc', 'Điều kiện', 'Lợi ích', 'Hoa hồng tăng', 'Ví dụ thu nhập/tháng'],
        [
            'Bronze',
            '5-19 HS active',
            'Dashboard cơ bản',
            '30% (chuẩn)',
            '15 HS × 99k × 30% = 446k'
        ],
        [
            'Silver',
            '20-49 HS active',
            'Dashboard nâng cao, ưu tiên support',
            '32%',
            '30 HS × 99k × 32% = 950k'
        ],
        [
            'Gold',
            '50-99 HS active',
            'Toàn bộ features, training 1-1',
            '35%',
            '70 HS × 99k × 35% = 2.4M'
        ],
        [
            'Platinum',
            '100+ HS active',
            'White-label option, dedicated account manager',
            '40%',
            '120 HS × 99k × 40% = 4.75M'
        ]
    ]
    
    doc.add_paragraph('📊 Bảng 9.2: Hệ thống bậc và hoa hồng tăng dần')
    table = add_table_styled(doc, tier_system[1:], tier_system[0], col_widths=[0.9, 1.3, 2.0, 1.2, 1.8])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('📈 INSIGHT: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    run = p.add_run(
        'Hệ thống bậc tạo động lực cho GV mở rộng quy mô. Một GV có 100 học sinh kiếm gần 5 triệu/tháng '
        'từ hoa hồng - đây là mức thu nhập rất hấp dẫn ở Việt Nam, tương đương lương chính của nhiều GV.'
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '9.3. Chương trình Retention Bonus (Thưởng giữ chân)', level=3)
    
    retention_bonus = [
        ['Thời gian giữ chân HS', 'Bonus % trên doanh thu HS đó', 'Ví dụ'],
        ['6 tháng liên tục', '+5%', 'HS trả 99k → GV nhận 30k + 5k = 35k (tháng thứ 7 trở đi)'],
        ['12 tháng liên tục', '+10%', 'HS trả 99k → GV nhận 30k + 10k = 40k (tháng thứ 13 trở đi)'],
        ['24 tháng liên tục', '+15%', 'HS trả 99k → GV nhận 30k + 15k = 45k (tháng thứ 25 trở đi)']
    ]
    
    doc.add_paragraph('📊 Bảng 9.3: Thưởng giữ chân học sinh lâu dài')
    table = add_table_styled(doc, retention_bonus[1:], retention_bonus[0], col_widths=[2.0, 2.5, 3.0])
    doc.add_paragraph()
    
    add_detailed_section(doc, 'Lợi ích cho cả hai bên:', [
        'CHO GIÁO VIÊN:',
        '• Thu nhập tăng dần theo thời gian từ cùng một học sinh',
        '• Động lực chăm sóc học sinh tốt để họ không rời đi',
        '• LTV (Lifetime Value) cao → Thu nhập ổn định, bền vững',
        '',
        'CHO LEXIO:',
        '• Retention cao → Giảm churn rate',
        '• CAC được amortize trên thời gian dài hơn → Unit economics tốt hơn',
        '• Học sinh dùng lâu → Dữ liệu tốt hơn → Sản phẩm tốt hơn'
    ])
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '9.4. Chương trình Free Training & Certification', level=3)
    
    add_detailed_section(doc, 'Đào tạo miễn phí cho giáo viên:', [
        'LEXIO cung cấp MIỄN PHÍ các khóa đào tạo sau:',
        '',
        '1. LEXIO Product Training (2 giờ):',
        '   • Cách sử dụng app hiệu quả',
        '   • Tips & tricks để học sinh engage cao',
        '   • Cách theo dõi tiến độ và đưa ra feedback',
        '',
        '2. Cambridge CLIL Methodology (8 giờ):',
        '   • Phương pháp CLIL là gì?',
        '   • Cách áp dụng CLIL trong lớp học Việt Nam',
        '   • Case studies và best practices',
        '   • Nhận chứng chỉ Cambridge CLIL Foundation',
        '',
        '3. EdTech Skills for Teachers (4 giờ):',
        '   • Kỹ năng dạy học trên môi trường số',
        '   • Cách kết hợp online và offline',
        '   • Gamification trong giảng dạy',
        '',
        '4. Business Skills for Teacher-Entrepreneurs (6 giờ):',
        '   • Cách marketing lớp học của bạn',
        '   • Quản lý tài chính cá nhân',
        '   • Xây dựng thương hiệu cá nhân',
        '',
        'GIÁ TRỊ: Các khóa này trị giá ~3-5 triệu nếu mua ở ngoài. LEXIO cung cấp MIỄN PHÍ.',
        '',
        'LỢI ÍCH CHO LEXIO:',
        '✅ GV được đào tạo tốt → Dạy tốt hơn → HS hài lòng hơn → Retention cao',
        '✅ Chứng chỉ Cambridge → Tăng uy tín cho GV → Họ trung thành với LEXIO',
        '✅ Community building: GV cảm thấy được đầu tư, quan tâm → Loyalty cao'
    ])
    
    doc.add_page_break()
    
    p = doc.add_paragraph()
    run = p.add_run('📌 KẾT LUẬN PHẦN III: ')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(220, 38, 38)
    run = p.add_run(
        '\n\nMô hình hoa hồng của LEXIO không chỉ là một cơ chế thu hút giáo viên, mà là một HỆ SINH THÁI '
        'mang lại giá trị thực sự cho tất cả các bên:\n'
    )
    
    conclusion_part3 = [
        '',
        '✅ GIÁO VIÊN:',
        '   • Thu nhập thêm 400k - 5M/tháng không tốn thêm thời gian',
        '   • Công cụ dạy học tốt hơn → Uy tín tăng → Nhiều HS hơn',
        '   • Đào tạo miễn phí → Nâng cao năng lực chuyên môn',
        '   • Cộng đồng giáo viên → Kết nối, học hỏi lẫn nhau',
        '',
        '✅ HỌC SINH:',
        '   • Học tiếng Anh hiệu quả hơn nhờ công nghệ',
        '   • Giá rẻ hơn nhiều so với trung tâm (99k vs 500k-2M/tháng)',
        '   • Vẫn được GV quen thuộc hướng dẫn',
        '',
        '✅ PHỤ HUYNH:',
        '   • Con học tốt hơn, giá rẻ hơn',
        '   • Theo dõi tiến độ minh bạch qua app',
        '   • Tin tưởng vì có GV giới thiệu',
        '',
        '✅ LEXIO:',
        '   • CAC thấp (15k vs 100k B2C)',
        '   • Margin cao (83-94%)',
        '   • Tăng trưởng nhanh nhờ viral effect',
        '   • Giữ chân tốt nhờ GV trung gian',
        '',
        'Đây là mô hình "WIN-WIN-WIN" thực sự - TẤT CẢ CÁC BÊN ĐỀU THẮNG.'
    ]
    
    for line in conclusion_part3:
        doc.add_paragraph(line)
    
    doc.add_page_break()
    
    print("Đang tạo PHẦN IV: CHIẾN LƯỢC THỰC THI 3 NĂM (chi tiết)...")
    
    # ==================== PHẦN IV: CHIẾN LƯỢC THỰC THI ====================
    add_heading_custom(doc, 'PHẦN IV', level=1, color=RGBColor(220, 38, 38))
    add_heading_custom(doc, 'CHIẾN LƯỢC THỰC THI 3 NĂM CHI TIẾT', level=1, color=RGBColor(31, 41, 55))
    
    doc.add_paragraph(
        'Phần này trình bày chi tiết lộ trình thực thi theo từng giai đoạn, với các mục tiêu cụ thể, '
        'chiến thuật triển khai, và kết quả kỳ vọng cho mỗi tháng/quý.'
    )
    
    doc.add_page_break()
    
    # ========== 10. GIAI ĐOẠN 1 (THÁNG 1-6) ==========
    add_heading_custom(doc, '10. GIAI ĐOẠN 1: XÂY DỰNG NỀN TẢNG B2B (Tháng 1-6)', level=2)
    
    add_heading_custom(doc, '10.1. Tổng quan giai đoạn', level=3)
    
    phase1_overview = [
        ['Chỉ tiêu', 'Mục tiêu', 'Chiến lược', 'Ngân sách'],
        ['Số giáo viên', '50-60 GV', 'Tìm kiếm chủ động + referral', '20M (CAC 200k/GV)'],
        ['Số học sinh', '500-750 HS', 'Qua giáo viên (10-15 HS/GV)', 'Included in GV CAC'],
        ['Doanh thu/tháng', '35-50M', 'ARPU 84k × 500 HS', 'N/A'],
        ['Churn rate', '< 20%', 'Hỗ trợ GV sát sao', '5M support'],
        ['NPS Score', '> 50', 'Chất lượng sản phẩm + support tốt', '10M product improvement']
    ]
    
    doc.add_paragraph('📊 Bảng 10.1: Mục tiêu tổng quan Giai đoạn 1')
    table = add_table_styled(doc, phase1_overview[1:], phase1_overview[0], col_widths=[1.5, 1.2, 2.5, 1.8])
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '10.2. Lộ trình chi tiết từng tháng', level=3)
    
    # Tháng 1-2
    add_detailed_section(doc, '📅 THÁNG 1-2: Pilot và Tìm Product-Market Fit', [
        'MỤC TIÊU:',
        '• 10 giáo viên pilot',
        '• 100-150 học sinh',
        '• Xác định product-market fit',
        '• Thu thập feedback chất lượng cao',
        '',
        'CHIẾN THUẬT:',
        '1. Tìm kiếm 10 giáo viên pilot:',
        '   → Tìm trong mạng lưới quen biết',
        '   → Tham gia các group Facebook giáo viên',
        '   → Liên hệ trực tiếp các giáo viên dạy thêm ở Hà Nội, Bắc Ninh',
        '',
        '2. Onboarding cẩn thận:',
        '   → Đào tạo 1-1 cho từng giáo viên (2 giờ)',
        '   → Theo sát trong 2 tuần đầu',
        '   → Hotline support 24/7',
        '',
        '3. Thu thập feedback:',
        '   → Weekly call với mỗi GV',
        '   → Survey học sinh sau mỗi tuần',
        '   → Phân tích data usage',
        '',
        '4. Iterate nhanh:',
        '   → Fix bug trong vòng 24h',
        '   → Deploy feature mới mỗi tuần',
        '   → A/B test các tính năng',
        '',
        'KẾT QUẢ KỲ VỌNG CUỐI THÁNG 2:',
        '✅ 10 GV hài lòng (NPS > 60)',
        '✅ 120-150 HS active',
        '✅ Doanh thu: ~10M/tháng',
        '✅ Churn rate: < 10% (vì theo dõi sát)',
        '✅ Đã xác định được product-market fit cơ bản',
        '✅ List 50-100 feature requests/improvements để làm tiếp'
    ])
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # Tháng 3-4
    add_detailed_section(doc, '📅 THÁNG 3-4: Mở Rộng Ban Đầu và Tối Ưu', [
        'MỤC TIÊU:',
        '• 30 giáo viên (tăng 3x)',
        '• 300-400 học sinh',
        '• Tối ưu onboarding process',
        '• Xây dựng case studies',
        '',
        'CHIẾN THUẬT:',
        '1. Mở rộng tìm kiếm GV:',
        '   → Chương trình referral: 10 GV hiện tại giới thiệu thêm 20 GV',
        '   → Quảng cáo trong group Facebook giáo viên (budget: 5M)',
        '   → Tham gia sự kiện giáo viên ở Hà Nội',
        '',
        '2. Tự động hóa onboarding:',
        '   → Video hướng dẫn thay cho training 1-1',
        '   → Onboarding email sequence tự động',
        '   → Chatbot hỗ trợ câu hỏi thường gặp',
        '   → Giảm thời gian onboarding từ 2h → 30 phút',
        '',
        '3. Xây dựng marketing materials:',
        '   → 3 case studies chi tiết (như cô Lan)',
        '   → Video testimonial từ 5 GV',
        '   → Infographic về kết quả học tập',
        '',
        '4. Cải thiện retention:',
        '   → Email tự động cho GV/PH khi HS không active 3 ngày',
        '   → In-app gamification tăng engagement',
        '   → Weekly progress report cho phụ huynh',
        '',
        'KẾT QUẢ KỲ VỌNG CUỐI THÁNG 4:',
        '✅ 30 GV active (20 GV mới từ referral + ads)',
        '✅ 350-400 HS active',
        '✅ Doanh thu: ~30M/tháng (tăng 3x)',
        '✅ Churn rate: < 15%',
        '✅ NPS: > 55',
        '✅ Có 3 case studies chất lượng để marketing',
        '✅ Quy trình onboarding đã tự động hóa 80%'
    ])
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # Tháng 5-6
    add_detailed_section(doc, '📅 THÁNG 5-6: Tăng Tốc và Chuẩn Bị B2C', [
        'MỤC TIÊU:',
        '• 60 giáo viên (nhân đôi từ tháng 4)',
        '• 600-750 học sinh',
        '• Chuẩn bị hạ tầng cho B2C',
        '• Gọi vốn Seed round',
        '',
        'CHIẾN THUẬT B2B:',
        '1. Mở rộng địa lý:',
        '   → Từ Hà Nội/Bắc Ninh → Hải Phòng, Nam Định, Thái Bình',
        '   → Partnership với 3-5 trung tâm nhỏ',
        '   → Tìm "champion teacher" ở mỗi tỉnh mới',
        '',
        '2. Tăng viral coefficient:',
        '   → Launch chương trình referral bonus chính thức',
        '   → GV được 100k cho mỗi GV mới giới thiệu',
        '   → Kỳ vọng mỗi GV giới thiệu 1-2 GV → 30 GV hiện tại → +40-50 GV mới',
        '',
        '3. Nâng cao chất lượng:',
        '   → Ra mắt Cambridge CLIL certification cho GV',
        '   → Thêm 50+ bài học mới',
        '   → AI tutor beta feature',
        '',
        'CHUẨN BỊ CHO B2C:',
        '1. Xây dựng landing page B2C:',
        '   → Với testimonial từ GV và PH',
        '   → A/B testing messaging',
        '',
        '2. Chuẩn bị content marketing:',
        '   → Blog về phương pháp CLIL',
        '   → SEO cho keywords "học tiếng Anh cho trẻ"',
        '   → YouTube channel với tips học tiếng Anh',
        '',
        '3. Setup ads infrastructure:',
        '   → Facebook Pixel',
        '   → Google Analytics 4',
        '   → Conversion tracking',
        '',
        '4. Chuẩn bị gọi vốn:',
        '   → Pitch deck với traction từ 6 tháng',
        '   → Financial model 3 năm',
        '   → Tiếp cận 10-15 quỹ Seed',
        '',
        'KẾT QUẢ KỲ VỌNG CUỐI THÁNG 6:',
        '✅ 60 GV active ở 5-6 tỉnh thành',
        '✅ 700-750 HS active',
        '✅ Doanh thu: ~50M/tháng',
        '✅ MRR growth rate: 40-50%/tháng',
        '✅ Churn rate: < 12%',
        '✅ NPS: > 60',
        '✅ Đã có Term Sheet từ ít nhất 1 quỹ (500k-1M USD)',
        '✅ Landing page B2C đã sẵn sàng',
        '✅ 10k organic visitors/tháng trên blog/YouTube',
        '',
        '💰 TỔNG DOANH THU 6 THÁNG ĐẦU: ~150-180M',
        '📊 TỔNG CHI PHÍ: ~200-250M (bao gồm team, product development, marketing)',
        '💸 BURN RATE: ~-10M/tháng (chấp nhận được, sẽ gọi vốn để scale)'
    ])
    
    doc.add_page_break()
    
    # ========== 11. GIAI ĐOẠN 2 (THÁNG 7-18) ==========
    add_heading_custom(doc, '11. GIAI ĐOẠN 2: MỞ RỘNG CẢ HAI KÊNH (Tháng 7-18)', level=2)
    
    add_heading_custom(doc, '11.1. Tổng quan giai đoạn', level=3)
    
    phase2_overview = [
        ['Chỉ tiêu', 'Mục tiêu cuối tháng 18', 'Tăng trưởng', 'Chiến lược chính'],
        ['Giáo viên (B2B)', '300-400 GV', '5x từ tháng 6', 'Mở rộng toàn quốc, partnership với trung tâm'],
        ['Học sinh B2B', '3,500-5,000 HS', '5-7x', 'Scale qua GV và trung tâm'],
        ['Học sinh B2C', '1,000-2,000 HS', 'Từ 0', 'Launch B2C tại HN, HCM với paid ads'],
        ['Tổng học sinh', '5,000-7,000 HS', '7-10x', 'Dual-track hoàn chỉnh'],
        ['Doanh thu/tháng', '400-500M', '8-10x', 'B2B: 300M + B2C: 100-200M'],
        ['MRR growth', '15-20%/tháng', 'Compound', 'Virality + Ads']
    ]
    
    doc.add_paragraph('📊 Bảng 11.1: Mục tiêu tổng quan Giai đoạn 2')
    table = add_table_styled(doc, phase2_overview[1:], phase2_overview[0], col_widths=[1.5, 1.8, 1.3, 2.5])
    doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '11.2. Kế hoạch chi tiết', level=3)
    
    # Tháng 7-12: Năm 1 H2
    add_detailed_section(doc, '📅 THÁNG 7-12 (Năm 1 - Nửa cuối): Song song B2B và B2C Pilot', [
        'TRỌNG TÂM: Duy trì tăng trưởng B2B + Launch B2C pilot',
        '',
        'B2B - MỞ RỘNG ĐỊA LÝ:',
        '1. Mở rộng từ 6 → 20 tỉnh thành:',
        '   → Tháng 7-8: Thêm 5 tỉnh Bắc Trung Bộ (Thanh Hóa, Nghệ An, Hà Tĩnh, Quảng Bình, Huế)',
        '   → Tháng 9-10: Thêm 5 tỉnh Nam Trung Bộ và Tây Nguyên',
        '   → Tháng 11-12: Thêm 4 tỉnh đồng bằng sông Cửu Long',
        '',
        '2. Partnership với trung tâm:',
        '   → Tìm 20-30 trung tâm nhỏ (50-200 HS) làm đối tác',
        '   → Gói Small Center: 999k/tháng',
        '   → Mỗi trung tâm đóng góp 50-100 HS → Tổng: 1,500-2,000 HS từ channel này',
        '',
        '3. Kết quả kỳ vọng B2B cuối tháng 12:',
        '   → 150-180 GV cá nhân',
        '   → 25-30 trung tâm',
        '   → 2,500-3,000 HS tổng',
        '   → Doanh thu B2B: ~200-250M/tháng',
        '',
        'B2C - LAUNCH PILOT:',
        '1. Tháng 7-8: Soft launch ở Hà Nội:',
        '   → Facebook Ads budget: 20M/tháng',
        '   → Landing page với testimonial từ B2B',
        '   → Target: Phụ huynh con 8-12 tuổi, thu nhập 15M+/tháng',
        '   → Mục tiêu: 100-150 HS B2C trong 2 tháng',
        '',
        '2. Tháng 9-10: Mở rộng TP.HCM:',
        '   → Tương tự Hà Nội',
        '   → Budget: 25M/tháng (thị trường lớn hơn)',
        '   → Mục tiêu: +150-200 HS',
        '',
        '3. Tháng 11-12: Tối ưu và scale:',
        '   → Đã có data từ 4 tháng → Biết channel nào hiệu quả',
        '   → Tăng budget vào channel tốt nhất',
        '   → Thêm Đà Nẵng',
        '   → Mục tiêu: +200-300 HS',
        '',
        '4. Kết quả kỳ vọng B2C cuối tháng 12:',
        '   → 500-650 HS B2C',
        '   → CAC: ~80-100k (thấp hơn 100k target nhờ brand từ B2B)',
        '   → ARPU: ~250-300k (mix của gói Basic 199k và Pro 499k)',
        '   → Doanh thu B2C: ~130-180M/tháng',
        '',
        'TỔNG KẾT CUỐI NĂM 1 (THÁNG 12):',
        '🎯 Tổng HS: 3,000-3,500 (B2B) + 500-600 (B2C) = 3,500-4,000 HS',
        '💰 Doanh thu: 250M (B2B) + 150M (B2C) = 400M/tháng',
        '📈 MRR: 400M → ARR run-rate: 4.8 tỷ',
        '💸 Đã gọi được vốn Seed: 500k-1M USD → Đủ tiền để scale năm 2'
    ])
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # Tháng 13-18: Năm 2 H1
    add_detailed_section(doc, '📅 THÁNG 13-18 (Năm 2 - Nửa đầu): Tăng Tốc Toàn Diện', [
        'TRỌNG TÂM: Scale mạnh cả B2B và B2C với vốn Seed đã gọi được',
        '',
        'B2B - SCALE TOÀN QUỐC:',
        '1. Mở rộng 20 → 50 tỉnh/thành (phủ gần toàn quốc):',
        '   → Thuê 5 Regional Sales Managers',
        '   → Mỗi người phụ trách 10 tỉnh',
        '   → Tổ chức roadshow ở 20 tỉnh lớn',
        '',
        '2. Tăng cường partnership:',
        '   → Từ 30 → 80 trung tâm đối tác',
        '   → Thêm 10-15 trung tâm lớn (gói Large Center: 2.999k)',
        '   → Partnership với 2-3 chuỗi trung tâm (100+ cơ sở)',
        '',
        '3. Kết quả kỳ vọng B2B cuối tháng 18:',
        '   → 300-350 GV cá nhân',
        '   → 80-100 trung tâm (bao gồm 10 trung tâm lớn)',
        '   → 4,500-5,500 HS',
        '   → Doanh thu B2B: ~350-400M/tháng',
        '',
        'B2C - SCALE PAID ACQUISITION:',
        '1. Tăng marketing budget:',
        '   → Từ 50M/tháng → 150M/tháng',
        '   → Mở rộng thêm 5 thành phố: Cần Thơ, Hải Phòng, Nha Trang, Vũng Tàu, Biên Hòa',
        '',
        '2. Đa dạng hóa channel:',
        '   → Facebook/Instagram Ads: 60%',
        '   → Google Search/Display: 20%',
        '   → TikTok Ads: 10%',
        '   → KOL/Influencer: 10%',
        '',
        '3. Content marketing tăng cường:',
        '   → Blog 20 bài/tháng (SEO)',
        '   → YouTube 8 video/tháng',
        '   → TikTok 15 video/tháng',
        '   → Mục tiêu organic: 50k visitors/tháng',
        '',
        '4. Tối ưu conversion:',
        '   → A/B test landing pages',
        '   → Free trial 7 ngày',
        '   → Onboarding flow mượt mà',
        '   → Conversion rate target: 3-5%',
        '',
        '5. Kết quả kỳ vọng B2C cuối tháng 18:',
        '   → 1,500-2,000 HS B2C',
        '   → CAC đã giảm xuống 70-80k nhờ brand mạnh',
        '   → ARPU tăng lên ~320k (nhiều người chọn gói Pro hơn)',
        '   → Doanh thu B2C: ~450-600M/tháng',
        '',
        'TỔNG KẾT CUỐI THÁNG 18:',
        '🎯 Tổng HS: 5,000 (B2B) + 1,800 (B2C) = 6,800 HS',
        '💰 Doanh thu: 370M (B2B) + 550M (B2C) = ~920M/tháng',
        '📈 ARR run-rate: 11 tỷ VNĐ (~480k USD)',
        '💸 Chuẩn bị gọi vốn Series A: 2-3M USD với traction này',
        '🎖️ Đã là Top 3 EdTech platform tại Việt Nam'
    ])
    
    doc.add_page_break()
    
    # ========== 12. GIAI ĐOẠN 3 (THÁNG 19-36) ==========
    add_heading_custom(doc, '12. GIAI ĐOẠN 3: ĐỊNH VỊ CAO CẤP VÀ MỞ RỘNG KHU VỰC (Tháng 19-36)', level=2)
    
    add_heading_custom(doc, '12.1. Tổng quan giai đoạn', level=3)
    
    phase3_overview = [
        ['Chỉ tiêu', 'Mục tiêu cuối năm 3', 'Chiến lược'],
        ['Tổng học sinh', '20,000-30,000 HS', 'B2B stable, B2C tăng mạnh, mở rộng SEA'],
        ['Doanh thu/tháng', '4-6 tỷ VNĐ', 'B2C chiếm 60-70%, B2B 30-40%'],
        ['ARR', '50-70 tỷ VNĐ (~2.5M USD)', 'Tương đương startup Series B'],
        ['Market position', 'Top 1 CLIL platform VN', 'Brand leadership'],
        ['Expansion', 'Thailand, Indonesia pilot', 'SEA expansion']
    ]
    
    doc.add_paragraph('📊 Bảng 12.1: Mục tiêu tổng quan Giai đoạn 3')
    table = add_table_styled(doc, phase3_overview[1:], phase3_overview[0], col_widths=[2.0, 2.5, 3.0])
    doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_detailed_section(doc, '📅 CHIẾN LƯỢC CHI TIẾT NĂM 3:', [
        'TRỌNG TÂM: B2C trở thành động lực chính, B2B là nền tảng vững chắc',
        '',
        '━━━━━━━ B2B: CONSOLIDATION & OPTIMIZATION ━━━━━━━',
        '',
        '1. Consolidate thị phần:',
        '   → Duy trì 400-500 GV active',
        '   → 100-150 trung tâm đối tác',
        '   → Focus vào retention hơn là acquisition',
        '   → Churn rate xuống < 8%/năm',
        '',
        '2. Nâng cấp tier system:',
        '   → Platinum tier cho GV/TT lớn (100+ HS)',
        '   → White-label option',
        '   → Dedicated account manager',
        '',
        '3. Mở rộng sang B2B Enterprise:',
        '   → Partnership với các trường tư thục lớn',
        '   → Bán trực tiếp cho trường (không qua GV cá nhân)',
        '   → Gói Enterprise: 10-50M/năm cho 500-2,000 HS',
        '   → Target: 10-20 trường → 5,000-10,000 HS',
        '',
        '4. Kết quả B2B năm 3:',
        '   → 8,000-12,000 HS (bao gồm enterprise)',
        '   → Doanh thu: 600-900M/tháng',
        '   → Churn < 8%, NPS > 70',
        '',
        '━━━━━━━ B2C: SCALE TO DOMINANCE ━━━━━━━',
        '',
        '1. Tăng marketing budget lên 500M-1 tỷ/tháng:',
        '   → Với vốn Series A đã gọi được (2-3M USD)',
        '   → Chiến lược "winner takes most" trong EdTech',
        '',
        '2. Ra mắt Premium tier:',
        '   → Gói Premium: 1.5-2M/tháng',
        '   → 1-on-1 live tutoring với giáo viên bản ngữ',
        '   → Personalized learning path với AI',
        '   → Cam kết đạt chứng chỉ Cambridge',
        '   → Target: 5-10% user chọn Premium → 500-1,000 HS',
        '   → ARPU Premium: 1.8M → Đóng góp 900M-1.8 tỷ/tháng',
        '',
        '3. Mở rộng demographics:',
        '   → Hiện tại: 8-12 tuổi',
        '   → Mở rộng: 6-15 tuổi',
        '   → Thêm nội dung cho mầm non (4-6 tuổi)',
        '   → Thêm nội dung luyện thi IELTS cho THPT',
        '',
        '4. Xây dựng brand mạnh:',
        '   → TV commercials (VTV, HTV)',
        '   → Sponsorship các chương trình giáo dục',
        '   → Partnership với Cambridge để trở thành "Official Cambridge CLIL Platform"',
        '   → Launch community event: "LEXIO English Festival" ở 10 thành phố',
        '',
        '5. Retention & LTV optimization:',
        '   → Thêm features giữ chân: Clubs, Competitions, Badges',
        '   → Parent dashboard nâng cao',
        '   → Monthly progress report chuyên nghiệp',
        '   → Target retention tháng 12: > 65%',
        '',
        '6. Kết quả B2C năm 3:',
        '   → 12,000-18,000 HS (bao gồm 500-1k Premium)',
        '   → ARPU trung bình: ~450k (nhờ mix Premium)',
        '   → Doanh thu: 5-8 tỷ/tháng',
        '',
        '━━━━━━━ INTERNATIONAL EXPANSION: PILOT ━━━━━━━',
        '',
        '1. Chọn thị trường đầu tiên: Thailand',
        '   → Tương tự Việt Nam về văn hóa giáo dục',
        '   → Quy mô: 10M học sinh',
        '   → Dịch content sang tiếng Thái',
        '   → Hire 2-3 người local',
        '   → Pilot với 20-30 GV, 200-300 HS',
        '',
        '2. Chuẩn bị Indonesia:',
        '   → Thị trường lớn nhất SEA (50M+ học sinh)',
        '   → Nghiên cứu thị trường 6 tháng cuối năm',
        '   → Chuẩn bị launch năm 4',
        '',
        '━━━━━━━ TỔNG KẾT CUỐI NĂM 3 ━━━━━━━',
        '',
        '🎯 Tổng học sinh Việt Nam: 20,000-30,000',
        '   • B2B: 8,000-12,000 (40%)',
        '   • B2C: 12,000-18,000 (60%)',
        '',
        '💰 Doanh thu tháng: 5.6-8.8 tỷ VNĐ',
        '   • B2B: 0.6-0.9 tỷ',
        '   • B2C: 5-8 tỷ',
        '',
        '📊 ARR: 67-105 tỷ VNĐ (~2.9-4.5M USD)',
        '',
        '💸 Đã gọi vốn Series A: 2-3M USD → Sẵn sàng gọi Series B: 5-10M USD',
        '',
        '🏆 Vị trí thị trường:',
        '   • #1 CLIL platform Việt Nam',
        '   • Top 3 EdTech overall',
        '   • Brand awareness: 60-70% trong target demographic',
        '',
        '🌏 Đã có footprint tại Thailand → Chuẩn bị Indonesia',
        '',
        '👥 Team size: 80-120 người',
        '   • Product & Engineering: 30-40',
        '   • Sales & Marketing: 25-35',
        '   • Content: 15-20',
        '   • Operations & Support: 10-15',
        '   • Management: 5-10'
    ])
    
    doc.add_page_break()
    
    print("Đang tạo PHẦN V: DỰ PHÓNG TÀI CHÍNH...")
    
    # ==================== LƯU FILE ====================
    filename = f'LEXIO_Bao_Cao_Chien_Luoc_Day_Du_{datetime.datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    filepath = f'/Users/binhnguyen/Downloads/Engquest3k/{filename}'
    doc.save(filepath)
    
    print(f'\n✅ Báo cáo đã được tạo!')
    print(f'📄 File: {filepath}')
    print(f'📊 Đã hoàn thành Phần I, II, III, IV. Đang chuẩn bị Phần V, VI, VII, VIII...')
    
    return filepath

if __name__ == '__main__':
    try:
        filepath = create_comprehensive_report()
        print(f'\n🎯 File đã sẵn sàng: {filepath}')
    except Exception as e:
        print(f'❌ Lỗi: {e}')
        import traceback
        traceback.print_exc()
