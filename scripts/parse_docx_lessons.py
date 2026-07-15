#!/usr/bin/env python3
"""
Parse W25-36 and W37-53 lesson plan DOCX files into lessonPlans.json.
Generates sessions (3-session default), sessions_2, sessions_5, answer_key, task_cards.

v3: Reads DOCX tables (Quick Ref + Vocab Tiers) and smart-splits merged paragraphs.
"""

import re, json, os, sys
from docx import Document

FILE1 = os.path.join(os.path.dirname(__file__), '../Production_FINAL/1. FINAL MASS PRODUCTION/2_REFERENCE_DOCS/0. NEW_FINAL_Lesson plans_W25-36.docx')
FILE2 = os.path.join(os.path.dirname(__file__), '../Production_FINAL/1. FINAL MASS PRODUCTION/2_REFERENCE_DOCS/0. NEW_FINAL_Lesson plans_W37-53.docx')
OUT   = os.path.join(os.path.dirname(__file__), '../public/data/lessonPlans.json')

# ── helpers ──────────────────────────────────────────────────────────────────

SKIP_RE = re.compile(r'^[━═─✂\-\s]{3,}$')

def load_doc(path):
    """Load DOCX and build body-item index for table-aware extraction."""
    doc = Document(path)
    paras = [(p.text.strip(), p.style.name) for p in doc.paragraphs]

    # Build ordered list of body children: ('p', para_idx) or ('t', tbl_idx)
    # and a reverse map  para_idx -> position in body_items
    body_items = []
    para_to_bi = {}
    p_count = 0
    t_count = 0
    for child in doc.element.body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            para_to_bi[p_count] = len(body_items)
            body_items.append(('p', p_count))
            p_count += 1
        elif tag == 'tbl':
            body_items.append(('t', t_count))
            t_count += 1

    return doc, paras, body_items, para_to_bi


def next_table_after_para(body_items, para_to_bi, para_idx, within_n=15):
    """Return table index of the first table that immediately follows para_idx."""
    bi_pos = para_to_bi.get(para_idx, -1)
    if bi_pos < 0:
        return None
    for j in range(bi_pos + 1, min(bi_pos + within_n, len(body_items))):
        typ, idx = body_items[j]
        if typ == 't':
            return idx
        # Stop scanning if we hit a paragraph with real text (means the table
        # is not immediately following the header – rare but safe guard)
    return None


def tables_in_range(doc, body_items, para_to_bi, start_para, end_para):
    """Yield (tbl_idx, table) for all tables between start_para and end_para."""
    bi_start = para_to_bi.get(start_para, 0)
    bi_end   = para_to_bi.get(end_para, len(body_items)) if end_para < len(para_to_bi) else len(body_items)
    for j in range(bi_start, bi_end):
        typ, idx = body_items[j]
        if typ == 'p' and idx >= end_para:
            break
        if typ == 't':
            yield idx, doc.tables[idx]

def find_week_ranges(paras):
    PAT = re.compile(r'TEACHER CONTENT PACK\s*[—–\-]+\s*WEEK\s+(\d+)', re.I)
    order, ranges = [], {}
    for i, (t, _) in enumerate(paras):
        m = PAT.search(t)  # search not match — handles ═══ prefix
        if m:
            wn = int(m.group(1))
            if order:
                ranges[order[-1]] = (ranges[order[-1]][0], i)
            ranges[wn] = (i, len(paras))
            order.append(wn)
    return ranges, order

def find_session_starts(paras, wstart, wend):
    """Returns {1: idx, 2: idx, 3: idx} for session worksheet starts."""
    found = {}
    PATS = {
        1: [r'SESSION\s+1\s+WORKSHEET', r'SESSION\s+1\s*[|]', r'WEEK\s+\d+\s*[|]\s*SESSION\s+1',
            r'WORKSHEETS\s*[—–\-]\s*SESSION\s+1', r'SECTION\s+4:\s*STUDENT WORKSHEETS\s*—\s*SESSION\s+1'],
        2: [r'SESSION\s+2\s+WORKSHEET', r'SESSION\s+2\s*[|]', r'WEEK\s+\d+\s*[|]\s*SESSION\s+2',
            r'WORKSHEETS\s*[—–\-]\s*SESSION\s+2', r'SECTION\s+5:\s*STUDENT WORKSHEETS'],
        3: [r'SESSION\s+3\s+WORKSHEET', r'SESSION\s+3\s*[|]', r'WEEK\s+\d+\s*[|]\s*SESSION\s+3',
            r'WORKSHEETS\s*[—–\-]\s*SESSION\s+3', r'SECTION\s+6:\s*STUDENT WORKSHEETS'],
    }
    for i in range(wstart, wend):
        t = paras[i][0]
        for sn, pats in PATS.items():
            if sn not in found:
                for pat in pats:
                    if re.search(pat, t, re.I):
                        found[sn] = i
                        break
    return found

def section_end_before(paras, after, wend, *section_pats):
    """Find the earliest paragraph after `after` that matches any section pattern."""
    for i in range(after + 1, wend):
        t = paras[i][0]
        for pat in section_pats:
            if re.search(pat, t, re.I):
                return i
    return wend

# Numbered exercise item pattern: lines like "1.", "2." etc.
ITEM_NUM_RE = re.compile(r'^\s*(\d+)[.)\s]')

def add_subtotal(content):
    """Append a sub-total score line based on count of numbered items in the part."""
    numbered = sum(1 for line in content if ITEM_NUM_RE.match(line))
    if numbered >= 3:
        content.append(f'[ Sub-total: _____ / {numbered} ]')
    return content

def parse_session(paras, start, end):
    """Parse session paragraphs into PART objects [{title, content}]."""
    PART_START = re.compile(
        r'^(SPIRAL REVIEW|PART\s+[1-9][\s:\-]|📚\s*WEEK|\*\*WEEK|\bName:\b)', re.I
    )
    parts, cur_title, cur_content = [], None, []

    for i in range(start, min(end, len(paras))):
        t = paras[i][0]
        if not t or SKIP_RE.match(t):
            continue
        if PART_START.match(t):
            if cur_title is not None:
                parts.append({'title': cur_title, 'content': add_subtotal(cur_content)})
            cur_title, cur_content = t, []
        elif cur_title is not None:
            for sub in smart_split(t):
                cur_content.append(sub)  # keep ALL content, split merged lines
        else:
            # Before first known part — treat as header
            if re.search(r'Name:|📚|WEEK\s+\d+', t, re.I):
                cur_title, cur_content = t, []

    if cur_title is not None:
        parts.append({'title': cur_title, 'content': add_subtotal(cur_content)})
    return parts

def extract_lines(paras, start, wend, stop_pats):
    """Extract non-empty lines from start+1 until a stop pattern is found."""
    lines = []
    for i in range(start + 1, wend):
        t = paras[i][0]
        if not t or SKIP_RE.match(t):
            continue
        for pat in stop_pats:
            if re.search(pat, t, re.I):
                return lines
        lines.append(t)  # no length limit
    return lines

def find_section(paras, wstart, wend, *patterns):
    for i in range(wstart, wend):
        t = paras[i][0]
        for pat in patterns:
            if re.search(pat, t, re.I):
                return i
    return None

def smart_split(text):
    """Split a merged DOCX paragraph into separate logical lines.

    Splits on structural sub-item markers while preserving the marker at the
    start of each resulting line.
    """
    if len(text) < 60:
        return [text]
    SEP = '\x00'
    t = text
    # Before [O ... (Oreo / Optional / Bonus markers)
    t = re.sub(r'\s*(?=\[O[\s\[])', SEP, t, flags=re.I)
    # Before Panel N:
    t = re.sub(r'\s*(?=Panel\s+\d+:)', SEP, t, flags=re.I)
    # Before Stage N —
    t = re.sub(r'\s*(?=Stage\s+\d+\s*[—–\-])', SEP, t, flags=re.I)
    # Before T / F:
    t = re.sub(r'\s*(?=T\s*/\s*F:)', SEP, t, flags=re.I)
    # Before [ALL] or [O only]
    t = re.sub(r'\s*(?=\[(ALL|O only)\])', SEP, t, flags=re.I)
    # → separator: split and put → at start of continuation
    t = re.sub(r'\s*→\s*', SEP + '→ ', t)
    # Alphabetic sub-items: "... a. word: " or "... b. word: "
    t = re.sub(r'(?<=\S)\s+(?=[a-e]\.\s)', SEP, t)
    parts = [p.strip() for p in t.split(SEP) if p.strip()]
    return parts if len(parts) > 1 else [text]


def extract_quick_ref(doc, body_items, para_to_bi, paras, wstart, wend):
    s1 = find_section(paras, wstart, wend, r'SECTION\s+1:\s*QUICK\s*REF')
    if s1 is None:
        return {}

    # Primary: read the Quick Ref table that immediately follows the SECTION 1 header
    tbl_idx = next_table_after_para(body_items, para_to_bi, s1)
    if tbl_idx is not None:
        tbl = doc.tables[tbl_idx]
        if len(tbl.columns) >= 2:
            result = {}
            for row in tbl.rows:
                k = row.cells[0].text.strip()
                v = row.cells[1].text.strip()
                if k and k.lower() not in ('category', ''):
                    result[k] = v
            if result:
                return result

    # Fallback: parse plain paragraphs between SECTION 1 and SECTION 2
    stop = section_end_before(paras, s1, wend, r'SECTION\s+[23456789]:')
    result = {}
    for i in range(s1 + 1, stop):
        t = paras[i][0]
        if not t or SKIP_RE.match(t):
            continue
        if ':' in t:
            k, v = t.split(':', 1)
            result[k.strip()[:80]] = v.strip()
        else:
            result[f'note{i}'] = t
    return result

def extract_methodology(paras, wstart, wend):
    s2 = find_section(paras, wstart, wend, r'SECTION\s+2:\s*TEACHING\s*METHODOLOGY')
    if s2 is None:
        return []
    stop = section_end_before(paras, s2, wend, r'SECTION\s+[3456789]:')
    sections, cur = [], None
    for i in range(s2 + 1, stop):
        t, sty = paras[i]
        if not t or SKIP_RE.match(t):
            continue
        is_heading = sty.startswith('Heading') or re.match(r'^2\.\d+\s+', t)
        if is_heading:
            if cur:
                sections.append(cur)
            cur = {'title': t[:120], 'content': []}
        elif cur:
            cur['content'].append(t)
        else:
            cur = {'title': 'Overview', 'content': [t[:300]]}
    if cur:
        sections.append(cur)
    return sections

def extract_vocab_tiers(doc, body_items, para_to_bi, paras, wstart, wend):
    s3 = find_section(paras, wstart, wend, r'SECTION\s+3:\s*VOCAB')
    if s3 is None:
        return []

    s4 = find_section(paras, s3 + 1, wend, r'SECTION\s+4:')
    end_para = s4 if s4 is not None else wend

    vocab_rows = []
    seen_words = set()
    for tbl_idx, tbl in tables_in_range(doc, body_items, para_to_bi, s3, end_para):
        if len(tbl.rows) < 2 or len(tbl.columns) < 2:
            continue
        headers = [c.text.strip() for c in tbl.rows[0].cells]
        # Only word tables (first column is 'Word')
        if not headers or headers[0].lower() != 'word':
            continue
        # Skip SRS Schedule (has S1/Homework columns — no educational vocab data)
        if 'S1' in headers or 'Homework' in headers:
            continue
        # Skip tables with no Vietnamese column
        if 'Vietnamese' not in headers:
            continue
        # Normalise column name: SRS Map uses 'Key Collocation' (no trailing (s))
        norm_headers = ['Key Collocation(s)' if h == 'Key Collocation' else h for h in headers]
        for row in tbl.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells[0]:
                continue
            # Deduplicate: Tier 2 Table comes first; skip same word from SRS Map
            if cells[0] in seen_words:
                continue
            seen_words.add(cells[0])
            entry = {norm_headers[j]: cells[j] for j in range(min(len(norm_headers), len(cells)))}
            vocab_rows.append(entry)

    # Fallback: paragraph lines if no tables found
    if not vocab_rows:
        lines = extract_lines(paras, s3, wend, [r'SECTION\s+[456789]:'])
        return [{'title': 'Vocabulary Tiers & SRS Schedule', 'content': lines[:50]}]

    return vocab_rows

def extract_answer_key(paras, wstart, wend):
    # Primary: look for dedicated ANSWER KEYS section
    ak = find_section(paras, wstart, wend,
        r'SECTION\s+\d+:\s*ANSWER\s*KEY',
        r'SECTION\s+\d+:\s*ANSWER\s*KEYS',
        r'S1\s+ANSWER\s+KEY',
    )
    if ak is not None:
        raw = extract_lines(paras, ak, wend, [
            r'SECTION\s+\d+:',           # stop at any next major section header
            r'APPENDIX|VEO3|CHANTING\s+VIDEO|SHADOWING\s+VIDEO'
        ])
        # Split merged "PART X (...): ... PART Y (...): ..." lines
        lines = []
        for r in raw:
            lines.extend(_split_ak_line(r))
        return lines[:150]
    # Fallback: collect all ✍️ Answer Key: and S1/S2/S3 ANSWER KEY inline markers
    inline_pats = re.compile(r'Answer\s+Key|S[123]\s+ANSWER\s+KEY', re.I)
    lines = []
    for i in range(wstart, wend):
        t = paras[i][0]
        if inline_pats.search(t):
            lines.extend(_split_ak_line(t))
    return lines[:60]

def _split_ak_line(text):
    """Split a long merged answer key line into per-part lines.

    Splits on 'SPIRAL REVIEW:', 'PART N (...):' and 'Session N' markers
    that appear mid-string (not at the very start).
    """
    # Split before ' PART N' or ' Session N' markers that appear in the middle
    SEP_PAT = re.compile(
        r'(?<=\S)\s+(PART\s+\d+|Session\s+[123]|SPIRAL\s+REVIEW|S[123]\s+ANSWER)',
        re.I
    )
    parts_raw = SEP_PAT.split(text)
    if len(parts_raw) <= 1:
        return [text.strip()] if text.strip() else []
    # Re-join split tokens: split() eats the capture group, so reassemble
    # re.split with capturing group interleaves separators
    tokens = SEP_PAT.split(text)
    # Use re.findall to get the separators
    seps = SEP_PAT.findall(text)
    result = [tokens[0].strip()]
    for sep, body in zip(seps, tokens[1:]):
        combined = (sep + ' ' + body).strip()
        if combined:
            result.append(combined)
    return [r for r in result if r]

def _split_video_lines(text):
    """Split a (possibly merged) paragraph into individual [WXX...] prompt lines."""
    # Remove leading bullet characters
    text = re.sub(r'^[•\-\*\s]+', '', text.strip())
    # Split on [W boundaries — each prompt starts with [WNN.
    parts = re.split(r'(?=\[W\d+\.(?:CHANT|SHADOW)\.)', text)
    result = []
    for part in parts:
        part = part.strip().lstrip('•').strip()
        if part.startswith('[') and ']' in part:
            result.append(part)
    return result


def extract_video_prompts(paras, wnum, wstart, wend):
    """Extract chanting and shadowing video prompts from after the answer key section."""
    CHANT_HDR  = re.compile(r'chanting\s*(video)?\s*[—\-]?\s*session\s*\d*|chanting\s+s\d|🎵', re.I)
    SHADOW_HDR = re.compile(r'shadowing\s*(video)?\s*[—\-]?\s*session\s*(\d)|shadowing\s+s(\d)|🎤', re.I)
    PROMPT_PAT = re.compile(r'\[W\d+\.(CHANT|SHADOW)\.', re.I)

    # Start scanning from after the answer key (AK stop includes CHANTING/SHADOWING keywords)
    ak_end = find_section(paras, wstart, wend,
        r'APPENDIX.*VIDEO\s+PROMPT',
        r'VEO3\s+VIDEO',
        r'CHANTING\s+VIDEO',
        r'Chanting\s+S\d',
    ) or wend

    chanting  = []
    shadowing = {}      # {session_int: [prompt_str, ...]}
    mode      = None    # 'chant' | 'shadow'
    shadow_sn = None

    for i in range(ak_end, wend):
        t = paras[i][0]
        if not t:
            continue

        has_chant_hdr  = CHANT_HDR.search(t)
        has_shadow_hdr = SHADOW_HDR.search(t)
        has_prompt     = PROMPT_PAT.search(t)

        # Detect chanting section header (may or may not contain first prompt)
        if has_chant_hdr and not has_shadow_hdr:
            mode = 'chant'
            shadow_sn = None
            if has_prompt:
                # Merged: heading + first prompt on same paragraph (W26 style)
                for p in _split_video_lines(t):
                    chanting.append(p)
            continue

        # Detect shadowing section header (may or may not contain first prompt)
        if has_shadow_hdr:
            sn_match = re.search(r'S(?:ession)?\s*(\d)', t, re.I)
            sn = None
            if sn_match:
                sn = int(sn_match.group(1))
            else:
                ms2 = SHADOW_HDR.search(t)
                g2, g3 = (ms2.group(2) if ms2 else None), (ms2.group(3) if ms2 else None)
                sn = int(g2 or g3) if (g2 or g3) else (shadow_sn + 1 if shadow_sn else 2)
            mode = 'shadow'
            shadow_sn = sn
            shadowing.setdefault(sn, [])
            if has_prompt:
                for p in _split_video_lines(t):
                    shadowing[sn].append(p)
            continue

        # Detect informal plain-text shadowing headers like "Shadowing S2:" or "Shadowing S3:"
        plain_shadow = re.match(r'^\s*Shadowing\s+S(\d)\s*:?\s*$', t, re.I)
        if plain_shadow:
            sn = int(plain_shadow.group(1))
            mode = 'shadow'
            shadow_sn = sn
            shadowing.setdefault(sn, [])
            continue

        # Collect prompts if in a video section
        if has_prompt:
            for p in _split_video_lines(t):
                if mode == 'chant':
                    chanting.append(p)
                elif mode == 'shadow' and shadow_sn is not None:
                    shadowing.setdefault(shadow_sn, []).append(p)

    if not chanting and not shadowing:
        return {}

    result = {}
    if chanting:
        result['chanting'] = [{'title': f'Chanting — W{wnum} ({len(chanting)} prompts)', 'script': '\n'.join(chanting)}]

    shadow_list = []
    for sn in sorted(shadowing.keys()):
        items = shadowing[sn]
        if items:
            shadow_list.append({'session': sn, 'title': f'Shadowing — Session {sn} ({len(items)} prompts)', 'sentences': items})
    if shadow_list:
        result['shadowing'] = shadow_list

    return result


def extract_task_cards(paras, wstart, wend):
    tc = find_section(paras, wstart, wend, r'SECTION\s+\d+:\s*COMMUNICATIVE\s+TASK\s+CARDS')
    if tc is None:
        return []
    lines = extract_lines(paras, tc, wend,
        [r'SECTION\s+\d+:\s*(QA|ANSWER|RUBRIC|APPENDIX)'])
    return lines[:150]

# ── session variant builders ──────────────────────────────────────────────────

def parts_by_num(parts, low, high):
    result = []
    for p in parts:
        m = re.search(r'PART\s+(\d+)', p['title'], re.I)
        if m:
            n = int(m.group(1))
            if low <= n <= high:
                result.append(p)
        elif re.search(r'SPIRAL REVIEW|📚|Name:', p['title'], re.I) and low == 0:
            result.append(p)
    return result

# ── task-card / answer-key splitters ─────────────────────────────────────────

def split_task_cards_by_session(lines):
    """Split flat task_cards list into {s1:[], s2:[], s3:[]} by session header."""
    groups = {'s1': [], 's2': [], 's3': []}
    current = None
    SESS_PAT = re.compile(r'\bS([123])\s*(Task\b|—|-|TASK\b)|^S([123])\b', re.I)
    for line in lines:
        m = SESS_PAT.search(line)
        if m:
            n = m.group(1) or m.group(3)
            if n:
                current = f's{n}'
        if current:
            groups[current].append(line)
    # If nothing was tagged, fallback everything to s1
    if not any(groups.values()):
        groups['s1'] = lines[:]
    return groups


def split_answer_key_by_session(lines):
    """Split flat answer_key list into {s1:[], s2:[], s3:[]} by session markers.

    Handles patterns:
      - 'S1 ANSWER KEY' / 'S2 ANSWER KEY' / 'S3 ANSWER KEY'
      - 'Session 1' / 'Session 2' / 'Session 3'  (bare, e.g. W31/W37-W43)
      - 'Session 1:' / 'Session 1 -' etc.
    """
    groups = {'s1': [], 's2': [], 's3': []}
    current = 's1'
    # Match session number from marker lines; handle both 'S1 ANSWER KEY' and 'Session N' forms
    SESS_PAT = re.compile(
        r'\bS([123])\s+ANSWER\s*KEY\b'          # S1 ANSWER KEY
        r'|^\s*Session\s+([123])\s*[:\-]?\s*$'  # "Session 1" or "Session 1:" as whole line
        r'|^\s*Session\s+([123])\s*[:\-]',       # "Session 1:" or "Session 1 -" at line start
        re.I | re.M
    )
    for line in lines:
        m = SESS_PAT.search(line)
        if m:
            n = m.group(1) or m.group(2) or m.group(3)
            if n:
                current = f's{n}'
        groups[current].append(line)
    return groups


def generate_games(wnum, unit_theme, quick_ref, vocab_list):
    """Generate 3 pedagogically-typed games per week, parameterised by weekly content."""
    theme = unit_theme or f'Week {wnum}'
    grammar_pt = quick_ref.get('Grammar Focus', quick_ref.get('Grammar',
                               "this week's grammar point"))
    vocab_words = [v.get('Word', '') for v in (vocab_list or [])[:5]
                   if isinstance(v, dict) and v.get('Word')]
    vocab_str = ', '.join(vocab_words) if vocab_words else 'week vocabulary'
    return [
        {
            'id': f'w{wnum}_g1',
            'name': 'Vocab Slam',
            'type': 'Vocabulary',
            'duration': '10–12 min',
            'players': '2–4 per group',
            'session_fit': 'All formats — best after vocab intro (P3) in any session',
            'materials': f'Vocab flashcards W{wnum}: {vocab_str}',
            'instructions': [
                'Round 1: Teacher calls a Vietnamese meaning — first student to say the English word '
                'AND use it correctly in a sentence wins the card.',
                'Round 2: Reverse — teacher calls English word, student gives Vietnamese + memory trick.',
                f'Bonus round: Students chain all vocab words ({vocab_str}) into one connected story '
                'using sequence connectors from this week.',
                'Team variant: 2 teams, 10 cards per round. Team with highest score wins.',
            ],
        },
        {
            'id': f'w{wnum}_g2',
            'name': 'Grammar Relay',
            'type': 'Grammar & Error Correction',
            'duration': '12–15 min',
            'players': '3–6 (team relay)',
            'session_fit': 'Best in Drill slot (Slot 2) or after P4/P6 in 3-session mode',
            'materials': f'10 sentence cards — 5 correct, 5 with errors based on: {grammar_pt}',
            'instructions': [
                f'Grammar focus this week: {grammar_pt}.',
                'Speed Read: Display sentence for 5 sec. Teams write C (correct) or E (error) on '
                'whiteboards simultaneously — no copying allowed.',
                'Fix-It: For error sentences, team that corrects fastest AND explains the rule '
                'earns double points.',
                'Build: Each team member adds 1 original sentence using the grammar target. '
                'The chain must be logically connected — each sentence must link to the previous.',
            ],
        },
        {
            'id': f'w{wnum}_g3',
            'name': 'Story Race',
            'type': 'Speaking & Production',
            'duration': '15–20 min',
            'players': 'Whole class (pairs)',
            'session_fit': 'All formats — best in final slot of the week (S3 / Slot 5 / Session B)',
            'materials': f'Story starter card: theme "{theme}"',
            'instructions': [
                f'Give each pair a story starter connected to this week\'s theme: "{theme}".',
                'Pair A adds 2 sentences — must include at least 2 vocab words from this week.',
                'Pass to Pair B: they MUST correct any grammar mistake from Pair A first, '
                'then add 2 more sentences using different vocab words.',
                f'Final round: Each pair presents their story ending aloud. Class votes on '
                f'most creative + most accurate ({vocab_str} usage). Teacher scores grammar precision.',
            ],
        },
    ]

# ── session variant builders ──────────────────────────────────────────────────

def build_sessions_2(s1, s2, s3, tc_groups=None, ak_groups=None):
    """
    2 sessions/week:

    Session A — Nền tảng (Foundation) · 120 min
      Full S1 — complete input→structure→production cycle at baseline level.
      Task Cards: S1.

    Session B — Nâng cao (Advanced) · 120 min
      Full S2 — same 9 PARTs + Spiral but at intermediate difficulty.
      Task Cards: S2.

    S3 is omitted; teachers can access S3 content via the Task Cards tab
    for extension/challenge activities.
    """
    tc = tc_groups or {}
    ak = ak_groups or {}
    return [
        {
            'session': 1,
            'session_label': 'Session A — Nền tảng (Foundation) · 120 min',
            'parts': s1 or [],
            'task_cards': tc.get('s1') or [],
            'answer_key': ak.get('s1') or [],
        },
        {
            'session': 2,
            'session_label': 'Session B — Nâng cao (Advanced) · 120 min',
            'parts': s2 or [],
            'task_cards': tc.get('s2') or [],
            'answer_key': ak.get('s2') or [],
        },
    ]


def build_sessions_5(s1, s2, s3, tc_groups=None, ak_groups=None):
    """
    5 sessions/week — Correct PART map:
      P1=Reading  P2=Vocab  P3=Sentence Building
      P4=Listening  P5=Error Correction  P6=STEM/CLIL  P7=Quick Production
      P8=Portfolio (Writing)  P9=Homework

    Slot 1 — Activate / Khởi động (~55–60 min):
      S1: session header + Spiral Review + P1 + P2 + P3
      → Pure Input phase: SRS refresh, new text, vocab intro. No grading.

    Slot 2 — Drill / Luyện tập (~80–85 min):
      S1: P4 + P5 + P6 + P7 + P8 + P9
      → Full Structure + Production + Homework. S1 Task Cards.
      → AK: S1 answer key.

    Slot 3 — Bridge / Kết nối (~55–60 min):
      S2: session header + Spiral Review + P1 + P2 + P3
      → S2 Input phase with harder text/vocab. No grading.

    Slot 4 — Challenge / Thử thách (~80–85 min):
      S2: P4 + P5 + P6 + P7 + P8 + P9
      → S2 Structure + Production + Homework. S2 Task Cards.
      → AK: S2 answer key.

    Slot 5 — Perform / Biểu diễn (~80–85 min):
      S3: P4 + P5 + P6 + P7 + P8 + P9
      → S3 abbreviated: skip S3 P1-P3 (same text/vocab reviewed twice already).
        Peak production: P7 Quick Production + P8 Portfolio at highest level.
        S3 Task Cards for communicative peak of the week.
      → AK: S3 answer key (or S2 if S3 not available).
    """
    tc = tc_groups or {}
    ak = ak_groups or {}

    def _pick(parts, lo, hi, with_prefix=False):
        """Select numbered PARTs lo..hi. with_prefix=True also grabs all non-PART
        items that appear before PART 1 (session header, Name: line, Spiral Review)."""
        prefix_parts = []
        numbered_parts = []
        seen_first_part = False
        for p in (parts or []):
            m = re.search(r'PART\s+(\d+)', p['title'], re.I)
            n = int(m.group(1)) if m else None
            if n is not None:
                seen_first_part = True
                if lo <= n <= hi:
                    numbered_parts.append(p)
            elif not seen_first_part and with_prefix:
                prefix_parts.append(p)
        return prefix_parts + numbered_parts

    g1 = _pick(s1, 1, 3, with_prefix=True)   # S1 Input
    g2 = _pick(s1, 4, 9)                      # S1 Structure+Production+HW
    g3 = _pick(s2, 1, 3, with_prefix=True)   # S2 Input
    g4 = _pick(s2, 4, 9)                      # S2 Structure+Production+HW
    g5 = _pick(s3, 4, 9)                      # S3 abbreviated (P4-P9 only)

    SLOTS = [
        ('Activate', 'Khởi động — Spiral, Reading & Vocab S1'),
        ('Drill',    'Luyện tập — Listening, Error Correction & Writing S1'),
        ('Bridge',   'Kết nối — Spiral, Reading & Vocab S2'),
        ('Challenge','Thử thách — Listening, Error Correction & Writing S2'),
        ('Perform',  'Biểu diễn — Error Correction, Writing & Task Cards S3'),
    ]
    # Task Cards: grading/communication only in production slots (2, 4, 5)
    TC_SLOTS = [None, tc.get('s1') or [], None, tc.get('s2') or [], tc.get('s3') or []]
    # Answer Keys: only needed when grading production work (Slots 2, 4, 5)
    s3_ak = ak.get('s3') or ak.get('s2') or []  # fallback if S3 AK not in source
    AK_SLOTS = [[], ak.get('s1') or [], [], ak.get('s2') or [], s3_ak]

    return [
        {
            'session': i + 1,
            'session_label': f'Slot {i+1} — {en} / {vi}',
            'parts': grp or [],
            'task_cards': tc_slot or [],
            'answer_key': ak_slot or [],
        }
        for i, ((en, vi), grp, tc_slot, ak_slot)
        in enumerate(zip(SLOTS, [g1, g2, g3, g4, g5], TC_SLOTS, AK_SLOTS))
    ]

# ── week title extraction ─────────────────────────────────────────────────────

def extract_unit_theme(paras, wstart):
    header = paras[wstart][0]
    # "TEACHER CONTENT PACK — WEEK 25 | BLOCK B\nUnit 5: My Weekend..."
    # Sometimes on same line: "TEACHER CONTENT PACK — WEEK 26 | BLOCK B Unit 9: Memories..."
    m = re.search(r'Unit\s+\d+[:\s]+(.+?)(?:\s+Integrated|\s+v\d|$)', header, re.I)
    if m:
        return m.group(1).strip()[:120]
    # Check next lines
    for i in range(wstart + 1, min(wstart + 5, len(paras))):
        t = paras[i][0]
        if not t or re.match(r'SECTION|Integrated|v\d+\.\d+|═', t, re.I):
            continue
        m2 = re.search(r'Unit\s+\d+[:\s]+(.+)', t, re.I)
        if m2:
            return m2.group(1).strip()[:120]
        if len(t) > 10 and len(t) < 120:
            return t
    return ''

# ── main processor ────────────────────────────────────────────────────────────

def process_week(doc, paras, body_items, para_to_bi, wnum, wstart, wend):
    unit_theme = extract_unit_theme(paras, wstart)

    # Find session ranges
    ss = find_session_starts(paras, wstart, wend)
    skeys = sorted(ss.keys())

    def sess_range(sn):
        s = ss.get(sn)
        if s is None:
            return None, None
        # end = start of next session, or start of SECTION 5 (portfolio ref), or wend
        next_sn = sn + 1
        e = ss.get(next_sn)
        if e is None:
            e = section_end_before(paras, s, wend,
                r'SECTION\s+[5-9]\d*:', r'SECTION\s+10:', r'SECTION\s+11:')
        return s, e

    s1s, s1e = sess_range(1)
    s2s, s2e = sess_range(2)
    s3s, s3e = sess_range(3)

    s1_parts = parse_session(paras, s1s, s1e) if s1s else []
    s2_parts = parse_session(paras, s2s, s2e) if s2s else []
    s3_parts = parse_session(paras, s3s, s3e) if s3s else []

    sessions = []
    if s1_parts:
        sessions.append({'session': 1, 'session_label': f'Session 1 (120 min)', 'parts': s1_parts})
    if s2_parts:
        sessions.append({'session': 2, 'session_label': f'Session 2 (120 min)', 'parts': s2_parts})
    if s3_parts:
        sessions.append({'session': 3, 'session_label': f'Session 3 (120 min)', 'parts': s3_parts})

    tc_raw    = extract_task_cards(paras, wstart, wend)
    ak_raw    = extract_answer_key(paras, wstart, wend)
    tc_groups = split_task_cards_by_session(tc_raw)
    ak_groups = split_answer_key_by_session(ak_raw)

    qr   = extract_quick_ref(doc, body_items, para_to_bi, paras, wstart, wend)
    meth = extract_methodology(paras, wstart, wend)
    vt   = extract_vocab_tiers(doc, body_items, para_to_bi, paras, wstart, wend)
    vp   = extract_video_prompts(paras, wnum, wstart, wend)

    sessions_2 = build_sessions_2(s1_parts, s2_parts, s3_parts, tc_groups, ak_groups)
    sessions_5 = build_sessions_5(s1_parts, s2_parts, s3_parts, tc_groups, ak_groups)
    games      = generate_games(wnum, unit_theme, qr, vt)

    return {
        'week': wnum,
        'unit_theme': unit_theme,
        'quick_ref': qr,
        'methodology': meth,
        'vocab_tiers': vt,
        'sessions': sessions,
        'sessions_2': sessions_2,
        'sessions_5': sessions_5,
        'answer_key': ak_raw,
        'answer_key_by_session': ak_groups,
        'task_cards': tc_raw,
        'task_cards_by_session': tc_groups,
        'games': games,
        'video_prompts': vp,
    }

# ── entry point ───────────────────────────────────────────────────────────────

def main():
    print('Loading DOCX files...', flush=True)
    doc1, paras1, body_items1, para_to_bi1 = load_doc(FILE1)
    doc2, paras2, body_items2, para_to_bi2 = load_doc(FILE2)
    print(f'  File 1: {len(paras1)} paragraphs', flush=True)
    print(f'  File 2: {len(paras2)} paragraphs', flush=True)

    ranges1, order1 = find_week_ranges(paras1)
    ranges2, order2 = find_week_ranges(paras2)
    print(f'  File 1 weeks: {order1}', flush=True)
    print(f'  File 2 weeks: {order2}', flush=True)

    result = {}
    for wnum in order1:
        wstart, wend = ranges1[wnum]
        print(f'  Processing W{wnum}...', flush=True, end='')
        result[str(wnum)] = process_week(doc1, paras1, body_items1, para_to_bi1, wnum, wstart, wend)
        s3 = len(result[str(wnum)]['sessions'])
        qr = len(result[str(wnum)]['quick_ref'])
        vt = len(result[str(wnum)]['vocab_tiers'])
        ak = len(result[str(wnum)]['answer_key'])
        tc = len(result[str(wnum)]['task_cards'])
        print(f' {s3} sessions, QR={qr}, vocab={vt}, {ak} AK, {tc} TC', flush=True)

    for wnum in order2:
        wstart, wend = ranges2[wnum]
        print(f'  Processing W{wnum}...', flush=True, end='')
        result[str(wnum)] = process_week(doc2, paras2, body_items2, para_to_bi2, wnum, wstart, wend)
        s3 = len(result[str(wnum)]['sessions'])
        qr = len(result[str(wnum)]['quick_ref'])
        vt = len(result[str(wnum)]['vocab_tiers'])
        ak = len(result[str(wnum)]['answer_key'])
        tc = len(result[str(wnum)]['task_cards'])
        print(f' {s3} sessions, QR={qr}, vocab={vt}, {ak} AK, {tc} TC', flush=True)

    # Write monolithic file (kept for backward compat)
    print(f'\nWriting {len(result)} weeks to {OUT}...', flush=True)
    with open(OUT, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    # Write per-week index for lazy-load — to BOTH public/data/ and mcp-server/data/
    index = {wk: {'week': v['week'], 'unit_theme': v['unit_theme']} for wk, v in result.items()}
    for idx_path in [
        os.path.join(os.path.dirname(OUT), 'lessonPlans_index.json'),
        os.path.join(os.path.dirname(__file__), '../mcp-server/data/lessonPlans_index.json'),
    ]:
        with open(idx_path, 'w', encoding='utf-8') as f:
            json.dump(index, f, ensure_ascii=False, indent=2)

    # Write per-week files to mcp-server/data/lessons/ (authenticated, behind API)
    mcp_dir = os.path.join(os.path.dirname(__file__), '../mcp-server/data/lessons')
    os.makedirs(mcp_dir, exist_ok=True)
    for wk, v in result.items():
        with open(os.path.join(mcp_dir, f'W{wk}.json'), 'w', encoding='utf-8') as f:
            json.dump(v, f, ensure_ascii=False, indent=2)
    print(f'  Per-week files written to {mcp_dir}/', flush=True)

    size_mb = os.path.getsize(OUT) / 1024 / 1024
    print(f'Done! {size_mb:.2f} MB, {len(result)} weeks', flush=True)

    # Validation
    missing_sessions = [k for k, v in result.items() if len(v['sessions']) < 3]
    missing_ak = [k for k, v in result.items() if len(v['answer_key']) == 0]
    missing_vp = [k for k, v in result.items() if not v.get('video_prompts')]
    def has_s2_s3_cards(cards):
        text = ' '.join(cards)
        return bool(re.search(r'S[23]\s*[—–\-]|Week\s+\d+\s+S[23]', text))
    tc_only_s1 = [k for k, v in result.items()
                  if v['task_cards'] and not has_s2_s3_cards(v['task_cards'])]
    if missing_sessions:
        print(f'WARNING — weeks with <3 sessions: {missing_sessions}', flush=True)
    if missing_ak:
        print(f'WARNING — weeks with no answer key: {missing_ak}', flush=True)
    if missing_vp:
        print(f'NOTE — weeks with no video prompts parsed: {missing_vp}', flush=True)
    if tc_only_s1:
        print(f'NOTE — weeks with only S1 task cards: {tc_only_s1}', flush=True)

if __name__ == '__main__':
    main()
