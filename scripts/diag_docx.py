#!/usr/bin/env python3
"""Diagnostic: print all SECTION headers and key structural markers per week."""
import re, sys
from docx import Document

FILE1 = "Production_FINAL/1. FINAL MASS PRODUCTION/2_REFERENCE_DOCS/0. NEW_FINAL_Lesson plans_W25-36.docx"
FILE2 = "Production_FINAL/1. FINAL MASS PRODUCTION/2_REFERENCE_DOCS/0. NEW_FINAL_Lesson plans_W37-53.docx"

TARGET_WEEKS = list(range(25, 54))  # all weeks

WEEK_PAT    = re.compile(r'TEACHER CONTENT PACK.*?WEEK\s+(\d+)', re.I)
SECTION_PAT = re.compile(r'SECTION\s+(\d+)\s*:', re.I)
SESSION_PAT = re.compile(r'SESSION\s+([123])\s+(WORKSHEET|ANSWER|TASK)', re.I)
CARD_PAT    = re.compile(r'Card\s+[AB]|Info.?Gap|STUDENT\s+[AB]|TASK\s+CARD', re.I)
AK_PAT      = re.compile(r'ANSWER\s+KEY|S[123]\s+ANSWER|SECTION\s+\d+.*ANSWER', re.I)
VOCAB_PAT   = re.compile(r'SECTION\s+3|VOCAB\s+MAP|VOCABULARY.*TIER', re.I)
TC_PAT      = re.compile(r'COMMUNICATIVE|TASK\s+CARD', re.I)

def get_paras_and_ranges(fpath):
    doc = Document(fpath)
    paras = [p.text.strip() for p in doc.paragraphs]
    week_ranges, order = {}, []
    for i, t in enumerate(paras):
        m = WEEK_PAT.search(t)
        if m:
            wn = int(m.group(1))
            if order:
                week_ranges[order[-1]] = (week_ranges[order[-1]][0], i)
            week_ranges[wn] = (i, len(paras))
            order.append(wn)
    return paras, week_ranges, order, doc

def inspect_week(paras, doc, ws, we, wn):
    print(f"\n{'='*70}")
    print(f"WEEK {wn}  paras {ws}–{we}")
    print(f"{'='*70}")

    # Count tables in week range using body element order
    body = doc.element.body
    children = list(body)
    p_count = 0
    t_count = 0
    tbl_positions = []  # (para_before, tbl_idx)
    last_para_before = 0
    for child in children:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            p_count += 1
        elif tag == 'tbl':
            tbl_positions.append((p_count, t_count))
            t_count += 1

    # Tables in week para range
    week_tables = [(pp, ti) for pp, ti in tbl_positions if ws <= pp < we]
    print(f"  Tables in range: {len(week_tables)}")
    for pp, ti in week_tables[:20]:
        tbl = doc.tables[ti]
        try:
            first_row = [c.text.strip()[:30] for c in tbl.rows[0].cells]
        except Exception:
            first_row = ['(error)']
        # para context
        ctx = paras[pp][:60] if pp < len(paras) else ''
        print(f"    tbl#{ti} after para[{pp}] '{ctx}' | headers: {first_row}")

    # Key structural paragraphs
    printed = set()
    for i in range(ws, min(we, ws+700)):
        t = paras[i]
        if not t:
            continue
        is_key = (SECTION_PAT.search(t) or SESSION_PAT.search(t) or
                  CARD_PAT.search(t) or AK_PAT.search(t) or TC_PAT.search(t))
        if is_key and i not in printed:
            printed.add(i)
            print(f"  [{i:5d}] {t[:120]}")

for fpath in [FILE1, FILE2]:
    print(f"\n\n{'#'*70}")
    print(f"FILE: {fpath}")
    print(f"{'#'*70}")
    paras, ranges, order, doc = get_paras_and_ranges(fpath)
    print(f"Total paras: {len(paras)}, weeks: {order}")
    for wn in order:
        ws, we = ranges[wn]
        inspect_week(paras, doc, ws, we, wn)
