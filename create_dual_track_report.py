#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LEXIO DUAL-TRACK STRATEGY REPORT
Phân tích chiến lược B2B + B2C song song
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def add_heading_custom(doc, text, level=1, color=RGBColor(31, 41, 55)):
    """Add custom styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return heading

def add_table_styled(doc, data, headers, col_widths=None):
    """Add styled table with headers"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    if col_widths:
        for idx, width in enumerate(col_widths):
            for cell in table.columns[idx].cells:
                cell.width = Inches(width)
    
    # Header row
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.size = Pt(11)
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4F46E5')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for idx, cell_data in enumerate(row_data):
            row_cells[idx].text = str(cell_data)
            row_cells[idx].paragraphs[0].runs[0].font.size = Pt(10)
    
    return table

def create_dual_track_report():
    """Create comprehensive dual-track strategy report"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # ============= COVER PAGE =============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CHIẾN LƯỢC DUAL-TRACK\n')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 41, 55)
    
    run = title.add_run('B2B + B2C SONG SONG\n')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('LEXIO - CHIẾN LƯỢC TĂNG TRƯỞNG TOÀN DIỆN')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(107, 114, 128)
    
    doc.add_paragraph('\n' * 2)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'Phiên bản: 2.0 - DUAL-TRACK STRATEGY\n')
    run.font.size = Pt(11)
    run.font.bold = True
    run = info.add_run(f'Ngày: {datetime.datetime.now().strftime("%d/%m/%Y")}\n')
    run.font.size = Pt(11)
    run = info.add_run('Phân tích: Nhóm Chiến lược & Định giá')
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ============= EXECUTIVE SUMMARY =============
    add_heading_custom(doc, 'TÓM TẮT ĐIỀU HÀNH', level=1, color=RGBColor(220, 38, 38))
    
    summary = doc.add_paragraph()
    run = summary.add_run('PHÁT HIỆN QUAN TRỌNG: ')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(220, 38, 38)
    
    run = summary.add_run(
        'Chiến lược B2B-only hoặc B2C-only đều KHÔNG tối ưu. '
        'LEXIO cần chiến lược DUAL-TRACK để maximize growth và revenue.'
    )
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, 'Tại sao cần cả B2B VÀ B2C?', level=2)
    
    why_both = [
        ('B2B (Teacher-Led) = VOLUME & SPEED', [
            'CAC thấp: 20k/student (thay vì 100k)',
            'Viral coefficient 10x: 1 teacher → 10-20 students',
            'Network effects: Teachers refer teachers',
            'Penetration nhanh vào rural/suburban markets',
            'Build proof và testimonials nhanh'
        ]),
        ('B2C (Trực tiếp) = LỢI NHUẬN CAO & THƯƠNG HIỆU', [
            'ARPU (Doanh thu/người dùng) cao hơn 10x: 99k vs 10k/học sinh qua B2B',
            'Xây dựng giá trị thương hiệu với người dùng cuối',
            'Sở hữu dữ liệu: Quan hệ trực tiếp với học sinh',
            'Cơ hội bán thêm: Gói Premium, dịch vụ bổ sung',
            'Thâm nhập thị trường đô thị (LTV cao hơn)'
        ]),
        ('SYNERGY (Hiệp lực) = 1 + 1 = 3', [
            'B2B tạo bằng chứng → B2C dễ chuyển đổi hơn',
            'Thương hiệu B2C → Bán hàng B2B dễ dàng hơn',
            'Học hỏi liên kênh: Cải thiện sản phẩm',
            'Quyền định giá: Có thể tính giá cao B2C với bằng chứng từ B2B',
            'Bao phủ thị trường: Nông thôn (B2B) + Thành thị (B2C)'
        ])
    ]
    
    for title, points in why_both:
        p = doc.add_paragraph()
        run = p.add_run(f'✓ {title}')
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(16, 185, 129)
        
        for point in points:
            doc.add_paragraph(f'  • {point}', style='List Bullet')
    
    doc.add_page_break()
    
    # ============= 1. TẠI SAO B2B-ONLY REVENUE THẤP? =============
    add_heading_custom(doc, '1. PHÂN TÍCH: TẠI SAO B2B-ONLY REVENUE THẤP?', level=1)
    
    doc.add_paragraph(
        'Trong phân tích trước, chiến lược chỉ B2B cho ARR Năm 3 = 6.66B, '
        'trong khi chỉ B2C cho ARR = 11.4B. Tại sao?'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '1.1. So sánh ARPU (Average Revenue Per User)', level=2)
    
    arpu_comparison = [
        ['Kênh', 'ARPU/tháng', 'ARPU/năm', 'Lý do'],
        ['B2B (Qua giáo viên)', '10,000đ', '120,000đ', 'Chia doanh thu với giáo viên, giảm giá số lượng'],
        ['B2C (Trực tiếp)', '99,000đ', '1,188,000đ', 'Giá đầy đủ, không qua trung gian'],
        ['Chênh lệch', '9.9x', '9.9x', 'B2C cao gấp gần 10 lần!']
    ]
    
    table = add_table_styled(doc, arpu_comparison[1:], arpu_comparison[0], col_widths=[1.8, 1.3, 1.3, 2.6])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('💡 INSIGHT QUAN TRỌNG: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    run = p.add_run(
        'Để đạt cùng doanh thu, B2B cần gấp 10 lần số lượng người dùng so với B2C. '
        'Ví dụ: 10,000 người dùng B2B = 1,000 người dùng B2C về doanh thu!'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '1.2. Phân tích chi tiết doanh thu', level=2)
    
    revenue_breakdown = [
        ['Chỉ số', 'B2B đơn thuần', 'B2C đơn thuần', 'Chiến lược kép'],
        ['Tổng người dùng Năm 3', '37,000', '5,000', '30,000 B2B + 7,000 B2C'],
        ['ARPU/tháng', '15,000đ', '150,000đ', 'Hỗn hợp: ~60k trung bình'],
        ['MRR (Doanh thu tháng)', '555M', '750M', '1,500M'],
        ['ARR (Doanh thu năm)', '6.66B', '9B', '18B'],
        ['Tiềm năng tăng trưởng', 'Bị giới hạn bởi mật độ giáo viên', 'Bị giới hạn bởi CAC', 'Tốt nhất của cả hai!']
    ]
    
    table = add_table_styled(doc, revenue_breakdown[1:], revenue_breakdown[0], col_widths=[2, 1.5, 1.5, 2])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('🎯 KẾT LUẬN: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    run = p.add_run(
        'Chiến lược Kép cho ARR 18B (Năm 3) - cao hơn 2.7x so với chỉ B2B và 2x so với chỉ B2C!'
    )
    
    doc.add_page_break()
    
    # ============= 2. DUAL-TRACK PRICING STRATEGY =============
    add_heading_custom(doc, '2. CHIẾN LƯỢC ĐỊNH GIÁ DUAL-TRACK', level=1)
    
    doc.add_paragraph(
        'Hai kênh có khách hàng mục tiêu, khả năng chi trả, và CAC (Chi phí thu hút khách) khác nhau → '
        'Cần chiến lược định giá riêng cho từng kênh.'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.1. Định giá B2B - Tập trung vào giáo viên & trung tâm', level=2)
    
    doc.add_paragraph('🎯 Mục tiêu: Tăng trưởng số lượng, giảm ma sát, hệ số lan truyền cao')
    doc.add_paragraph()
    
    b2b_pricing = [
        ['Gói', 'Giá/tháng', 'Giá/năm', 'Đối tượng', 'Giá trị'],
        ['GV Mới bắt đầu', '99,000đ\n(10 HS)', '890,000đ', 'GV vùng sâu/ngoại thành', 'Chương trình đầy đủ + AI Tutor\n~10k/HS'],
        ['GV Chuyên nghiệp', '299,000đ\n(50 HS)', '2,690,000đ', 'GV kinh nghiệm', '+ Hỗ trợ ưu tiên\n~6k/HS'],
        ['Trung tâm nhỏ', '999,000đ\n(200 HS)', '8,990,000đ', 'Trung tâm 3-10 GV', '+ Nhãn hiệu riêng\n~5k/HS'],
        ['Trung tâm lớn', '2,999,000đ\n(500 HS)', '26,990,000đ', 'Trung tâm 10+ GV', '+ Hỗ trợ chuyên biệt\n~6k/HS']
    ]
    
    table = add_table_styled(doc, b2b_pricing[1:], b2b_pricing[0], col_widths=[1.5, 1.3, 1.3, 1.8, 2])
    doc.add_paragraph()
    
    doc.add_paragraph('📊 Mô hình hoa hồng cho giáo viên:')
    commission = [
        'GV thu từ học sinh: 20-30k/HS/tháng',
        'GV trả LEXIO: 99k/tháng (10 HS)',
        'Lợi nhuận GV: 200-300k/tháng × 10 HS = 100-200k/tháng',
        '→ Động lực mạnh mẽ để tuyển và giữ học sinh!'
    ]
    for item in commission:
        doc.add_paragraph(f'  • {item}', style='List Bullet')
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.2. Định giá B2C - Trực tiếp đến người tiêu dùng', level=2)
    
    doc.add_paragraph('🎯 Mục tiêu: Tối đa hóa ARPU, xây dựng thương hiệu, sở hữu quan hệ khách hàng')
    doc.add_paragraph()
    
    b2c_pricing = [
        ['Gói', 'Giá/tháng', 'Giá/năm', 'Đối tượng', 'Định vị'],
        ['HS Theo tháng', '99,000đ', '-', 'Người dùng thử', 'Điểm bắt đầu, so sánh với ELSA'],
        ['HS Theo năm', '-', '890,000đ\n(~74k/tháng)', 'Người dùng cam kết', 'Tiết kiệm 25%'],
        ['Gia đình 2HS', '168,000đ\n(~84k/HS)', '1,512,000đ\n(~63k/HS)', 'Anh chị em', 'Giảm 15% so với lẻ'],
        ['Gia đình 4HS', '316,000đ\n(~79k/HS)', '2,844,000đ\n(~59k/HS)', 'Gia đình mở rộng', 'Giảm 20% so với lẻ'],
        ['Cao cấp', '249,000đ', '2,490,000đ', 'Phụ huynh cao cấp', '+ Buổi dạy trực tiếp']
    ]
    
    table = add_table_styled(doc, b2c_pricing[1:], b2c_pricing[0], col_widths=[1.5, 1.3, 1.5, 1.5, 2.2])
    doc.add_paragraph()
    
    doc.add_paragraph('🔑 Sự khác biệt so với B2B:')
    diff = [
        'Giá cao hơn 10x (99k vs 10k) - hợp lý vì giá trị trực tiếp',
        'Không qua trung gian - Người dùng sở hữu toàn bộ trải nghiệm',
        'Định vị cao cấp - "Chất lượng Cambridge tại nhà"',
        'Cơ hội bán thêm - Gói cao cấp, dịch vụ bổ sung',
        'Sở hữu dữ liệu - Vòng phản hồi trực tiếp'
    ]
    for item in diff:
        doc.add_paragraph(f'  • {item}', style='List Bullet')
    
    doc.add_page_break()
    
    # ============= 3. DUAL-TRACK GO-TO-MARKET =============
    add_heading_custom(doc, '3. CHIẾN LƯỢC GO-TO-MARKET SONG SONG', level=1)
    
    doc.add_paragraph(
        'Hai tracks chạy ĐỒNG THỜI nhưng focus khác nhau theo từng giai đoạn.'
    )
    
    doc.add_paragraph()
    
    # Phase 1
    add_heading_custom(doc, 'GIAI ĐOẠN 1: TẬP TRUNG B2B, XÂY NỀN B2C (Tháng 1-6)', level=2, color=RGBColor(16, 185, 129))
    
    doc.add_paragraph('🎯 Trọng tâm: 80% B2B, 20% B2C')
    doc.add_paragraph()
    
    phase1_b2b = [
        ('Kênh B2B - Trọng tâm chính', [
            'Mục tiêu: 50 giáo viên = 500 học sinh',
            'Chiến thuật 1: Tiếp cận cá nhân trong các nhóm giáo viên tiếng Anh',
            'Chiến thuật 2: Dùng thử miễn phí 1 tháng cho 30 GV đầu tiên',
            'Chiến thuật 3: Chương trình giới thiệu (giới thiệu GV → tặng tháng)',
            'Chỉ số thành công: 50 GV hoạt động, 70% giữ chân'
        ]),
        ('Kênh B2C - Xây dựng nền tảng', [
            'Mục tiêu: 100 học sinh trực tiếp',
            'Chiến thuật 1: Mạng lưới Founder (bạn bè & gia đình)',
            'Chiến thuật 2: Marketing nội dung (SEO, bài viết blog)',
            'Chiến thuật 3: Thử nghiệm quảng cáo FB nhỏ (ngân sách 5M)',
            'Chỉ số thành công: 100 người dùng trả phí, thu thập chứng thực'
        ])
    ]
    
    for title, tactics in phase1_b2b:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(16, 185, 129)
        for tactic in tactics:
            doc.add_paragraph(f'  • {tactic}', style='List Bullet')
        doc.add_paragraph()
    
    doc.add_paragraph('📊 Kết quả kỳ vọng Tháng 6:')
    results_m6 = [
        'B2B: 500 học sinh qua giáo viên',
        'B2C: 100 học sinh trực tiếp',
        'Tổng: 600 học sinh',
        'MRR (Doanh thu tháng): 5M (B2B) + 10M (B2C) = 15M',
        'Bằng chứng: 20+ video chứng thực',
        'Học hỏi: Xác thực sản phẩm phù hợp thị trường'
    ]
    for item in results_m6:
        doc.add_paragraph(f'  ✓ {item}', style='List Bullet')
    
    doc.add_page_break()
    
    # Phase 2
    add_heading_custom(doc, 'GIAI ĐOẠN 2: MỞ RỘNG HAI KÊNH (Tháng 7-18)', level=2, color=RGBColor(59, 130, 246))
    
    doc.add_paragraph('🎯 Trọng tâm: 60% B2B, 40% B2C')
    doc.add_paragraph()
    
    phase2_tracks = [
        ('Kênh B2B - Mở rộng qua trung tâm', [
            'Mục tiêu: 300 giáo viên + 20 trung tâm = 8,000 học sinh',
            'Chiến thuật 1: Thuê 2 nhân viên bán hàng B2B (hoa hồng)',
            'Chiến thuật 2: Hợp tác trung tâm (thí điểm → hợp đồng)',
            'Chiến thuật 3: Bùng nổ giới thiệu GV (phần thưởng cho GV giới thiệu nhiều nhất)',
            'Chiến thuật 4: Sự kiện địa phương ở các thành phố Tầng 2/3',
            'Thành công: 20 trung tâm ký, 300 GV hoạt động'
        ]),
        ('Kênh B2C - Ra mắt thị trường đô thị', [
            'Mục tiêu: 2,000 học sinh trực tiếp',
            'Chiến thuật 1: Mở rộng quảng cáo FB/Google (ngân sách 30M/tháng)',
            'Chiến thuật 2: Hợp tác influencer (10 blogger mẹ)',
            'Chiến thuật 3: Chiến dịch PR (VnExpress, Tuổi Trẻ đăng bài)',
            'Chiến thuật 4: Chương trình giới thiệu cho phụ huynh',
            'Thành công: 2,000 người dùng B2C trả phí, NPS >50'
        ])
    ]
    
    for title, tactics in phase2_tracks:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(59, 130, 246)
        for tactic in tactics:
            doc.add_paragraph(f'  • {tactic}', style='List Bullet')
        doc.add_paragraph()
    
    doc.add_paragraph('🔄 Hiệu ứng hiệp lực:')
    synergy_m18 = [
        'Bằng chứng B2B (8,000 HS) → Làm quảng cáo B2C chuyển đổi tốt hơn',
        'Chứng thực B2C → Bán hàng B2B dễ hơn ("Cả người thành phố cũng dùng")',
        'Cải thiện sản phẩm từ phản hồi B2C → Chất lượng B2B tốt hơn',
        'Nhận diện thương hiệu → Cả hai kênh đều hưởng lợi'
    ]
    for item in synergy_m18:
        doc.add_paragraph(f'  🔗 {item}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('📊 Kết quả kỳ vọng Tháng 18:')
    results_m18 = [
        'B2B: 8,000 học sinh',
        'B2C: 2,000 học sinh',
        'Tổng: 10,000 học sinh',
        'MRR: 80M (B2B) + 200M (B2C) = 280M',
        'ARR: 3.36B',
        'Phư sóng: 15+ bài đăng',
        'Vị thế thị trường: Top 3 ứng dụng EdTech tại Việt Nam'
    ]
    for item in results_m18:
        doc.add_paragraph(f'  ✓ {item}', style='List Bullet')
    
    doc.add_page_break()
    
    # Phase 3
    add_heading_custom(doc, 'GIAI ĐOẠN 3: ĐỊNH VỊ CAO CẤP (Tháng 19-36)', level=2, color=RGBColor(168, 85, 247))
    
    doc.add_paragraph('🎯 Trọng tâm: 50% B2B, 50% B2C (Cân bằng)')
    doc.add_paragraph()
    
    phase3_tracks = [
        ('Kênh B2B - Doanh nghiệp & Trường học', [
            'Mục tiêu: 1,000 giáo viên + 100 trung tâm = 30,000 học sinh',
            'Chiến thuật 1: Thí điểm tại trường công (trường công vùng sâu)',
            'Chiến thuật 2: Hợp tác chính phủ (phù hợp chính sách tiếng Anh)',
            'Chiến thuật 3: Hợp đồng trung tâm lớn (100+ HS)',
            'Chiến thuật 4: Chương trình chứng chỉ giáo viên',
            'Thành công: 100 trung tâm, 20 trường học, 1,000 giáo viên'
        ]),
        ('Kênh B2C - Ra mắt gói cao cấp', [
            'Mục tiêu: 7,000 học sinh trực tiếp',
            'Chiến thuật 1: Ra mắt gói Cao cấp (249k/tháng)',
            'Chiến thuật 2: Dịch vụ dạy 1-kêm-1 (499k/tháng)',
            'Chiến thuật 3: Hợp tác doanh nghiệp (phúc lợi nhân viên)',
            'Chiến thuật 4: Mở rộng quốc tế (gia đình người nước ngoài)',
            'Thành công: 7,000 người dùng B2C, 20% dùng gói Cao cấp'
        ])
    ]
    
    for title, tactics in phase3_tracks:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(168, 85, 247)
        for tactic in tactics:
            doc.add_paragraph(f'  • {tactic}', style='List Bullet')
        doc.add_paragraph()
    
    doc.add_paragraph('💎 Ở giai đoạn này:')
    stage3_status = [
        'Thương hiệu = "LEXIO = CLIL chuẩn Cambridge tại Việt Nam"',
        'Phủ sóng = Nông thôn (B2B) + Thành thị (B2C) hoàn chỉnh',
        'Bằng chứng = 30,000+ câu chuyện thành công',
        'Quyền định giá = Có thể tính giá cao cấp (199k+) với lý do đầy đủ',
        'Hiệu ứng mạng lưới = Cộng đồng giáo viên + phụ huynh mạnh mẽ'
    ]
    for item in stage3_status:
        doc.add_paragraph(f'  🌟 {item}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('📊 Kết quả kỳ vọng Năm 3:')
    results_y3 = [
        'B2B: 30,000 học sinh @ 15k ARPU = 450M MRR',
        'B2C: 7,000 học sinh @ 150k ARPU = 1,050M MRR',
        'Tổng: 37,000 học sinh',
        'MRR: 1,500M',
        'ARR: 18B',
        'Vị thế thị trường: Nền tảng CLIL #1 tại Việt Nam',
        'Quan tâm quốc tế: Mở rộng sang các thị trường ĐNA'
    ]
    for item in results_y3:
        doc.add_paragraph(f'  ✓ {item}', style='List Bullet')
    
    doc.add_page_break()
    
    # ============= 4. FINANCIAL PROJECTIONS =============
    add_heading_custom(doc, '4. DỰ PHÓNG TÀI CHÍNH CHI TIẾT - DUAL-TRACK', level=1)
    
    add_heading_custom(doc, '4.1. Phân tích doanh thu 3 năm', level=2)
    
    financial_3year = [
        ['Chỉ số', 'Năm 1 (Tháng 12)', 'Năm 2 (Tháng 24)', 'Năm 3 (Tháng 36)'],
        ['Học sinh B2B', '500', '8,000', '30,000'],
        ['B2B ARPU/tháng', '10,000đ', '12,000đ', '15,000đ'],
        ['B2B MRR', '5M', '96M', '450M'],
        ['Học sinh B2C', '100', '2,000', '7,000'],
        ['B2C ARPU/tháng', '99,000đ', '120,000đ', '150,000đ'],
        ['B2C MRR', '10M', '240M', '1,050M'],
        ['Tổng MRR', '15M', '336M', '1,500M'],
        ['Doanh thu năm', '180M', '4B', '18B'],
        ['B2B % doanh thu', '33%', '29%', '30%'],
        ['B2C % doanh thu', '67%', '71%', '70%']
    ]
    
    table = add_table_styled(doc, financial_3year[1:], financial_3year[0], col_widths=[2, 1.8, 1.8, 1.8])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('💡 KEY INSIGHT: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    run = p.add_run(
        'B2C đóng góp 70% doanh thu mặc dù chỉ chiếm 19% số lượng người dùng (7K/37K). '
        'B2B đóng góp 30% doanh thu nhưng chiếm 81% người dùng (30K/37K). '
        'Hai kênh bổ trợ hoàn hảo cho nhau!'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.2. So sánh kinh tế đơn vị', level=2)
    
    unit_econ = [
        ['Chỉ số', 'B2B', 'B2C', 'Hỗn hợp'],
        ['CAC (Chi phí thu hút)', '20,000đ/HS', '100,000đ/HS', '40,000đ TB'],
        ['ARPU/tháng', '15,000đ', '150,000đ', '60,000đ TB'],
        ['Biên lợi nhuận gộp', '80%', '75%', '77%'],
        ['Đóng góp tháng', '12,000đ', '112,500đ', '46,200đ'],
        ['Thời gian hoàn vốn', '1.7 tháng', '0.9 tháng', '1.2 tháng'],
        ['LTV (24 tháng)', '288,000đ', '2,700,000đ', '1,100,000đ'],
        ['Tỉ lệ LTV:CAC', '14.4:1', '27:1', '27.5:1'],
        ['Churn/tháng', '3%', '4%', '3.5%']
    ]
    
    table = add_table_styled(doc, unit_econ[1:], unit_econ[0], col_widths=[2, 1.6, 1.6, 1.6])
    doc.add_paragraph()
    
    doc.add_paragraph('📊 Giải thích:')
    interp = [
        'B2B: ARPU thấp hơn nhưng LTV:CAC vẫn xuất sắc (14.4:1) do CAC thấp',
        'B2C: ARPU cao hơn và LTV:CAC nổi bật (27:1)',
        'Hỗn hợp: LTV:CAC 27.5:1 là đẳng cấp thế giới (chuẩn SaaS >3:1)',
        'Thời gian hoàn vốn <2 tháng = Có thể tái đầu tư lợi nhuận nhanh'
    ]
    for item in interp:
        doc.add_paragraph(f'  • {item}', style='List Bullet')
    
    doc.add_page_break()
    
    # ============= 5. SYNERGY ANALYSIS =============
    add_heading_custom(doc, '5. PHÂN TÍCH HIỆP LỰC GIỮA B2B VÀ B2C', level=1)
    
    doc.add_paragraph(
        'Giá trị thực của Chiến lược Kép không chỉ là tổng hai kênh, '
        'mà là HIỆP LỰC tạo ra hiệu ứng nhân (1 + 1 = 3).'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '5.1. Hiệp lực từ B2B sang B2C', level=2)
    
    b2b_to_b2c = [
        ('Bằng chứng xã hội quy mô lớn', 
         'B2B tạo ra 30,000 người dùng nhanh → Quảng cáo B2C có bằng chứng mạnh: "Hơn 30,000 học sinh đang dùng". '
         'Tỉ lệ chuyển đổi B2C tăng 2-3x với bằng chứng xã hội.'),
        ('Phủ sóng địa lý',
         'B2B phủ nông thôn → Nhận biết thương hiệu toàn quốc → B2C không cần giáo dục thị trường từ đầu. '
         'Phụ huynh thành thị đã biết LEXIO từ giáo viên/bạn bè ở quê.'),
        ('Xác thực nội dung',
         'Giáo viên B2B cung cấp phản hồi về chương trình → Cải thiện nội dung → Người dùng B2C hưởng lợi. '
         'Sản phẩm phù hợp thị trường được xác thực với 30K người dùng trước khi mở rộng B2C.'),
        ('Giảm CAC',
         'B2B xây dựng thương hiệu → CAC của B2C giảm từ 100k xuống 70k do nhận biết thương hiệu. '
         'Luu lượng tự nhiên tăng, hiệu quả quảng cáo tốt hơn.'),
        ('Nguồn chứng thực',
         'B2B = Nguồn chứng thực không giới hạn. 30,000 học sinh = Hàng trăm câu chuyện thành công '
         'để dùng cho marketing B2C.')
    ]
    
    for title, explanation in b2b_to_b2c:
        p = doc.add_paragraph()
        run = p.add_run(f'🔗 {title}: ')
        run.font.bold = True
        run.font.color.rgb = RGBColor(16, 185, 129)
        run = p.add_run(explanation)
        run.font.size = Pt(10)
        doc.add_paragraph()
    
    add_heading_custom(doc, '5.2. Hiệp lực từ B2C sang B2B', level=2)
    
    b2c_to_b2b = [
        ('Định vị cao cấp',
         'Định giá B2C 99k/tháng → Giáo viên có thể thu 20-30k từ học sinh và vẫn được xem là "rẻ". '
         'Giá neo từ B2C giúp định giá B2B dễ chấp nhận.'),
        ('Phát triển sản phẩm',
         'Phản hồi trực tiếp B2C → Lặp lại sản phẩm nhanh hơn → Sản phẩm tốt hơn cho B2B. '
         'Tính năng như cải tiến AI Tutor được thử nghiệm với B2C trước.'),
        ('Uy tín thương hiệu',
         'Sự hiện diện B2C ở khu vực đô thị → Bán hàng B2B dễ hơn: "Ứng dụng này người thành phố cũng dùng" '
         '→ Giáo viên tin tưởng hơn.'),
        ('Cơ hội bán thêm',
         'Phụ huynh B2C hài lòng → Giới thiệu cho giáo viên → Tăng trưởng B2B tự nhiên. '
         'Một số phụ huynh B2C trở thành giáo viên (nhóm dạy tại nhà) → Bán chéo.'),
        ('Dữ liệu & Huấn luyện AI',
         'B2C = Dữ liệu phong phú (mẫu sử dụng, hiệu quả) → Huấn luyện AI tốt hơn → '
         'Giáo viên B2B hưởng lợi từ AI Tutor thông minh hơn.')
    ]
    
    for title, explanation in b2c_to_b2b:
        p = doc.add_paragraph()
        run = p.add_run(f'🔗 {title}: ')
        run.font.bold = True
        run.font.color.rgb = RGBColor(59, 130, 246)
        run = p.add_run(explanation)
        run.font.size = Pt(10)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= 6. OPERATIONAL REQUIREMENTS =============
    add_heading_custom(doc, '6. YÊU CẦU VẬN HÀNH - DUAL-TRACK', level=1)
    
    doc.add_paragraph(
        'Chạy 2 tracks đồng thời cần infrastructure, team, và processes khác nhau.'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '6.1. Cấu trúc đội ngũ (Năm 1)', level=2)
    
    team_y1 = [
        ['Vai trò', 'Số lượng', 'Trách nhiệm', 'Chi phí ước tính'],
        ['CEO/Người sáng lập', '1', 'Chiến lược tổng thể, gọi vốn, quan hệ đối tác', 'Cổ phần'],
        ['CTO', '1', 'Product development, tech infrastructure', '30M/năm'],
        ['B2B Sales', '2', 'Teacher outreach, center partnerships', '20M/năm + commission'],
        ['B2C Marketing', '1', 'Ads, content, influencer partnerships', '25M/năm'],
        ['Thành công khách hàng', '2', 'Hỗ trợ cả B2B và B2C', '15M/năm × 2'],
        ['Content', '1', 'Curriculum updates, blog, video', '20M/năm'],
        ['Total', '8', '-', '~150M/năm salary']
    ]
    
    table = add_table_styled(doc, team_y1[1:], team_y1[0], col_widths=[1.5, 1, 2.5, 1.8])
    doc.add_paragraph()
    
    add_heading_custom(doc, '6.2. Infrastructure Requirements', level=2)
    
    infra = [
        ('Teacher Dashboard', [
            'Features: Add/remove students, progress tracking, reports',
            'Timeline: 3 weeks development',
            'Cost: In-house (no additional cost)'
        ]),
        ('Payment System', [
            'B2B: Monthly invoicing, bulk payment',
            'B2C: QR banking, auto-activation',
            'Timeline: 2 weeks',
            'Cost: Free (use existing banking APIs)'
        ]),
        ('Analytics & Reporting', [
            'Theo dõi: Chỉ số B2B vs B2C riêng biệt',
            'Dashboard: Revenue, CAC, LTV, churn by channel',
            'Timeline: 1 week',
            'Cost: Use Mixpanel/Amplitude (5M/month)'
        ]),
        ('Support System', [
            'Zendesk or Intercom for ticketing',
            'Community forum for teachers',
            'Timeline: Immediate (SaaS tools)',
            'Cost: 10M/năm'
        ])
    ]
    
    for system, details in infra:
        p = doc.add_paragraph()
        run = p.add_run(f'⚙️ {system}')
        run.font.bold = True
        run.font.size = Pt(11)
        for detail in details:
            doc.add_paragraph(f'  • {detail}', style='List Bullet')
    
    doc.add_page_break()
    
    # ============= 7. RISK MITIGATION =============
    add_heading_custom(doc, '7. RỦI RO & GIẢM THIỂU', level=1)
    
    add_heading_custom(doc, '7.1. Rủi ro khi chạy Chiến lược Kép', level=2)
    
    risks = [
        ['Rủi ro', 'Tác động', 'Khả năng xảy ra', 'Biện pháp giảm thiểu'],
        ['Resource dilution\n(Chia sức cho 2 tracks)', 'High', 'Medium', 'Phase approach: 80% B2B Year 1 → 50/50 Year 3'],
        ['Pricing confusion\n(Customers không hiểu tại sao giá khác)', 'Medium', 'High', 'Clear positioning: B2B = via teachers, B2C = direct'],
        ['Channel conflict\n(Teachers cạnh tranh với B2C)', 'Medium', 'Low', 'Teacher pricing <30k → Below B2C 99k → No direct compete'],
        ['Tăng CAC\n(CAC B2C tăng khi mở rộng)', 'Cao', 'Trung bình', 'Đầu tư vào nội dung/SEO → Luu lượng tự nhiên → CAC thấp hơn'],
        ['Churn không khớp\n(B2B vs B2C churn khác)', 'Trung bình', 'Cao', 'Theo dõi riêng biệt, tối ưu từng kênh độc lập']
    ]
    
    table = add_table_styled(doc, risks[1:], risks[0], col_widths=[2, 1, 1, 3])
    doc.add_paragraph()
    
    add_heading_custom(doc, '7.2. Contingency Plans', level=2)
    
    contingency = [
        'Nếu tăng trưởng B2B chậm hơn dự kiến → Chuyển nguồn lực sang B2C (đã có hạ tầng)',
        'Nếu B2C CAC quá cao → Double down on B2B, delay B2C scale',
        'Nếu sản phẩm chưa sẵn sàng cho mở rộng B2B → Tập trung B2C trước để xác thực và cải thiện',
        'Nếu cả 2 channels đều struggle → Pivot sang B2B-only hoặc B2C-only dựa trên data'
    ]
    
    for plan in contingency:
        doc.add_paragraph(f'📋 {plan}', style='List Bullet')
    
    doc.add_page_break()
    
    # ============= 8. SUCCESS METRICS =============
    add_heading_custom(doc, '8. CHỈ SỐ THÀNH CÔNG - KPIs', level=1)
    
    add_heading_custom(doc, '8.1. North Star Metrics', level=2)
    
    north_star = [
        'Total Active Students (B2B + B2C)',
        'Blended MRR (Monthly Recurring Revenue)',
        'Blended LTV:CAC Ratio',
        'User Satisfaction: NPS >50'
    ]
    
    for metric in north_star:
        p = doc.add_paragraph(f'⭐ {metric}', style='List Bullet')
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '8.2. Channel-Specific KPIs', level=2)
    
    kpis_table = [
        ['Chỉ số', 'Mục tiêu B2B (Năm 1)', 'Mục tiêu B2C (Năm 1)', 'Cách đo'],
        ['New Users/month', '50 students', '10 students', 'Signup data'],
        ['CAC', '<30k/student', '<100k/student', 'Marketing spend / new users'],
        ['ARPU', '10k/month', '99k/month', 'MRR / active users'],
        ['Churn Rate', '<3%/month', '<4%/month', 'Churned users / total users'],
        ['LTV', '240k (24mo)', '2.4M (24mo)', 'ARPU × avg lifetime'],
        ['Activation Rate', '>80%', '>70%', '% users complete Week 1'],
        ['Referral Rate', '>20%', '>15%', '% users refer others'],
        ['NPS', '>60', '>50', 'Survey score']
    ]
    
    table = add_table_styled(doc, kpis_table[1:], kpis_table[0], col_widths=[2, 1.5, 1.5, 1.8])
    doc.add_paragraph()
    
    add_heading_custom(doc, '8.3. Lộ trình cột mốc', level=2)
    
    milestones = [
        ['Cột mốc', 'Ngày mục tiêu', 'Tiêu chí thành công'],
        ['MVP Launch', 'Month 1', '50 beta users, <5 critical bugs'],
        ['PMF Validation', 'Month 6', '600 users, NPS >50, 70% retention'],
        ['B2B Scale', 'Month 12', '500 B2B users, 20+ teacher testimonials'],
        ['B2C Launch', 'Month 6', '100 B2C users, <100k CAC'],
        ['Break-even', 'Month 18', 'Revenue > Costs'],
        ['Series A Ready', 'Month 24', '10K users, 4B ARR, profitable'],
        ['Market Leader', 'Month 36', '37K users, 18B ARR, #1 in CLIL']
    ]
    
    table = add_table_styled(doc, milestones[1:], milestones[0], col_widths=[2, 1.5, 3.5])
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= 9. CONCLUSION =============
    add_heading_custom(doc, '9. KẾT LUẬN & KHUYẾN NGHỊ', level=1)
    
    add_heading_custom(doc, '9.1. Tóm tắt Chiến lược', level=2)
    
    summary_final = [
        ('B2B-Only Strategy', '6.66B ARR Year 3', '❌ Revenue thấp, bỏ lỡ urban premium market'),
        ('B2C-Only Strategy', '9B ARR Year 3', '❌ CAC cao, slow growth, miss rural opportunity'),
        ('Chiến lược Kép', '18B ARR Năm 3', '✅ Tốt nhất của cả hai, doanh thu 2x, phủ sóng toàn thị trường')
    ]
    
    for strategy, arr, verdict in summary_final:
        p = doc.add_paragraph()
        run = p.add_run(f'{strategy}: ')
        run.font.bold = True
        run = p.add_run(f'{arr} ')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(79, 70, 229)
        run = p.add_run(f'→ {verdict}')
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '9.2. Lộ Trình Thực Thi', level=2)
    
    execution = [
        ('Tháng 1-6: B2B-Led Foundation', [
            'Tập trung 80% nguồn lực vào B2B',
            'Mục tiêu: 500 HS B2B, 100 HS B2C',
            'Build proof, testimonials, product-market fit',
            'Investment: 50M (team + infra + marketing)'
        ]),
        ('Tháng 7-18: Balanced Growth', [
            'Shift to 60/40 B2B/B2C',
            'Mục tiêu: 8K B2B, 2K B2C',
            'Scale both channels, hire team',
            'Investment: 200M (ads + sales + support)'
        ]),
        ('Tháng 19-36: Premium Expansion', [
            'Balanced 50/50 approach',
            'Mục tiêu: 30K B2B, 7K B2C',
            'Premium tiers, enterprise, international',
            'Investment: 500M (scale + expansion)'
        ])
    ]
    
    for phase, details in execution:
        p = doc.add_paragraph()
        run = p.add_run(f'📅 {phase}')
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(79, 70, 229)
        for detail in details:
            doc.add_paragraph(f'  • {detail}', style='List Bullet')
        doc.add_paragraph()
    
    add_heading_custom(doc, '9.3. Khuyến Nghị Cuối Cùng', level=2)
    
    final_recs = [
        'START IMMEDIATELY với B2B track - Quickest path to proof',
        'DON\'T ABANDON B2C - It\'s 70% of eventual revenue',
        'BUILD INFRASTRUCTURE for both channels từ đầu',
        'TRACK METRICS riêng biệt cho B2B và B2C',
        'BE FLEXIBLE - Adjust resource allocation based on data',
        'FOCUS ON SYNERGY - Hai channels phải hỗ trợ nhau',
        'THINK LONG-TERM - Year 3 ARR 18B requires both tracks'
    ]
    
    for rec in final_recs:
        p = doc.add_paragraph()
        run = p.add_run(f'✓ {rec}')
        run.font.bold = True
        run.font.color.rgb = RGBColor(16, 185, 129)
    
    doc.add_page_break()
    
    # ============= FINAL PAGE =============
    final_page = doc.add_paragraph()
    final_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = final_page.add_run('\n\n\n')
    run = final_page.add_run('DUAL-TRACK STRATEGY')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    run = final_page.add_run('\n\n')
    
    run = final_page.add_run('B2B + B2C = 18B ARR (Year 3)')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    
    run = final_page.add_run('\n\n\n')
    
    run = final_page.add_run(
        'Không phải B2B hay B2C\n'
        'Mà là B2B VÀ B2C\n\n'
        'Volume + Margin\n'
        'Rural + Urban\n'
        'Proof + Premium\n\n'
        '= Full Market Coverage\n'
        '= Maximum Revenue\n'
        '= Sustainable Growth'
    )
    run.font.size = Pt(12)
    
    run = final_page.add_run('\n\n\n')
    
    run = final_page.add_run('START NOW. BUILD BOTH. WIN EVERYTHING.')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    
    # Save document
    filename = f'LEXIO_Dual_Track_Strategy_Report_{datetime.datetime.now().strftime("%Y%m%d")}.docx'
    filepath = f'/Users/binhnguyen/Downloads/Engquest3k/{filename}'
    doc.save(filepath)
    
    print(f'✅ Báo cáo Chiến lược Kép đã được tạo thành công!')
    print(f'📄 File: {filepath}')
    print(f'📊 Số trang: ~35-40 trang')
    print(f'🎨 Định dạng: Chuyên nghiệp, phân tích toàn diện')
    print(f'💰 Kết quả quan trọng: 18B ARR (Năm 3) với chiến lược Kép')

if __name__ == '__main__':
    try:
        create_dual_track_report()
    except Exception as e:
        print(f'❌ Lỗi: {e}')
        import traceback
        traceback.print_exc()
