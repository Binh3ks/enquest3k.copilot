"""
Read Syllabus V5 and extract grammar/vocab targets per week.
Output: JSON mapping week -> {grammar, vocab_topics, reading_topics, exercise_types}
"""
import json, re, sys
from pathlib import Path

DOCX = Path("/Users/binhnguyen/Downloads/Engquest3k/Production_FINAL/Syllabus upgrade/Syllabus_V5_PublicationReady.docx")

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

doc = Document(str(DOCX))

# Extract all text
lines = []
for para in doc.paragraphs:
    t = para.text.strip()
    if t:
        lines.append(t)

# Also from tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            t = cell.text.strip()
            if t and t not in lines:
                lines.append(t)

# Save raw text for inspection
out = Path("/Users/binhnguyen/Downloads/Engquest3k/_syllabus_v5_raw.txt")
out.write_text("\n".join(lines), encoding="utf-8")
print(f"Extracted {len(lines)} lines -> {out}")
print("\n--- FIRST 100 LINES ---")
for l in lines[:100]:
    print(l)
