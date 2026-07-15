"""
create_w27_exam_v2.py
=====================
Generates the revised W27 Mock Exam document (two parts):
  PART A — Curriculum-Aligned Assessment (vocabulary & grammar W1-26 only)
  PART B — Flyers Gap Diagnostic (Flyers-level content; low score is expected & normal)

Output: Production_FINAL/1. FINAL MASS PRODUCTION/2_REFERENCE_DOCS/lesson_plans/output/
        W27_Mock_Exam_v2_Curriculum_Aligned.docx
        W27_Mock_Exam_v2_Answer_Key.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = (
    "Production_FINAL/1. FINAL MASS PRODUCTION/"
    "2_REFERENCE_DOCS/lesson_plans/output"
)
os.makedirs(OUT_DIR, exist_ok=True)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1, color="1F3864", size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def subheading(doc, text, color="2E74B5", size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def body(doc, text, size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def blank_line(doc, label="", size=11):
    """Print a numbered answer blank."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(label + "  _______________________________________________")
    run.font.size = Pt(size)
    return p


def divider(doc, char="─", length=60):
    p = doc.add_paragraph(char * length)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(150, 150, 150)


def section_box(doc, title, marks, bg="D6E4F0"):
    """Create a shaded section header with marks."""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    left_cell = t.cell(0, 0)
    right_cell = t.cell(0, 1)
    set_cell_bg(left_cell, bg)
    set_cell_bg(right_cell, bg)
    left_p = left_cell.paragraphs[0]
    left_run = left_p.add_run(title)
    left_run.bold = True
    left_run.font.size = Pt(11)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_p.add_run(f"[ ___ / {marks} marks ]")
    right_run.font.size = Pt(11)
    doc.add_paragraph()  # spacing
    return t


def mcq_item(doc, num, question, options):
    """Print a MCQ item with A/B/C options."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"{num}.  {question}")
    r.font.size = Pt(11)
    op_p = doc.add_paragraph()
    op_p.paragraph_format.left_indent = Cm(1.5)
    labels = ["A", "B", "C"]
    for i, opt in enumerate(options):
        op_p.add_run(f"  {labels[i]}) {opt}     ")
    op_p.runs[0].font.size = Pt(11)


def tf_item(doc, num, statement):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"{num}.  {statement}")
    r.font.size = Pt(11)
    answer_p = doc.add_paragraph()
    answer_p.paragraph_format.left_indent = Cm(1.5)
    answer_p.add_run("Answer:  ☐ TRUE     ☐ FALSE     ☐ DOESN'T SAY").font.size = Pt(10)


def gap_item(doc, num, before, after=""):
    """Part 4 style cloze item."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"{num}.  {before}  _____________  {after}")
    r.font.size = Pt(11)


def note_box(doc, text, bg="FFF2CC"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD EXAM DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build_exam():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── COVER ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("═" * 62)
    r.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WEEK 27  |  BLOCK B  |  ASSESSMENT DAY")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string("1F3864")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mock Exam 1 — Curriculum Edition  (v2.0)")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Integrated English Programme  |  Vocabulary & Grammar: W1–W26 Only")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("═" * 62)
    r.font.size = Pt(10)

    doc.add_paragraph()
    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    labels = ["Name:", "Class:", "Date:", "Teacher:"]
    for i, lbl in enumerate(labels):
        info.cell(i, 0).paragraphs[0].add_run(lbl).bold = True
        info.cell(i, 1).paragraphs[0].add_run("_" * 35)
    doc.add_paragraph()

    note_box(
        doc,
        "📌 TO THE STUDENT: This exam uses ONLY vocabulary and grammar from Weeks 1–26.\n"
        "Read each question carefully. Never leave an answer blank — always make your best guess.\n"
        "After the main exam you will complete a short 'Flyers Gap Tracker' (Part B, 15 min). "
        "Part B uses NEW words we haven't learned yet — a low score there is completely normal!",
        "D6E4F0",
    )

    # Timed exam overview table
    p = doc.add_paragraph()
    r = p.add_run("EXAM TIMETABLE  (90 minutes total)")
    r.bold = True; r.font.size = Pt(11)

    timed = doc.add_table(rows=8, cols=3)
    timed.style = "Table Grid"
    for i, h in enumerate(["Section", "Time", "Marks"]):
        c = timed.cell(0, i)
        set_cell_bg(c, "1F3864")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    time_rows = [
        ("Reading & Writing — Parts 1–6", "0–52 min",  "42 marks"),
        ("  Part 1  Word Definitions",     "  5 min",   "5"),
        ("  Part 2  MCQ Reading",          "  7 min",   "5"),
        ("  Part 3  TRUE/FALSE/DS",        "  7 min",   "7"),
        ("  Part 4  MCQ Cloze",            "  8 min",   "10"),
        ("  Part 5  Word Box",             "  6 min",   "7"),
        ("  Part 6  Letter",               "  7 min",   "8"),
    ]
    for i, (sec, tm, mk) in enumerate(time_rows, 1):
        timed.cell(i, 0).paragraphs[0].add_run(sec).font.size = Pt(9)
        timed.cell(i, 1).paragraphs[0].add_run(tm).font.size = Pt(9)
        timed.cell(i, 2).paragraphs[0].add_run(mk).font.size = Pt(9)

    doc.add_paragraph()
    timed2 = doc.add_table(rows=5, cols=3)
    timed2.style = "Table Grid"
    for i, h in enumerate(["Section", "Time", "Marks"]):
        c = timed2.cell(0, i)
        set_cell_bg(c, "1F3864")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    time_rows2 = [
        ("Part 7  Story Writing",         "52–64 min  (12 min)",  "20 marks"),
        ("Listening Parts 1–4",           "64–84 min  (20 min)",  "20 marks"),
        ("Part B  Flyers Gap Tracker",    "84–90 min  ( 6 min)",  "10 marks*"),
        ("TOTAL",                         "90 min",               "92 marks"),
    ]
    for i, (sec, tm, mk) in enumerate(time_rows2, 1):
        c0 = timed2.cell(i, 0); c1 = timed2.cell(i, 1); c2 = timed2.cell(i, 2)
        c0.paragraphs[0].add_run(sec).font.size = Pt(9)
        c1.paragraphs[0].add_run(tm).font.size = Pt(9)
        c2.paragraphs[0].add_run(mk).font.size = Pt(9)
        if sec == "TOTAL":
            set_cell_bg(c0, "D6E4F0"); set_cell_bg(c1, "D6E4F0"); set_cell_bg(c2, "D6E4F0")

    body(doc, "* Part B marks do not count toward your PART A band. They are diagnostic only.", size=9)

    # Score summary table
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SCORE SUMMARY")
    r.bold = True; r.font.size = Pt(11)

    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    headers = ["R&W Total", "Listening Total", "PART A Total", "PART B (Gap)", "Level"]
    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        set_cell_bg(c, "1F3864")
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(10)
    for i, v in enumerate(["___ / 62", "___ / 20", "___ / 82", "___ / 10", "___________"]):
        tbl.cell(1, i).paragraphs[0].add_run(v).font.size = Pt(10)

    doc.add_paragraph()
    bands = doc.add_table(rows=2, cols=4)
    bands.style = "Table Grid"
    band_headers = ["DISTINCTION", "MERIT", "PASS", "NEEDS PRACTICE"]
    band_ranges = ["70–82", "57–69", "41–56", "Below 41"]
    band_colors = ["1E8449", "2874A6", "E67E22", "CB4335"]
    for i, (bh, br, bc) in enumerate(zip(band_headers, band_ranges, band_colors)):
        c = bands.cell(0, i)
        set_cell_bg(c, bc)
        r = c.paragraphs[0].add_run(bh)
        r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(9)
        bands.cell(1, i).paragraphs[0].add_run(br).font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PART A — CURRICULUM-ALIGNED EXAM
    # ══════════════════════════════════════════════════════════════════════

    p = doc.add_paragraph()
    r = p.add_run("PART A  —  CURRICULUM-ALIGNED ASSESSMENT  (82 marks)")
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("1F3864")
    body(doc, "Vocabulary & Grammar covered in Weeks 1–26 of your programme.", size=10)
    divider(doc)
    doc.add_paragraph()

    # ── R&W Section Header ─────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("READING & WRITING  (62 marks  |  ~50 minutes)")
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("2E74B5")
    doc.add_paragraph()

    # ── PART 1 ─────────────────────────────────────────────────────────────
    section_box(doc, "PART 1 — Word Definitions  (5 marks)  ⏱ 5 min", 5)
    body(doc,
         "Read each definition. Choose the correct word from the box and write it on the line.\n"
         "There are 3 extra words you do not need.", size=11)
    doc.add_paragraph()
    # Word box
    word_box_p = doc.add_paragraph()
    word_box_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = word_box_p.add_run(
        "bedroom  |  kitchen  |  bathroom  |  park  |  library"
        "  |  supermarket  |  living room  |  school"
    )
    r.bold = True; r.font.size = Pt(11)
    doc.add_paragraph()

    defs = [
        "This is the room where you sleep at night.",
        "This is the place where you can borrow books for free.",
        "This is the room where you or your parents cook food.",
        "This is the place where you buy food and other things.",
        "This is the room where you wash your face and hands.",
    ]
    answers_p1 = ["bedroom", "library", "kitchen", "supermarket", "bathroom"]
    for i, d in enumerate(defs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{i}.  {d}").font.size = Pt(11)
        blank_line(doc, f"   Answer:")

    doc.add_paragraph()
    divider(doc)

    # ── PART 2 ─────────────────────────────────────────────────────────────
    section_box(doc, "PART 2 — Reading Comprehension: MCQ  (5 marks)  ⏱ 7 min", 5)
    body(doc, "Read the text. Choose the best answer — A, B or C.", size=11)
    doc.add_paragraph()

    part2_text = """\
A Day at the Farm

Last Saturday, Luna and Max visited their uncle's farm in the countryside. The farm was very quiet and clean. In the morning, they walked around the farm and watched the animals. Max was very excited. Luna was happy because she liked the countryside. It was not busy or noisy like the city.

In the afternoon, they played games in the park near the farm. They were a little tired after the games. In the evening, Uncle Tom cooked a delicious dinner for everyone. Max helped him in the kitchen. Luna cleaned the table after dinner. They walked home under the stars. It was a wonderful day!"""

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    r = p.add_run(part2_text)
    r.font.size = Pt(11)
    r.font.italic = True
    doc.add_paragraph()

    mcqs = [
        ("Where did Luna and Max go last Saturday?",
         ["the city", "the farm", "the library"]),
        ("What was the farm like?",
         ["noisy and busy", "quiet and clean", "old and new"]),
        ("How did Max feel at the farm in the morning?",
         ["bored", "tired", "excited"]),
        ("Who cooked dinner in the evening?",
         ["Luna", "Max", "Uncle Tom"]),
        ("What did Luna do after dinner?",
         ["played games", "cleaned the table", "walked to the park"]),
    ]
    for i, (q, opts) in enumerate(mcqs, 1):
        mcq_item(doc, i, q, opts)
        doc.add_paragraph()

    divider(doc)

    # ── PART 3 ─────────────────────────────────────────────────────────────
    section_box(doc, "PART 3 — TRUE / FALSE / DOESN'T SAY  (7 marks)  ⏱ 7 min", 7)
    body(doc,
         "Read the text. Write TRUE, FALSE or DOESN'T SAY next to each statement.",
         size=11)
    doc.add_paragraph()

    part3_text = """\
My Busy Saturday

My name is Sarah. Last Saturday was a very busy day! In the morning, I woke up early. I brushed my teeth and ate breakfast. After breakfast, I walked to the park with my brother. We played in the park for one hour. I was happy because the park was quiet and clean.

In the afternoon, I helped my mother cook lunch. We cooked rice and vegetables. My brother watched TV in the living room. I read my book on my bed. In the evening, I cleaned my bedroom and did my homework. I was a little tired, but I was also very happy!"""

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    r = p.add_run(part3_text)
    r.font.size = Pt(11)
    r.font.italic = True
    doc.add_paragraph()

    tf_statements = [
        "Sarah woke up early on Saturday morning.",
        "Sarah walked to the park with her sister.",
        "The park was noisy.",
        "Sarah and her mother cooked lunch together.",
        "Sarah's brother read a book after lunch.",
        "Sarah cleaned her bedroom in the evening.",
        "Sarah was happy at the end of the day.",
    ]
    tf_answers = ["TRUE", "FALSE", "FALSE", "TRUE", "FALSE", "TRUE", "TRUE"]
    for i, stmt in enumerate(tf_statements, 1):
        tf_item(doc, i, stmt)
        doc.add_paragraph()

    divider(doc)
    doc.add_page_break()

    # ── PART 4 ─────────────────────────────────────────────────────────────
    section_box(doc, "PART 4 — Story Gap-fill: Choose A, B or C  (10 marks)  ⏱ 8 min", 10)
    body(doc,
         "Read the story. Choose the best word — A, B or C — for each gap.",
         size=11)
    doc.add_paragraph()

    part4_intro = doc.add_paragraph()
    part4_intro.add_run("Luna's Art Project").bold = True
    part4_intro.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    p4_story_parts = [
        ("Last Thursday, the weather (1)___ very sunny and warm. Luna (2)___ to school with her brother. "
         "She (3)___ her art bag because it (4)___ Art class day. At school, Luna and her classmates "
         "(5)___ beautiful pictures. Luna (6)___ her picture of a tall building in the city. "
         "Her classmates (7)___ at her painting and (8)___. After class, the teacher (9)___ at all "
         "the pictures. Luna (10)___ her desk and then walked home.")
    ]
    for part in p4_story_parts:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)
        r = p.add_run(part)
        r.font.size = Pt(11)
        r.font.italic = True

    doc.add_paragraph()

    p4_items = [
        (1,  "A) is",        "B) was",       "C) were"),
        (2,  "A) walking",   "B) walk",       "C) walked"),
        (3,  "A) carries",   "B) carrying",   "C) carried"),
        (4,  "A) was",       "B) were",       "C) is"),
        (5,  "A) paint",     "B) painting",   "C) painted"),
        (6,  "A) colored",   "B) color",      "C) coloring"),
        (7,  "A) look",      "B) looking",    "C) looked"),
        (8,  "A) clapping",  "B) clapped",    "C) clap"),
        (9,  "A) looked",    "B) look",       "C) looking"),
        (10, "A) cleaned",   "B) clean",      "C) cleaning"),
    ]
    p4_answers = ["B", "C", "C", "A", "C", "A", "C", "B", "A", "A"]
    for num, a, b, c in p4_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{num:2}.    {a}          {b}          {c}").font.size = Pt(11)

    doc.add_paragraph()
    divider(doc)

    # ── PART 5 ─────────────────────────────────────────────────────────────
    section_box(doc, "PART 5 — Read and Choose from the Word Box  (7 marks)  ⏱ 6 min", 7)
    body(doc,
         "Choose the best word from the box to fill each gap. There are 3 extra words.",
         size=11)
    doc.add_paragraph()

    p5_box = doc.add_paragraph()
    p5_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p5_box.add_run(
        "because  |  First  |  Then  |  Next  |  Finally  |  and  |  When  |  but  |  so  |  after"
    )
    r.bold = True; r.font.size = Pt(11)
    doc.add_paragraph()

    p5_title = doc.add_paragraph()
    p5_title.add_run("How I Make My Breakfast").bold = True
    p5_title.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    p5_text = (
        "(1)___, I wake up and go to the kitchen. (2)___, I wash my hands with soap and water. "
        "(3)___, I take some bread and spread jam on it with a knife. My brother (4)___ I eat "
        "breakfast together every morning. I like jam (5)___ it is sweet. (6)___ I finish my "
        "breakfast, I brush my teeth. (7)___, I am ready for school!"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    r = p.add_run(p5_text)
    r.font.size = Pt(11)
    r.font.italic = True

    doc.add_paragraph()
    p5_answers = ["First", "Next / Then", "Then / Next", "and", "because", "When", "Finally"]
    p5_ans_labels = ["(1)", "(2)", "(3)", "(4)", "(5)", "(6)", "(7)"]
    for lbl in p5_ans_labels:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.add_run(f"{lbl}  ____________________").font.size = Pt(11)

    doc.add_paragraph()
    divider(doc)
    doc.add_page_break()

    # ── PART 6 ─────────────────────────────────────────────────────────────
    section_box(doc, "PART 6 — Complete the Letter  (8 marks)  ⏱ 7 min", 8)
    body(doc,
         "Complete the letter. Write ONE word for each gap. Use words from Weeks 1–26.",
         size=11)
    doc.add_paragraph()

    p6_lines = [
        ("Dear (1)_________,", False),
        ("", False),
        ("How are you? I am writing to tell you about my Saturday!", False),
        ("", False),
        ("In the morning, my family and I (2)_________ to the park. The weather (3)_________ sunny and warm.", False),
        ("My brother and I (4)_________ football together. I (5)_________ very happy.", False),
        ("After the park, we (6)_________ home and my mother (7)_________ a delicious lunch for us.", False),
        ("", False),
        ("In the evening, we (8)_________ a film on TV together.", False),
        ("", False),
        ("Please write back soon! I miss you!", False),
        ("Your friend,", False),
        ("Luna", False),
    ]
    for line_text, indent in p6_lines:
        p = doc.add_paragraph(line_text)
        p.paragraph_format.left_indent = Cm(1)
        r = p.runs[0] if p.runs else p.add_run(line_text)
        r.font.size = Pt(11)
        r.font.italic = True

    doc.add_paragraph()
    body(doc, "Write your answers here:", size=10)
    for i in range(1, 9):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.add_run(f"({i}) ____________________").font.size = Pt(11)

    doc.add_paragraph()
    divider(doc)

    # ── PART 7 ─────────────────────────────────────────────────────────────
    section_box(doc, "PART 7 — Story Writing from 4 Pictures  (20 marks)  ⏱ 12 min", 20)
    body(doc,
         "Look at the four pictures. Write a story about Luna and Max's Saturday adventure.\n"
         "Use past tense verbs and sequence words. Write at least 5 sentences.",
         size=11)
    doc.add_paragraph()

    picture_prompts = [
        "Picture 1: Luna and Max wake up and look outside — it is sunny!",
        "Picture 2: They walk to the park with their bags.",
        "Picture 3: They play football together in the park.",
        "Picture 4: They help clean the park (picking up rubbish) — they are very happy and tired!",
    ]
    for pp in picture_prompts:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run("📷  " + pp)
        r.font.size = Pt(11)
        r.bold = True

    doc.add_paragraph()
    body(doc, "Write your story here:", size=11)
    for i in range(7):
        p = doc.add_paragraph("_" * 72)
        p.paragraph_format.left_indent = Cm(0.5)
        p.runs[0].font.size = Pt(11)

    # Part 7 rubric
    doc.add_paragraph()
    body(doc, "TEACHER MARKING RUBRIC — Part 7", size=10)
    rub = doc.add_table(rows=4, cols=3)
    rub.style = "Table Grid"
    for row, (cat, max_m, criteria) in enumerate([
        ("Communication & Content", "8", "Story clear, all 4 pictures covered, beginning–middle–end"),
        ("Vocabulary", "6", "Uses W1–26 words correctly; adjectives/feelings included"),
        ("Grammar", "6", "Past simple regular (-ed) + was/were; sequence words used"),
    ]):
        rub.cell(row + 1, 0).paragraphs[0].add_run(cat).font.size = Pt(9)
        rub.cell(row + 1, 1).paragraphs[0].add_run(f"/ {max_m}").font.size = Pt(9)
        rub.cell(row + 1, 2).paragraphs[0].add_run(criteria).font.size = Pt(9)
    for h, txt in zip(rub.cell(0, 0).paragraphs, [""]):
        pass
    rub.cell(0, 0).paragraphs[0].add_run("Category").bold = True
    rub.cell(0, 1).paragraphs[0].add_run("Marks").bold = True
    rub.cell(0, 2).paragraphs[0].add_run("Criteria").bold = True
    set_cell_bg(rub.cell(0, 0), "1F3864"); rub.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    set_cell_bg(rub.cell(0, 1), "1F3864"); rub.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    set_cell_bg(rub.cell(0, 2), "1F3864"); rub.cell(0, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  LISTENING SECTION
    # ══════════════════════════════════════════════════════════════════════

    p = doc.add_paragraph()
    r = p.add_run("LISTENING  (20 marks)  ⏱ 20 min  — starts at minute 64")
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("2E74B5")
    doc.add_paragraph()
    note_box(doc,
             "📌 TEACHER: Read each listening script at approximately 100–105 words per minute. "
             "Read each part TWICE with a 10-second pause between readings.",
             "FFF2CC")

    # ── Listening Part 1 ───────────────────────────────────────────────────
    section_box(doc, "LISTENING PART 1 — 5 Dialogues: Choose A, B or C  (5 marks)  ⏱ ~5 min", 5)
    body(doc,
         "Listen to each short conversation. Circle the best picture or answer — A, B or C.",
         size=11)
    doc.add_paragraph()

    l1_items = [
        ("What is the weather like today?",
         ["A) It is raining.", "B) It is sunny.", "C) It is snowing."]),
        ("Which sport does Emma like?",
         ["A) football", "B) swimming", "C) running"]),
        ("How did Tom feel after school?",
         ["A) happy", "B) excited", "C) tired"]),
        ("Where is the book?",
         ["A) on the desk", "B) under the chair", "C) in the bag"]),
        ("How did Anna go to school today?",
         ["A) by bus", "B) by car", "C) she walked"]),
    ]
    l1_answers = ["B", "B", "C", "B", "C"]
    for i, (q, opts) in enumerate(l1_items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{i}.  {q}").font.size = Pt(11)
        op_p = doc.add_paragraph()
        op_p.paragraph_format.left_indent = Cm(1.5)
        op_p.add_run("    ".join(opts)).font.size = Pt(11)
        doc.add_paragraph()

    divider(doc)

    # ── Listening Part 2 ───────────────────────────────────────────────────
    section_box(doc, "LISTENING PART 2 — School Sports Day Information  (5 marks)  ⏱ ~5 min", 5)
    body(doc,
         "Listen and complete the information. Write ONE word or number for each gap.",
         size=11)
    doc.add_paragraph()

    body(doc, "SCHOOL SPORTS DAY", size=12)
    l2_items = [
        ("Day:", "Saturday"),
        ("Start time:", "9 o'clock"),
        ("Cost:", "free"),
        ("Students should wear:", "sports clothes"),
        ("Please bring:", "a water bottle"),
    ]
    l2_answers = [v for _, v in l2_items]
    for label, _ in l2_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.add_run(f"  {label:<25} _________________________").font.size = Pt(11)
    doc.add_paragraph()
    divider(doc)

    # ── Listening Part 3 ───────────────────────────────────────────────────
    section_box(doc, "LISTENING PART 3 — Match the Children to Activities  (5 marks)  ⏱ ~5 min", 5)
    body(doc,
         "Listen. Draw a line to match each child to what they did on Saturday.",
         size=11)
    doc.add_paragraph()

    l3_tbl = doc.add_table(rows=6, cols=3)
    l3_tbl.style = "Table Grid"
    l3_tbl.cell(0, 0).paragraphs[0].add_run("Child").bold = True
    l3_tbl.cell(0, 1).paragraphs[0].add_run("").bold = True
    l3_tbl.cell(0, 2).paragraphs[0].add_run("Activity").bold = True
    children = ["Anna", "Ben", "Carlos", "Diana", "Emi"]
    activities = ["played football", "cleaned her bedroom", "painted a picture",
                  "walked to the library", "watched TV with her family"]
    for i, (ch, act) in enumerate(zip(children, activities), 1):
        l3_tbl.cell(i, 0).paragraphs[0].add_run(ch).font.size = Pt(11)
        l3_tbl.cell(i, 1).paragraphs[0].add_run("→").font.size = Pt(11)
        l3_tbl.cell(i, 2).paragraphs[0].add_run(act).font.size = Pt(11)

    doc.add_paragraph()
    body(doc, "Write the correct activity next to each name:", size=10)
    l3_answers = dict(zip(children, activities))
    for ch in children:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.add_run(f"{ch}:  ___________________________________").font.size = Pt(11)

    doc.add_paragraph()
    divider(doc)

    # ── Listening Part 4 ───────────────────────────────────────────────────
    section_box(doc, "LISTENING PART 4 — Luna and Max's Weekend  (5 marks)  ⏱ ~5 min", 5)
    body(doc,
         "Listen to the conversation. Answer TRUE or FALSE for each statement.",
         size=11)
    doc.add_paragraph()

    l4_items = [
        "Max walked to the park on Sunday.",
        "Luna played football at the park.",
        "Max felt tired after his activity.",
        "Luna cleaned her room on Saturday.",
        "Max and Luna both had a great weekend.",
    ]
    l4_answers = ["TRUE", "FALSE", "TRUE", "TRUE", "TRUE"]
    for i, stmt in enumerate(l4_items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{i}.  {stmt}     →   ☐ TRUE     ☐ FALSE").font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  PART B — FLYERS GAP DIAGNOSTIC TRACKER
    # ══════════════════════════════════════════════════════════════════════

    p = doc.add_paragraph()
    r = p.add_run("PART B  —  FLYERS GAP DIAGNOSTIC TRACKER  (10 marks)  ⏱ 6 min")
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("1E8449")

    note_box(
        doc,
        "⭐ IMPORTANT NOTE FOR STUDENTS & PARENTS:\n"
        "This section uses vocabulary and grammar that we have NOT yet taught.\n"
        "These words come from the Cambridge Flyers exam — the exam we are working towards by Week 66–70.\n"
        "A LOW SCORE HERE IS COMPLETELY NORMAL AND EXPECTED AT WEEK 27.\n"
        "The purpose is to show us which areas we will focus on in Weeks 28–70.\n"
        "Part B marks do NOT count toward your exam band — they are diagnostic only.",
        "D5F5E3",
    )
    doc.add_paragraph()

    # Category 3 — Advanced Connectors (Flyers scope, Phase 2)
    subheading(doc, "SECTION B1 — Advanced Connectors  (5 marks)  — Coming: W55+", color="1E8449")
    body(doc,
         "These connectors are used in Cambridge Flyers Reading Part 5. "
         "Choose the best word for each gap.",
         size=10)
    doc.add_paragraph()

    b3_box_p = doc.add_paragraph()
    b3_box_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = b3_box_p.add_run("although  |  however  |  which  |  unless  |  while")
    r.bold = True; r.font.size = Pt(11)
    doc.add_paragraph()

    b3_items = [
        ("I like football. ___, my sister prefers swimming.", "However"),
        ("___ it was raining, we played outside.", "Although"),
        ("This is the book ___ I bought last week.", "which"),
        ("You cannot enter ___ you have a ticket.", "unless"),
        ("She was reading ___ her brother was watching TV.", "while"),
    ]
    for i, (sentence, _ans) in enumerate(b3_items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{i}.  {sentence}").font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        p2.add_run("Answer: ____________________").font.size = Pt(11)
        doc.add_paragraph()

    # Category 4 — Irregular Verbs Preview (W28-35 scope)
    subheading(doc, "SECTION B2 — Irregular Verb Preview  (5 marks)  — Coming: W28–35", color="1E8449")
    body(doc,
         "We will learn these verbs in Weeks 28–35! Try to guess the past form.",
         size=10)
    doc.add_paragraph()

    b4_items = [
        ("go",   "went",  "We ___ to the park. (go)"),
        ("eat",  "ate",   "She ___ an apple. (eat)"),
        ("see",  "saw",   "I ___ a bird in the tree. (see)"),
        ("run",  "ran",   "Max ___ very fast. (run)"),
        ("make", "made",  "We ___ a sandcastle. (make)"),
    ]
    for i, (base, past, sentence) in enumerate(b4_items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{i}.  {sentence}").font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        p2.add_run("Past form: ____________________").font.size = Pt(11)
        doc.add_paragraph()

    # Gap Tracker Summary
    doc.add_paragraph()
    heading(doc, "FLYERS GAP TRACKER — Summary Chart", level=2, color="1E8449", size=12)
    body(doc,
         "After marking, your teacher will colour-code this chart to show your current Flyers readiness.",
         size=10)
    doc.add_paragraph()

    tracker = doc.add_table(rows=4, cols=4)
    tracker.style = "Table Grid"
    tracker_headers = ["Topic Area", "Tested in Week", "Your Score", "Status"]
    for i, h in enumerate(tracker_headers):
        c = tracker.cell(0, i)
        set_cell_bg(c, "1E8449")
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(10)

    tracker_rows = [
        ("Advanced Connectors (B1)", "W55+",   "___ / 5",  "☐ Gap  ☐ OK"),
        ("Irregular Verbs (B2)",     "W28–35", "___ / 5",  "☐ Gap  ☐ OK"),
        ("PART A Total",             "W1–26",  "___ / 82", "☐ Gap  ☐ OK"),
    ]
    for i, row_data in enumerate(tracker_rows, 1):
        for j, val in enumerate(row_data):
            tracker.cell(i, j).paragraphs[0].add_run(val).font.size = Pt(10)

    doc.add_paragraph()
    note_box(
        doc,
        "📌 TEACHER: After marking Part B, use the Gap Tracker table above to plan focus for W28–W70.\n"
        "• B2 (Irregular Verbs) < 3/5 → prioritise W28–35 irregular verb work immediately\n"
        "• B1 (Connectors) < 3/5 → add although/however mini-lessons from W37 onwards",
        "FFF2CC",
    )

    # ══════════════════════════════════════════════════════════════════════
    #  LISTENING SCRIPTS (for teacher)
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()

    heading(doc, "TEACHER ONLY — LISTENING SCRIPTS", color="CB4335", size=12)
    note_box(doc,
             "📌 Read each script at ~100 wpm. Read TWICE. Pause 10 seconds between readings. "
             "Do not show this page to students.",
             "FADBD8")

    # L1 scripts
    subheading(doc, "LISTENING PART 1 Scripts", color="CB4335")
    l1_scripts = [
        ("Dialogue 1 (Weather)",
         "Boy: What is the weather like today?\nGirl: It is sunny! No clouds at all. Let's go to the park!\nBoy: Great idea!"),
        ("Dialogue 2 (Sport)",
         "Man: Emma, do you like football?\nEmma: No, I don't like football. I like swimming.\nMan: Do you swim every week?\nEmma: Yes, every Saturday morning."),
        ("Dialogue 3 (Feelings)",
         "Mum: Tom, how are you? How was school?\nTom: I am tired, Mum. We had PE today and I ran very fast.\nMum: Come and sit down. I'll make some tea."),
        ("Dialogue 4 (Location)",
         "Girl: Have you seen my book?\nBoy: Yes! It's under the chair. There, next to the table.\nGirl: Oh! Thank you!"),
        ("Dialogue 5 (Transport)",
         "Dad: Anna, did you take the bus today?\nAnna: No, Dad. I walked to school. It was sunny and not far.\nDad: Good girl!"),
    ]
    for title, script in l1_scripts:
        p = doc.add_paragraph()
        r = p.add_run(f"▶ {title}")
        r.bold = True; r.font.size = Pt(10)
        p2 = doc.add_paragraph(script)
        p2.paragraph_format.left_indent = Cm(1)
        p2.runs[0].font.size = Pt(10)
        p2.runs[0].font.italic = True
        doc.add_paragraph()

    # L2 script
    subheading(doc, "LISTENING PART 2 Script", color="CB4335")
    l2_script = """\
Teacher: Good morning, class! I have some exciting news. Our school Sports Day is this Saturday. It starts at 9 o'clock in the morning. Please tell your parents it is completely free — there is no cost at all. You must wear your sports clothes — not your school uniform. And please bring a water bottle because it might be hot. We hope to see all families there. It will be a wonderful day!"""
    p = doc.add_paragraph(l2_script)
    p.paragraph_format.left_indent = Cm(1)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.italic = True
    doc.add_paragraph()

    # L3 script
    subheading(doc, "LISTENING PART 3 Script", color="CB4335")
    l3_script = """\
Max: Hi! I'm asking everyone what they did last Saturday for our class project.
Anna: I played football with my friends in the park. It was really fun!
Ben: I stayed home. My room was very messy, so I cleaned it all morning.
Carlos: I painted a picture of my family. My mum put it on the wall!
Diana: I walked to the library and borrowed three books. I love reading!
Emi: I was a little tired, so I watched TV with my family in the evening."""
    p = doc.add_paragraph(l3_script)
    p.paragraph_format.left_indent = Cm(1)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.italic = True
    doc.add_paragraph()

    # L4 script
    subheading(doc, "LISTENING PART 4 Script", color="CB4335")
    l4_script = """\
Luna: Max! How was your weekend?
Max: It was great! On Sunday I walked to the park with my brother. We played there for two hours. I was so tired after!
Luna: Ha! I was tired too. On Saturday I cleaned my room — it took forever!
Max: Did you do anything fun?
Luna: Yes! After cleaning, I painted a picture in my room. It was really nice.
Max: So we both had a great weekend then!
Luna: Definitely! Best weekend ever."""
    p = doc.add_paragraph(l4_script)
    p.paragraph_format.left_indent = Cm(1)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.italic = True

    # Save
    out_path = os.path.join(OUT_DIR, "W27_Mock_Exam_v2_Curriculum_Aligned.docx")
    doc.save(out_path)
    print(f"✅ Exam saved: {out_path}")
    return out_path, {
        "p1": answers_p1, "p2": [v[0] for _, v in [(q, [o.split(") ")[1] if ") " in o else o for o in opts]) for q, opts in mcqs]],
        "p3_tf": tf_answers, "p4": p4_answers, "p5": p5_answers, "p6": ["[name]", "walked", "was", "played", "was", "walked", "cooked", "watched"],
        "l1": l1_answers, "l2": l2_answers, "l3": l3_answers, "l4": l4_answers,
        "b3": [a for _, a in b3_items], "b4": [past for _, past, _ in b4_items],
    }


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD ANSWER KEY DOCUMENT  (full rubric + gap analysis guide)
# ══════════════════════════════════════════════════════════════════════════════

def build_answer_key():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Cover ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("═" * 62); r.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WEEK 27 — MOCK EXAM v2")
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string("CB4335")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TEACHER ANSWER KEY  +  RUBRIC  +  GAP ANALYSIS GUIDE")
    r.bold = True; r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("═" * 62); r.font.size = Pt(10)

    note_box(doc,
        "🔒 CONFIDENTIAL — Do not distribute to students.\n"
        "This pack contains: (1) Full answer key with explanations  "
        "(2) Part 7 writing rubric with band descriptors  "
        "(3) Skill-by-skill error analysis  "
        "(4) Gap diagnostic guide for W28–W70 planning.",
        "FADBD8")
    doc.add_paragraph()

    # ── Scoring Summary ────────────────────────────────────────────────────
    heading(doc, "SECTION 1 — QUICK SCORING GUIDE", color="1F3864", size=12)

    score_tbl = doc.add_table(rows=10, cols=4)
    score_tbl.style = "Table Grid"
    for i, h in enumerate(["Part", "Topic / Skill", "Marks", "W1–26 Link"]):
        c = score_tbl.cell(0, i)
        set_cell_bg(c, "1F3864")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    rows_data = [
        ("P1",  "Word Definitions",              "5",  "W5, W11 (rooms & places)"),
        ("P2",  "MCQ Reading Comprehension",     "5",  "W10, W21 (farm/city, past -ed)"),
        ("P3",  "TRUE / FALSE / DOESN'T SAY",    "7",  "W21–24 (past simple, feelings)"),
        ("P4",  "MCQ Cloze (was/were + -ed)",    "10", "W19–23 (was/were, regular -ed)"),
        ("P5",  "Word Box — Connectors",         "7",  "W25 (First/Then/Finally/because)"),
        ("P6",  "Letter Gap-fill",               "8",  "W21–24 (past verbs in context)"),
        ("P7",  "Story Writing (4 pictures)",    "20", "W21–25 (past + connectors + adj)"),
        ("L1–4","Listening (4 parts)",           "20", "W10–17, W21–24"),
        ("",    "PART A TOTAL",                  "82", ""),
    ]
    for i, row in enumerate(rows_data, 1):
        for j, val in enumerate(row):
            c = score_tbl.cell(i, j)
            rr = c.paragraphs[0].add_run(val)
            rr.font.size = Pt(10)
            if val == "PART A TOTAL":
                rr.bold = True
                set_cell_bg(c, "D6E4F0")
                set_cell_bg(score_tbl.cell(i, 2), "D6E4F0")

    doc.add_paragraph()
    body(doc,
         "Band thresholds (PART A only):  DISTINCTION 70–82  |  MERIT 57–69  "
         "|  PASS 41–56  |  NEEDS PRACTICE < 41",
         size=10)
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 2 — FULL ANSWER KEY WITH EXPLANATIONS
    # ══════════════════════════════════════════════════════════════════════
    heading(doc, "SECTION 2 — FULL ANSWER KEY WITH EXPLANATIONS", color="1F3864", size=12)
    divider(doc)
    doc.add_paragraph()

    # P1
    subheading(doc, "PART 1 — Word Definitions  (5 marks, 1 each)", color="2E74B5")
    p1_data = [
        ("1", "bedroom",     "W5 — rooms vocabulary. Key: 'the room where you sleep'."),
        ("2", "library",     "W11 — places in the city. Key: 'borrow books for free'."),
        ("3", "kitchen",     "W5 — rooms vocabulary. Key: 'cook food'."),
        ("4", "supermarket", "W11 — places. Key: 'buy food and other things'."),
        ("5", "bathroom",    "W5 — rooms vocabulary. Key: 'wash your face and hands'."),
    ]
    ak_tbl = doc.add_table(rows=len(p1_data)+1, cols=3)
    ak_tbl.style = "Table Grid"
    for i, h in enumerate(["#", "Answer", "Teaching Note"]):
        c = ak_tbl.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(p1_data, 1):
        ak_tbl.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        rr = ak_tbl.cell(i, 1).paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        ak_tbl.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    # P2
    subheading(doc, "PART 2 — MCQ Reading  (5 marks, 1 each)", color="2E74B5")
    p2_data = [
        ("1", "B — the farm",         "Paragraph 1: 'visited their uncle's farm'. A/C not mentioned."),
        ("2", "B — quiet and clean",  "Paragraph 1: 'very quiet and clean'. A is city description (W9)."),
        ("3", "C — excited",          "Paragraph 1: 'Max was very excited'. A/B not stated."),
        ("4", "C — Uncle Tom",        "Paragraph 2: 'Uncle Tom cooked'. A/B wrong character."),
        ("5", "B — cleaned the table","Paragraph 2: 'Luna cleaned the table'. A is what they did at park."),
    ]
    ak2 = doc.add_table(rows=len(p2_data)+1, cols=3)
    ak2.style = "Table Grid"
    for i, h in enumerate(["#", "Answer", "Common Error / Note"]):
        c = ak2.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(p2_data, 1):
        ak2.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        rr = ak2.cell(i, 1).paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        ak2.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    # P3
    subheading(doc, "PART 3 — TRUE / FALSE / DOESN'T SAY  (7 marks, 1 each)", color="2E74B5")
    p3_data = [
        ("1", "TRUE",          "Paragraph 1: 'I woke up early.' — directly stated."),
        ("2", "FALSE",         "'walked to the park with my brother' — NOT sister. Common error: students skim."),
        ("3", "FALSE",         "Text says 'quiet and clean', not noisy. Direct contradiction."),
        ("4", "TRUE",          "'I helped my mother cook lunch.' — directly stated."),
        ("5", "FALSE",         "Brother 'watched TV', NOT read a book. Students often confuse the two actions."),
        ("6", "TRUE",          "'I cleaned my bedroom' — stated in paragraph 2."),
        ("7", "TRUE",          "'I was also very happy!' — last sentence."),
    ]
    note_box(doc,
        "⚠️ DOESN'T SAY vs FALSE: Remind students — FALSE means the text says the opposite. "
        "DOESN'T SAY means the text simply does not mention it. "
        "In this exam, all 7 are TRUE or FALSE (no DOESN'T SAY used) — "
        "a useful debrief point: 'Did you look for evidence in the text?'",
        "FFF2CC")
    ak3 = doc.add_table(rows=len(p3_data)+1, cols=3)
    ak3.style = "Table Grid"
    for i, h in enumerate(["#", "Answer", "Evidence in Text / Teaching Note"]):
        c = ak3.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(p3_data, 1):
        ak3.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        c = ak3.cell(i, 1)
        rr = c.paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        if ans == "FALSE": set_cell_bg(c, "FADBD8")
        ak3.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    # P4
    subheading(doc, "PART 4 — MCQ Cloze: Luna's Art Project  (10 marks, 1 each)", color="2E74B5")
    p4_data = [
        ("1",  "B — was",      "Past state: weather description → was (W19)"),
        ("2",  "C — walked",   "Past action, regular -ed (W21). 'walking' = present; 'walk' = present simple"),
        ("3",  "C — carried",  "Past -ed; 'carry→carried' (y→ied). Common error: A 'carries' is present."),
        ("4",  "A — was",      "Past state 'it was Art class day' (W19). 'were' needs plural subject."),
        ("5",  "C — painted",  "Past -ed (W23 art verbs). A 'paint' = present; B 'painting' = continuous."),
        ("6",  "A — colored",  "Past -ed art verb (W23). B 'color' = present; C 'coloring' = continuous."),
        ("7",  "C — looked",   "Past -ed (W21). A 'look' = present; B 'looking' = continuous."),
        ("8",  "B — clapped",  "Past -ed. A 'clapping' = continuous; C 'clap' = present."),
        ("9",  "A — looked",   "Past -ed. Context: teacher reviewed pictures after class."),
        ("10", "A — cleaned",  "Past -ed (W21 'cleaned'). B/C wrong tense."),
    ]
    note_box(doc,
        "📌 DIAGNOSTIC: If a student misses items 1, 4 (was/were) → revisit W19–20.\n"
        "If student misses 2, 3, 5–10 (regular -ed verbs) → revisit W21–23 verb conjugation.",
        "FFF2CC")
    ak4 = doc.add_table(rows=len(p4_data)+1, cols=3)
    ak4.style = "Table Grid"
    for i, h in enumerate(["#", "Answer", "Grammar Rule + W Reference"]):
        c = ak4.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(p4_data, 1):
        ak4.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        rr = ak4.cell(i, 1).paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        ak4.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    # P5
    subheading(doc, "PART 5 — Word Box Connectors  (7 marks, 1 each)", color="2E74B5")
    p5_data = [
        ("1", "First",         "Sequence opener (W25). Accept 'First' only."),
        ("2", "Next / Then",   "Either accepted. Both taught W25."),
        ("3", "Then / Next",   "Either accepted (swap with gap 2 is fine)."),
        ("4", "and",           "Coordinating conjunction — joining two subjects (W4)."),
        ("5", "because",       "Reason connector (W25 + W4). 'I like jam because...'"),
        ("6", "When",          "Time connector. 'When I finish...' = after I finish. Accept 'After'."),
        ("7", "Finally",       "Sequence closer (W25). Must be Finally here — sequence is complete."),
    ]
    ak5 = doc.add_table(rows=len(p5_data)+1, cols=3)
    ak5.style = "Table Grid"
    for i, h in enumerate(["#", "Answer", "Notes / Acceptable Alternatives"]):
        c = ak5.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(p5_data, 1):
        ak5.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        rr = ak5.cell(i, 1).paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        ak5.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    # P6
    subheading(doc, "PART 6 — Letter Gap-fill  (8 marks, 1 each)", color="2E74B5")
    p6_data = [
        ("1", "[any name]",  "Accept any plausible name. Penalise only if blank."),
        ("2", "walked",      "Past -ed (W21). 'We walked to the park.'"),
        ("3", "was",         "Past state: weather (W19). 'The weather was sunny.'"),
        ("4", "played",      "Past -ed (W21). 'played football together.'"),
        ("5", "was",         "Past state: feeling (W24). 'I was very happy.'"),
        ("6", "walked",      "Past -ed (W21). 'we walked home.' Accept 'went' — not yet taught but logical."),
        ("7", "cooked",      "Past -ed (W21/W25). 'my mother cooked a delicious lunch.'"),
        ("8", "watched",     "Past -ed (W21). 'we watched a film.' Accept 'saw' — not yet taught."),
    ]
    note_box(doc,
        "📌 MARKING NOTE: Accept any grammatically correct past tense verb that fits the context, "
        "even if not from the exact word list (e.g., 'went' for gap 6, 'saw' for gap 8). "
        "The skill being assessed is past tense formation, not vocabulary recall alone.",
        "FFF2CC")
    ak6 = doc.add_table(rows=len(p6_data)+1, cols=3)
    ak6.style = "Table Grid"
    for i, h in enumerate(["#", "Answer", "Notes / Acceptable Alternatives"]):
        c = ak6.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(p6_data, 1):
        ak6.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        rr = ak6.cell(i, 1).paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        ak6.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    doc.add_page_break()

    # ── Part 7 Rubric ──────────────────────────────────────────────────────
    subheading(doc, "PART 7 — STORY WRITING RUBRIC  (20 marks total)", color="CB4335", size=12)
    note_box(doc,
        "Picture prompts: (1) Luna & Max see sunny weather  (2) Walk to park  "
        "(3) Play football  (4) Clean up the park — happy & tired\n"
        "Expected story elements: past tense verbs (-ed, was/were), "
        "sequence words (First/Then/Finally), at least 1 feeling adjective.",
        "D6E4F0")
    doc.add_paragraph()

    # Communication rubric
    body(doc, "CATEGORY 1 — COMMUNICATION & CONTENT  (8 marks)", size=11)
    comm_tbl = doc.add_table(rows=5, cols=3)
    comm_tbl.style = "Table Grid"
    for i, h in enumerate(["Band", "Marks", "Descriptor"]):
        c = comm_tbl.cell(0, i); set_cell_bg(c, "CB4335")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    comm_bands = [
        ("4 — Excellent",    "7–8", "All 4 pictures covered. Clear beginning, middle, end. Story is easy to follow. Reader understands the full event."),
        ("3 — Good",         "5–6", "3–4 pictures covered. Story mostly clear. Minor gaps in logic but meaning is intact."),
        ("2 — Developing",   "3–4", "2 pictures covered or story is unclear/jumbled. Reader can partially understand."),
        ("1 — Beginning",    "1–2", "Only 1 picture addressed or story is very limited. Meaning hard to follow."),
    ]
    for i, (band, marks, desc) in enumerate(comm_bands, 1):
        comm_tbl.cell(i, 0).paragraphs[0].add_run(band).font.size = Pt(10)
        comm_tbl.cell(i, 1).paragraphs[0].add_run(marks).font.size = Pt(10)
        comm_tbl.cell(i, 2).paragraphs[0].add_run(desc).font.size = Pt(9)
    doc.add_paragraph()

    # Vocabulary rubric
    body(doc, "CATEGORY 2 — VOCABULARY  (6 marks)", size=11)
    vocab_tbl = doc.add_table(rows=4, cols=3)
    vocab_tbl.style = "Table Grid"
    for i, h in enumerate(["Band", "Marks", "Descriptor"]):
        c = vocab_tbl.cell(0, i); set_cell_bg(c, "1E8449")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    vocab_bands = [
        ("3 — Strong",      "5–6", "Uses 6+ words from W1–26 correctly. Includes adjectives (sunny, tired, happy, clean) and/or feeling words."),
        ("2 — Adequate",    "3–4", "Uses 3–5 W1–26 words correctly. Mostly nouns/verbs; limited adjectives."),
        ("1 — Limited",     "1–2", "Fewer than 3 W1–26 words used correctly, or words used incorrectly throughout."),
    ]
    for i, (band, marks, desc) in enumerate(vocab_bands, 1):
        vocab_tbl.cell(i, 0).paragraphs[0].add_run(band).font.size = Pt(10)
        vocab_tbl.cell(i, 1).paragraphs[0].add_run(marks).font.size = Pt(10)
        vocab_tbl.cell(i, 2).paragraphs[0].add_run(desc).font.size = Pt(9)
    doc.add_paragraph()

    # Grammar rubric
    body(doc, "CATEGORY 3 — GRAMMAR  (6 marks)", size=11)
    gram_tbl = doc.add_table(rows=4, cols=3)
    gram_tbl.style = "Table Grid"
    for i, h in enumerate(["Band", "Marks", "Descriptor"]):
        c = gram_tbl.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    gram_bands = [
        ("3 — Strong",   "5–6", "Consistent past simple (-ed verbs + was/were). Uses at least 2 sequence connectors (First / Then / Finally). Very few errors that affect meaning."),
        ("2 — Adequate", "3–4", "Uses past simple mostly correctly but inconsistent (mixes present/past). Uses at least 1 sequence connector."),
        ("1 — Limited",  "1–2", "Mostly present tense or no connectors. Errors are frequent and affect meaning."),
    ]
    for i, (band, marks, desc) in enumerate(gram_bands, 1):
        gram_tbl.cell(i, 0).paragraphs[0].add_run(band).font.size = Pt(10)
        gram_tbl.cell(i, 1).paragraphs[0].add_run(marks).font.size = Pt(10)
        gram_tbl.cell(i, 2).paragraphs[0].add_run(desc).font.size = Pt(9)
    doc.add_paragraph()

    note_box(doc,
        "📌 SAMPLE HIGH-SCORING RESPONSE (for teacher reference):\n"
        "\"First, Luna and Max woke up. It was sunny and warm. They were very happy!\n"
        "Then, they walked to the park together. They played football. Luna kicked the ball.\n"
        "Finally, they cleaned the park. They picked up rubbish. They were tired but happy.\"\n"
        "→ Communication: 8  |  Vocabulary: 6 (sunny, warm, happy, walked, park, kicked, rubbish, tired)  "
        "|  Grammar: 6 (was, walked, played, kicked, cleaned, picked, were — all past; First/Then/Finally used)",
        "D5F5E3")

    doc.add_page_break()

    # ── Listening Answer Keys ──────────────────────────────────────────────
    heading(doc, "SECTION 3 — LISTENING ANSWER KEY", color="1F3864", size=12)
    divider(doc)
    doc.add_paragraph()

    subheading(doc, "LISTENING PART 1  (5 marks, 1 each)", color="2E74B5")
    l1_data = [
        ("1", "B — sunny",           "Dialogue 1: 'It is sunny! No clouds at all.' A/C not mentioned."),
        ("2", "B — swimming",        "Dialogue 2: 'I like swimming.' Not football (A) or running (C)."),
        ("3", "C — tired",           "Dialogue 3: 'I am tired, Mum.' Not happy(A) or excited(B)."),
        ("4", "B — under the chair", "Dialogue 4: 'It's under the chair.' Not on desk(A) or in bag(C)."),
        ("5", "C — she walked",      "Dialogue 5: 'I walked to school.' Not bus(A) or car(B)."),
    ]
    ak_l1 = doc.add_table(rows=len(l1_data)+1, cols=3)
    ak_l1.style = "Table Grid"
    for i, h in enumerate(["#", "Answer", "Key phrase in script"]):
        c = ak_l1.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(l1_data, 1):
        ak_l1.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        rr = ak_l1.cell(i, 1).paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        ak_l1.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    subheading(doc, "LISTENING PART 2  (5 marks, 1 each)", color="2E74B5")
    l2_data = [
        ("Day",          "Saturday",     "'Sports Day is this Saturday.'"),
        ("Start time",   "9 o'clock",    "'starts at 9 o'clock in the morning.'"),
        ("Cost",         "free",         "'completely free — there is no cost at all.'"),
        ("Wear",         "sports clothes","'You must wear your sports clothes.'"),
        ("Bring",        "a water bottle","'please bring a water bottle.'"),
    ]
    ak_l2 = doc.add_table(rows=len(l2_data)+1, cols=3)
    ak_l2.style = "Table Grid"
    for i, h in enumerate(["Gap", "Answer", "Key phrase in script"]):
        c = ak_l2.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (gap, ans, note) in enumerate(l2_data, 1):
        ak_l2.cell(i, 0).paragraphs[0].add_run(gap).font.size = Pt(10)
        rr = ak_l2.cell(i, 1).paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        ak_l2.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    subheading(doc, "LISTENING PART 3  (5 marks, 1 each)", color="2E74B5")
    l3_data = [
        ("Anna",   "played football",           "'I played football with my friends.'"),
        ("Ben",    "cleaned his bedroom",        "'My room was very messy, so I cleaned it.'"),
        ("Carlos", "painted a picture",          "'I painted a picture of my family.'"),
        ("Diana",  "walked to the library",      "'I walked to the library and borrowed three books.'"),
        ("Emi",    "watched TV with her family", "'I watched TV with my family in the evening.'"),
    ]
    note_box(doc,
        "⚠️ NOTE: Ben = 'cleaned his bedroom' (not her). Emi = watched TV (not read/played). "
        "Common student error: confusing Ben (cleaned) with Diana (library).",
        "FFF2CC")
    ak_l3 = doc.add_table(rows=len(l3_data)+1, cols=3)
    ak_l3.style = "Table Grid"
    for i, h in enumerate(["Child", "Activity", "Key phrase"]):
        c = ak_l3.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (child, act, note) in enumerate(l3_data, 1):
        ak_l3.cell(i, 0).paragraphs[0].add_run(child).font.size = Pt(10)
        rr = ak_l3.cell(i, 1).paragraphs[0].add_run(act)
        rr.bold = True; rr.font.size = Pt(10)
        ak_l3.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    subheading(doc, "LISTENING PART 4  (5 marks, 1 each)", color="2E74B5")
    l4_data = [
        ("1", "TRUE",  "Max: 'On Sunday I walked to the park with my brother.'"),
        ("2", "FALSE", "Luna painted, she did NOT play football. Common trap."),
        ("3", "TRUE",  "Max: 'I was so tired after!'"),
        ("4", "TRUE",  "Luna: 'On Saturday I cleaned my room.'"),
        ("5", "TRUE",  "Max: 'So we both had a great weekend then!' / Luna: 'Definitely!'"),
    ]
    ak_l4 = doc.add_table(rows=len(l4_data)+1, cols=3)
    ak_l4.style = "Table Grid"
    for i, h in enumerate(["#", "Answer", "Evidence in script"]):
        c = ak_l4.cell(0, i); set_cell_bg(c, "2E74B5")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(l4_data, 1):
        ak_l4.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        c = ak_l4.cell(i, 1)
        rr = c.paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        if ans == "FALSE": set_cell_bg(c, "FADBD8")
        ak_l4.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    doc.add_page_break()

    # ── Part B ─────────────────────────────────────────────────────────────
    heading(doc, "SECTION 4 — PART B ANSWER KEY  (Diagnostic — not in band)", color="1E8449", size=12)
    divider(doc)
    doc.add_paragraph()

    subheading(doc, "B1 — Advanced Connectors  (5 marks)  [Taught: W55+]", color="1E8449")
    b1_data = [
        ("1", "However",  "Contrast. 'I like football. However, my sister prefers swimming.'"),
        ("2", "Although", "Concession. 'Although it was raining, we played outside.'"),
        ("3", "which",    "Relative clause. 'the book which I bought last week.'"),
        ("4", "unless",   "Condition (negative). 'cannot enter unless you have a ticket.'"),
        ("5", "while",    "Simultaneous. 'reading while her brother was watching TV.'"),
    ]
    ak_b1 = doc.add_table(rows=len(b1_data)+1, cols=3)
    ak_b1.style = "Table Grid"
    for i, h in enumerate(["#", "Answer", "Connector type + explanation"]):
        c = ak_b1.cell(0, i); set_cell_bg(c, "1E8449")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(b1_data, 1):
        ak_b1.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        rr = ak_b1.cell(i, 1).paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        ak_b1.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    subheading(doc, "B2 — Irregular Verb Preview  (5 marks)  [Taught: W28–35]", color="1E8449")
    b2_data = [
        ("1", "went",  "go → went  (W29: Movement verbs)"),
        ("2", "ate",   "eat → ate  (W30: Consumption verbs)"),
        ("3", "saw",   "see → saw  (W31: Perception verbs)"),
        ("4", "ran",   "run → ran  (W29: Movement verbs)"),
        ("5", "made",  "make → made  (W32: Task verbs)"),
    ]
    ak_b2 = doc.add_table(rows=len(b2_data)+1, cols=3)
    ak_b2.style = "Table Grid"
    for i, h in enumerate(["#", "Past Form", "Verb group + taught in"]):
        c = ak_b2.cell(0, i); set_cell_bg(c, "1E8449")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    for i, (num, ans, note) in enumerate(b2_data, 1):
        ak_b2.cell(i, 0).paragraphs[0].add_run(num).font.size = Pt(10)
        rr = ak_b2.cell(i, 1).paragraphs[0].add_run(ans)
        rr.bold = True; rr.font.size = Pt(10)
        ak_b2.cell(i, 2).paragraphs[0].add_run(note).font.size = Pt(9)
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 5 — SKILL-BY-SKILL GAP ANALYSIS
    # ══════════════════════════════════════════════════════════════════════
    heading(doc, "SECTION 5 — SKILL-BY-SKILL ERROR ANALYSIS  (per student)", color="CB4335", size=12)
    note_box(doc,
        "HOW TO USE: After marking, record each student's score per part below. "
        "Any part with score ≤ 50% = flag as a GAP. Use the Action column to plan W28+ content.",
        "D6E4F0")
    doc.add_paragraph()

    error_tbl = doc.add_table(rows=12, cols=5)
    error_tbl.style = "Table Grid"
    for i, h in enumerate(["Part", "Skill Tested", "Max", "Score ≤50% = GAP", "Recommended Action"]):
        c = error_tbl.cell(0, i); set_cell_bg(c, "CB4335")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(9)
    error_rows = [
        ("P1",  "Vocabulary: rooms & places",        "5",  "≤2/5",  "Re-teach W5/W11 word wall; flashcard games W28"),
        ("P2",  "Reading: scan for specific info",   "5",  "≤2/5",  "Explicit scan strategy: underline key words; W28 warm-up"),
        ("P3",  "Reading: TRUE/FALSE distinction",   "7",  "≤3/7",  "Re-teach DOESN'T SAY vs FALSE rule; W28 mini-drill"),
        ("P4",  "Grammar: was/were",                 "10", "≤4/10", "Revisit W19–20 was/were drills; daily starter W28–30"),
        ("P4",  "Grammar: regular -ed verbs",        "10", "≤4/10", "Revisit W21–23 -ed endings; pronunciation /t/ /d/ /ɪd/"),
        ("P5",  "Connectors: First/Then/Finally",    "7",  "≤3/7",  "Re-sequence activity W28; connector card game"),
        ("P6",  "Grammar: past verbs in context",    "8",  "≤4/8",  "Guided writing W28–29; letter frame scaffold"),
        ("P7",  "Writing: story + structure",        "20", "≤10/20","Story map scaffold; W28 journal entry daily writing"),
        ("L1",  "Listening: MCQ short dialogues",    "5",  "≤2/5",  "Dictation warm-ups W28+; listen-and-repeat exercises"),
        ("L2",  "Listening: note completion",        "5",  "≤2/5",  "Number & date dictation; W28 daily info-listening"),
        ("L3–4","Listening: matching + T/F",         "10", "≤5/10", "Partner interview tasks; listen-and-order activities"),
    ]
    for i, row in enumerate(error_rows, 1):
        for j, val in enumerate(row):
            error_tbl.cell(i, j).paragraphs[0].add_run(val).font.size = Pt(9)
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 6 — FLYERS GAP TRACKER + W28–70 PLANNING
    # ══════════════════════════════════════════════════════════════════════
    heading(doc, "SECTION 6 — FLYERS GAP TRACKER & W28–W70 TEACHING PLAN", color="1E8449", size=12)
    note_box(doc,
        "Target: Students sit Cambridge Flyers (15 shields) at Week 66–70.\n"
        "Current status (Week 27): ~200–250 words learned from Flyers wordlist (~500 total needed).\n"
        "Grammar gap: ~60% of Flyers grammar not yet taught (irregular verbs, comparatives, modals, "
        "present perfect, relative clauses).\n"
        "This section tracks progress toward that target and guides weekly decisions.",
        "D5F5E3")
    doc.add_paragraph()

    # Vocabulary gap
    subheading(doc, "VOCABULARY GAP TRACKER  (Flyers ~500 words needed)", color="1E8449")
    voc_tbl = doc.add_table(rows=8, cols=4)
    voc_tbl.style = "Table Grid"
    for i, h in enumerate(["Topic Area", "Flyers Words Needed", "Taught by W27", "Taught by..."]):
        c = voc_tbl.cell(0, i); set_cell_bg(c, "1E8449")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    voc_rows = [
        ("Rooms & House",          "~30 words", "✅ W5–6 (~20/30)",    "W37+"),
        ("Places in town",         "~25 words", "✅ W9–11 (~15/25)",   "W37+"),
        ("Animals & Nature",       "~60 words", "🟡 W10 (~10/60)",     "W37–42 CLIL"),
        ("Actions (irregular)",    "~40 verbs", "❌ 0/40",             "W28–35"),
        ("Feelings & Descriptions","~30 words", "✅ W24 (~15/30)",     "W37+"),
        ("School & Hobbies",       "~35 words", "🟡 W7–8 (~15/35)",   "W49–54"),
        ("Connectors & Grammar",   "~20 forms", "🟡 W25 (~5/20)",     "W55+"),
    ]
    for i, row in enumerate(voc_rows, 1):
        for j, val in enumerate(row):
            c = voc_tbl.cell(i, j)
            rr = c.paragraphs[0].add_run(val)
            rr.font.size = Pt(10)
            if "✅" in val: set_cell_bg(c, "D5F5E3")
            if "❌" in val: set_cell_bg(c, "FADBD8")
            if "🟡" in val: set_cell_bg(c, "FEF9E7")
    doc.add_paragraph()

    # Grammar gap
    subheading(doc, "GRAMMAR GAP TRACKER", color="1E8449")
    gram_gap = doc.add_table(rows=9, cols=4)
    gram_gap.style = "Table Grid"
    for i, h in enumerate(["Grammar Point", "Flyers Need?", "Status at W27", "Taught in"]):
        c = gram_gap.cell(0, i); set_cell_bg(c, "1E8449")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    gg_rows = [
        ("Past Simple — regular -ed",     "✅ Essential", "✅ W21–25",        "Complete"),
        ("Past Simple — irregular verbs", "✅ Essential", "❌ Not yet",       "W28–35"),
        ("Was / Were",                    "✅ Essential", "✅ W19–20",        "Complete"),
        ("Comparatives (-er than)",       "✅ Essential", "❌ Not yet",       "W40"),
        ("Modal verbs (can/can't)",       "✅ Essential", "✅ W12, W38",      "Partial"),
        ("Modal verbs (should/must)",     "✅ Essential", "❌ Not yet",       "Phase 2"),
        ("Present Perfect (have + -ed)",  "✅ Essential", "❌ Not yet",       "W52 (chunks)"),
        ("Relative clauses (which/who)",  "✅ Important", "❌ Not yet",       "W55+"),
    ]
    for i, row in enumerate(gg_rows, 1):
        for j, val in enumerate(row):
            c = gram_gap.cell(i, j)
            rr = c.paragraphs[0].add_run(val)
            rr.font.size = Pt(10)
            if "✅" in val and j == 2: set_cell_bg(c, "D5F5E3")
            if "❌" in val and j == 2: set_cell_bg(c, "FADBD8")
    doc.add_paragraph()

    # Mock exam schedule
    subheading(doc, "RECOMMENDED MOCK EXAM SCHEDULE  (W27 → W70)", color="1E8449")
    mock_tbl = doc.add_table(rows=5, cols=4)
    mock_tbl.style = "Table Grid"
    for i, h in enumerate(["Mock Exam", "Timing", "Scope", "Goal"]):
        c = mock_tbl.cell(0, i); set_cell_bg(c, "1F3864")
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255); rr.font.size = Pt(10)
    mock_rows = [
        ("Mock 1 (this exam)", "W27",   "W1–26 aligned + Gap Tracker",           "Baseline + Flyers format introduction"),
        ("Mock 2",             "W42",   "W1–42 (+ irregular verbs, CLIL vocab)", "Measure growth; adjust W43+ content"),
        ("Mock 3",             "W54",   "Phase 1 full scope + Flyers wordlist",  "Phase 1 exit check; set Phase 2 priorities"),
        ("Mock 4 (Flyers)",    "W66–70","Full Cambridge Flyers format",           "Target: 15 shields (Pass+)"),
    ]
    for i, row in enumerate(mock_rows, 1):
        for j, val in enumerate(row):
            mock_tbl.cell(i, j).paragraphs[0].add_run(val).font.size = Pt(10)
    doc.add_paragraph()

    note_box(doc,
        "📌 TEACHER ACTION PLAN AFTER W27 EXAM:\n"
        "1. Students scoring DISTINCTION (70–82): Begin 'Flyers Preview' vocabulary list from W28 onwards\n"
        "2. Students scoring MERIT (57–69): Continue W28 programme as planned; add irregular verb focus\n"
        "3. Students scoring PASS (41–56): Add was/were + -ed consolidation warm-up every session W28–30\n"
        "4. Students scoring NEEDS PRACTICE (<41): 1:1 review of P1+P4 gaps; consider W26 revision tasks\n"
        "5. ALL STUDENTS: B2 score < 3/5 → make irregular verb exposure a priority from W28",
        "FFF2CC")

    out_path = os.path.join(OUT_DIR, "W27_Mock_Exam_v2_Answer_Key.docx")
    doc.save(out_path)
    print(f"✅ Answer key saved: {out_path}")
    return out_path


if __name__ == "__main__":
    exam_path, answers = build_exam()
    ak_path = build_answer_key()
    print("\nDone! Files generated:")
    print(f"  Exam:       {exam_path}")
    print(f"  Answer Key: {ak_path}")
