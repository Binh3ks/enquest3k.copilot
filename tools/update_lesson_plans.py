#!/usr/bin/env python3
"""
update_lesson_plans.py — Lesson Plan Extractor for EngQuest Mass Production
============================================================================

USAGE:
  # Auto-scan default folder
  python3 tools/update_lesson_plans.py

  # Specify a custom docx folder
  python3 tools/update_lesson_plans.py --docx-dir "/Volumes/MY DOCUMENT/CT TA Ứng dụng_3yrs/FINAL Syllabus/"

  # Check a specific week number
  python3 tools/update_lesson_plans.py --check-week 28

  # Dry run — show what would be added, no file changes
  python3 tools/update_lesson_plans.py --dry-run

WORKFLOW (called by agent at BƯỚC 0.5):
  1. Scans docx folder for files matching "Lesson_plans_W*.docx" or similar
  2. Parses only weeks NOT already in lessonPlans.json   ← incremental, safe
  3. Merges new weeks into existing JSON (old weeks NOT overwritten)
  4. Prints "Added weeks: [28, 54, 55]"  OR  "No new weeks found"
  5. Prints instruction to update LP_AVAILABLE in TeacherPanel.jsx

OUTPUT: public/data/lessonPlans.json  (merged, not replaced)
"""
from docx import Document
import json, re, sys, argparse, glob, os

# ── Paths ───────────────────────────────────────────────────────────────────
WORKSPACE   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH    = os.path.join(WORKSPACE, "public", "data", "lessonPlans.json")
TEACHER_JSX = os.path.join(WORKSPACE, "src", "components", "teacher", "TeacherPanel.jsx")
DEFAULT_DOCX_DIR = "/Volumes/MY DOCUMENT/CT TA Ứng dụng_3yrs/FINAL Syllabus/"

# Known docx files (filename → week range hint); used when auto-scanning
KNOWN_FILE_PATTERNS = [
    r"Lesson.plans.W(\d+)[_-](\d+)\.docx",   # "Lesson_plans_W25-36.docx"
    r"W(\d+)[_-](\d+).*\.docx",               # "W37-53_lessons.docx"
]

# ── Table helper ────────────────────────────────────────────────────────────
def table_to_rows(tbl):
    rows = []
    for row in tbl.rows:
        rows.append([c.text.strip() for c in row.cells])
    return rows


# ── Main parser ─────────────────────────────────────────────────────────────
def parse_doc(path):
    """Parse a .docx file and return dict of {week_num_str: week_data}."""
    doc = Document(path)
    all_paras = [(p.style.name, p.text.strip()) for p in doc.paragraphs]
    tables    = doc.tables

    # Locate each week block
    week_blocks = []
    for i, (style, text) in enumerate(all_paras):
        m = re.match(r"TEACHER CONTENT PACK.*?—\s*WEEK\s+(\d+)", text)
        if m:
            week_blocks.append([int(m.group(1)), i, None])

    if not week_blocks:
        print(f"  ⚠️  No week blocks found in {os.path.basename(path)}")
        return {}

    # Assign end index (stop before APPENDIX)
    for bi in range(len(week_blocks)):
        next_start = week_blocks[bi+1][1] if bi+1 < len(week_blocks) else len(all_paras)
        app_cut = next_start
        for j in range(week_blocks[bi][1], next_start):
            if all_paras[j][1].startswith("APPENDIX"):
                app_cut = j
                break
        week_blocks[bi][2] = app_cut

    # Map body element order (para vs table)
    body_elements = []
    p_idx, t_idx = 0, 0
    for child in doc.element.body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            if p_idx < len(all_paras):
                body_elements.append(('para', p_idx))
            p_idx += 1
        elif tag == 'tbl':
            body_elements.append(('table', t_idx))
            t_idx += 1

    result = {}

    for week_num, wstart, wend in week_blocks:
        block_paras = all_paras[wstart:wend]

        def idx_of(pattern, start=0):
            for i in range(start, len(block_paras)):
                if re.search(pattern, block_paras[i][1], re.IGNORECASE):
                    return i
            return None

        def section_text(from_idx, to_idx):
            if from_idx is None:
                return []
            end = to_idx if to_idx is not None else len(block_paras)
            return [t for _, t in block_paras[from_idx+1:end] if t]

        # Header / unit theme
        header = block_paras[0][1] if block_paras else ""
        unit_theme = ""
        for _, t in block_paras[1:6]:
            t = t.strip()
            if t and not t.startswith("═") and not t.startswith("Integrated") \
               and not t.startswith("Generat") and not t.startswith("SECTION"):
                unit_theme = t
                break

        bm = re.search(r"BLOCK ([A-Z])", header)
        block_letter = bm.group(1) if bm else "?"

        # Section indices
        s1  = idx_of(r"SECTION 1:")
        s2  = idx_of(r"SECTION 2:")
        s3  = idx_of(r"SECTION 3:")
        s4  = idx_of(r"SECTION 4:")
        s10 = idx_of(r"SECTION 10:")
        s11 = idx_of(r"SECTION 11:")
        s12 = idx_of(r"SECTION 12:")

        # Quick Reference (Section 1 table)
        quick_ref = {}
        if s1 is not None:
            abs_s1 = wstart + s1
            for be_idx, (be_type, be_i) in enumerate(body_elements):
                if be_type == 'table':
                    prev_para = next((body_elements[j][1] for j in range(be_idx-1, -1, -1) if body_elements[j][0]=='para'), 0)
                    next_para = next((body_elements[j][1] for j in range(be_idx+1, len(body_elements)) if body_elements[j][0]=='para'), len(all_paras))
                    if prev_para >= abs_s1 and next_para <= wstart + (s2 or s3 or s4 or len(block_paras)):
                        rows = table_to_rows(tables[be_i])
                        if rows:
                            for row in rows[1:]:
                                if len(row) >= 2 and row[0]:
                                    quick_ref[row[0]] = row[1]
                        break

        # Teaching Methodology (Section 2)
        methodology = []
        if s2 is not None:
            s2_end = s3 if s3 is not None else (s4 or len(block_paras))
            raw = section_text(s2, s2_end)
            cur = None
            for line in raw:
                if re.match(r"2\.\d+\s+\w", line):
                    cur = {"title": line, "content": []}
                    methodology.append(cur)
                elif cur:
                    cur["content"].append(line)
                else:
                    methodology.append({"title": line, "content": []})

        # Vocab tiers (Section 3 tables)
        vocab_tiers = []
        if s3 is not None:
            s3_end = s4 if s4 is not None else (s10 or len(block_paras))
            abs_s3 = wstart + s3
            abs_s3_end = wstart + s3_end
            for be_idx, (be_type, be_i) in enumerate(body_elements):
                if be_type == 'table':
                    prev_para = next((body_elements[j][1] for j in range(be_idx-1, -1, -1) if body_elements[j][0]=='para'), 0)
                    next_para = next((body_elements[j][1] for j in range(be_idx+1, len(body_elements)) if body_elements[j][0]=='para'), len(all_paras))
                    if abs_s3 <= prev_para < abs_s3_end:
                        rows = table_to_rows(tables[be_i])
                        if rows and len(rows) > 1:
                            headers = rows[0]
                            for row in rows[1:]:
                                entry = {headers[j]: row[j] for j in range(min(len(headers), len(row)))}
                                if entry.get("Word") or entry.get("Tier"):
                                    vocab_tiers.append(entry)

        # Sessions 1, 2, 3
        session_markers = []
        for i, (_, t) in enumerate(block_paras):
            m = re.match(r"WEEK\s+\d+\s*\|\s*SESSION\s+(\d+)", t)
            if m:
                session_markers.append((i, int(m.group(1))))

        sessions = []
        for si, (sm_idx, sess_num) in enumerate(session_markers):
            sess_end = session_markers[si+1][0] if si+1 < len(session_markers) else (s10 or len(block_paras))
            sess_lines = [t for _, t in block_paras[sm_idx+1:sess_end] if t]
            parts = []
            cur_part = None
            for line in sess_lines:
                if re.match(r"PART\s+\d+:", line) or re.match(r"\[(O|X|\?)\]\s+[A-Z]", line):
                    cur_part = {"title": line, "content": []}
                    parts.append(cur_part)
                elif cur_part:
                    cur_part["content"].append(line)
                else:
                    parts.append({"title": line, "content": []})
            sessions.append({"session": sess_num, "parts": parts, "raw_lines": sess_lines})

        # Answer Key, Rubrics, Task Cards
        answer_key = section_text(s10, s11) if s10 is not None else []
        rubrics    = section_text(s11, s12) if s11 is not None else []
        task_cards = section_text(s12, None) if s12 is not None else []

        result[str(week_num)] = {
            "week":        week_num,
            "block":       block_letter,
            "header":      header,
            "unit_theme":  unit_theme,
            "quick_ref":   quick_ref,
            "methodology": methodology,
            "vocab_tiers": vocab_tiers,
            "sessions":    sessions,
            "answer_key":  answer_key,
            "rubrics":     rubrics,
            "task_cards":  task_cards,
        }

    return result


# ── Auto-discover docx files in a folder ────────────────────────────────────
def discover_docx_files(docx_dir):
    """Return list of .docx file paths found in docx_dir."""
    if not os.path.isdir(docx_dir):
        print(f"  ⚠️  Folder not found: {docx_dir}")
        return []
    paths = glob.glob(os.path.join(docx_dir, "*.docx"))
    # Filter to plausible lesson plan files
    lesson_files = [
        p for p in paths
        if re.search(r"(lesson.plan|W\d+)", os.path.basename(p), re.IGNORECASE)
    ]
    if not lesson_files and paths:
        # Fall back: any .docx in folder
        lesson_files = paths
    return sorted(lesson_files)


# ── Update LP_AVAILABLE instructions ────────────────────────────────────────
def print_lp_available_instructions(new_weeks, all_weeks):
    """Tell the agent what to change in TeacherPanel.jsx."""
    sorted_weeks = sorted(int(w) for w in all_weeks.keys())
    formatted = ", ".join(str(w) for w in sorted_weeks)
    print("\n" + "="*60)
    print("📋 ACTION REQUIRED: Update LP_AVAILABLE in TeacherPanel.jsx")
    print("="*60)
    print(f"\nFile: src/components/teacher/TeacherPanel.jsx")
    print(f"Find the line:  const LP_AVAILABLE = new Set([")
    print(f"Replace its contents with all {len(sorted_weeks)} weeks:")
    print(f"\n  const LP_AVAILABLE = new Set([")

    # Format in rows of ~12 for readability
    rows = [sorted_weeks[i:i+12] for i in range(0, len(sorted_weeks), 12)]
    for row in rows:
        print("    " + ",".join(str(w) for w in row) + ",")
    print(f"  ]);\n")
    print(f"Newly added weeks: {sorted(int(w) for w in new_weeks)}")
    print("="*60 + "\n")


# ── main ─────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="Update public/data/lessonPlans.json from .docx lesson plan files."
    )
    parser.add_argument("--docx-dir", default=DEFAULT_DOCX_DIR,
                        help=f"Folder containing .docx lesson plan files (default: {DEFAULT_DOCX_DIR})")
    parser.add_argument("--check-week", type=int, metavar="N",
                        help="Just check if week N is already in lessonPlans.json and exit")
    parser.add_argument("--dry-run", action="store_true",
                        help="Parse and report what would be added, but do not write files")
    parser.add_argument("--force", action="store_true",
                        help="Re-parse and overwrite even weeks already present in lessonPlans.json")
    args = parser.parse_args()

    # ── Check-week mode ──────────────────────────────────────────────────
    if args.check_week:
        if os.path.exists(OUT_PATH):
            with open(OUT_PATH, encoding="utf-8") as f:
                existing = json.load(f)
            week_key = str(args.check_week)
            if week_key in existing:
                theme = existing[week_key].get("unit_theme", "")
                print(f"✅ W{args.check_week} already in lessonPlans.json — '{theme}'")
            else:
                print(f"⚠️  W{args.check_week} NOT found in lessonPlans.json")
                print(f"   → Run  python3 tools/update_lesson_plans.py  to add it (if docx exists)")
        else:
            print(f"⚠️  lessonPlans.json not found at {OUT_PATH}")
        return

    # ── Load existing data ───────────────────────────────────────────────
    if os.path.exists(OUT_PATH):
        with open(OUT_PATH, encoding="utf-8") as f:
            existing = json.load(f)
        print(f"📂 Loaded existing lessonPlans.json — {len(existing)} weeks")
    else:
        existing = {}
        print(f"📂 No existing lessonPlans.json — will create fresh")

    # ── Discover docx files ──────────────────────────────────────────────
    docx_files = discover_docx_files(args.docx_dir)
    if not docx_files:
        print(f"\n⚠️  No .docx lesson plan files found in: {args.docx_dir}")
        print("   Check that the folder is mounted / accessible.")
        print("   Weeks not covered by lessonPlans.json will show 🔒 in Teacher Panel (acceptable).")
        return

    print(f"\n📄 Found {len(docx_files)} docx file(s):")
    for p in docx_files:
        print(f"   {os.path.basename(p)}")

    # ── Parse docs (skip weeks already present unless --force) ──────────
    new_data = {}
    for docx_path in docx_files:
        print(f"\n  Parsing: {os.path.basename(docx_path)} ...")
        parsed = parse_doc(docx_path)
        for wk, data in parsed.items():
            if wk in existing and not args.force:
                print(f"    W{wk}: already present, skipping (use --force to overwrite)")
            else:
                new_data[wk] = data
        print(f"    Got {len(parsed)} weeks: {sorted(parsed.keys(), key=int)}")

    # ── Report ───────────────────────────────────────────────────────────
    if not new_data:
        print("\n✅ No new weeks to add — lessonPlans.json is up to date.")
        # Still print LP_AVAILABLE status
        print(f"\n📊 Current coverage: {len(existing)} weeks — {sorted(existing.keys(), key=int)}")
        return

    print(f"\n🆕 New weeks to add: {sorted(new_data.keys(), key=int)}")

    if args.dry_run:
        print("\n[DRY RUN] No files written.")
        print_lp_available_instructions(new_data.keys(), {**existing, **new_data})
        return

    # ── Merge and write ──────────────────────────────────────────────────
    merged = {**existing, **new_data}
    # Sort keys numerically
    sorted_merged = {k: merged[k] for k in sorted(merged.keys(), key=int)}

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    with open(OUT_PATH, "w", encoding="utf-8") as f:
        json.dump(sorted_merged, f, ensure_ascii=False, indent=2)

    size_kb = os.path.getsize(OUT_PATH) / 1024
    print(f"\n✅ Saved to {OUT_PATH}")
    print(f"   {len(sorted_merged)} total weeks · {size_kb:.0f} KB")

    # ── LP_AVAILABLE update instructions ────────────────────────────────
    print_lp_available_instructions(new_data.keys(), sorted_merged)


if __name__ == "__main__":
    main()
